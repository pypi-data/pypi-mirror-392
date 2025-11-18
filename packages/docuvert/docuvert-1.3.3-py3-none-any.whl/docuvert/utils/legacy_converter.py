"""
Utility functions for converting legacy formats (.doc, .xls) to modern equivalents.
"""

import os
import tempfile
import shutil
from typing import Optional, Tuple
import subprocess
import sys


class LegacyFormatConverter:
    """Handles conversion of legacy Office formats to modern equivalents."""
    
    @staticmethod
    def convert_doc_to_docx(doc_path: str, temp_dir: Optional[str] = None) -> str:
        """
        Convert .doc file to .docx format.
        
        Args:
            doc_path: Path to the .doc file
            temp_dir: Optional temp directory, creates one if not provided
            
        Returns:
            Path to the converted .docx file
            
        Raises:
            Exception: If conversion fails
        """
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Input file not found: {doc_path}")
            
        if temp_dir is None:
            temp_dir = tempfile.mkdtemp()
            
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(temp_dir, f"{base_name}.docx")
        
        # Try multiple conversion methods in order of preference
        conversion_methods = [
            LegacyFormatConverter._convert_with_libreoffice,
            LegacyFormatConverter._convert_with_pandoc,
            LegacyFormatConverter._convert_with_python_docx
        ]
        
        for method in conversion_methods:
            try:
                method(doc_path, docx_path)
                if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
                    return docx_path
            except Exception as e:
                print(f"Warning: {method.__name__} failed: {e}", file=sys.stderr)
                continue
                
        raise Exception("All conversion methods failed for .doc to .docx conversion")
    
    @staticmethod
    def convert_xls_to_xlsx(xls_path: str, temp_dir: Optional[str] = None) -> str:
        """
        Convert .xls file to .xlsx format.
        
        Args:
            xls_path: Path to the .xls file
            temp_dir: Optional temp directory, creates one if not provided
            
        Returns:
            Path to the converted .xlsx file
            
        Raises:
            Exception: If conversion fails
        """
        if not os.path.exists(xls_path):
            raise FileNotFoundError(f"Input file not found: {xls_path}")
            
        if temp_dir is None:
            temp_dir = tempfile.mkdtemp()
            
        base_name = os.path.splitext(os.path.basename(xls_path))[0]
        xlsx_path = os.path.join(temp_dir, f"{base_name}.xlsx")
        
        # Try multiple conversion methods
        conversion_methods = [
            LegacyFormatConverter._convert_xls_with_libreoffice,
            LegacyFormatConverter._convert_xls_with_pandas
        ]
        
        for method in conversion_methods:
            try:
                method(xls_path, xlsx_path)
                if os.path.exists(xlsx_path) and os.path.getsize(xlsx_path) > 0:
                    return xlsx_path
            except Exception as e:
                print(f"Warning: {method.__name__} failed: {e}", file=sys.stderr)
                continue
                
        raise Exception("All conversion methods failed for .xls to .xlsx conversion")
    
    @staticmethod
    def _convert_with_libreoffice(input_path: str, output_path: str):
        """Convert using LibreOffice command line."""
        output_dir = os.path.dirname(output_path)
        
        # Try common LibreOffice executable names
        libreoffice_commands = ['libreoffice', 'soffice', '/Applications/LibreOffice.app/Contents/MacOS/soffice']
        
        for cmd in libreoffice_commands:
            try:
                result = subprocess.run([
                    cmd, '--headless', '--convert-to', 'docx',
                    '--outdir', output_dir, input_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    # LibreOffice creates file with same name but .docx extension
                    base_name = os.path.splitext(os.path.basename(input_path))[0]
                    created_file = os.path.join(output_dir, f"{base_name}.docx")
                    
                    if os.path.exists(created_file) and created_file != output_path:
                        shutil.move(created_file, output_path)
                    return
                    
            except (subprocess.TimeoutExpired, FileNotFoundError):
                continue
                
        raise Exception("LibreOffice not found or conversion failed")
    
    @staticmethod
    def _convert_with_pandoc(input_path: str, output_path: str):
        """Convert using Pandoc."""
        try:
            result = subprocess.run([
                'pandoc', input_path, '-o', output_path
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"Pandoc failed: {result.stderr}")
                
        except FileNotFoundError:
            raise Exception("Pandoc not found")
    
    @staticmethod 
    def _convert_with_python_docx(input_path: str, output_path: str):
        """Attempt basic conversion using python-docx (limited functionality)."""
        try:
            # This is a fallback that creates an empty docx with error message
            from docx import Document
            
            doc = Document()
            doc.add_paragraph(f"Error: Unable to convert {input_path}")
            doc.add_paragraph("Legacy .doc format requires LibreOffice or Pandoc for proper conversion.")
            doc.save(output_path)
            
        except ImportError:
            raise Exception("python-docx not available")
    
    @staticmethod
    def _convert_xls_with_libreoffice(input_path: str, output_path: str):
        """Convert XLS using LibreOffice command line."""
        output_dir = os.path.dirname(output_path)
        
        libreoffice_commands = ['libreoffice', 'soffice', '/Applications/LibreOffice.app/Contents/MacOS/soffice']
        
        for cmd in libreoffice_commands:
            try:
                result = subprocess.run([
                    cmd, '--headless', '--convert-to', 'xlsx',
                    '--outdir', output_dir, input_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    base_name = os.path.splitext(os.path.basename(input_path))[0] 
                    created_file = os.path.join(output_dir, f"{base_name}.xlsx")
                    
                    if os.path.exists(created_file) and created_file != output_path:
                        shutil.move(created_file, output_path)
                    return
                    
            except (subprocess.TimeoutExpired, FileNotFoundError):
                continue
                
        raise Exception("LibreOffice not found or XLS conversion failed")
    
    @staticmethod
    def _convert_xls_with_pandas(input_path: str, output_path: str):
        """Convert XLS using pandas (requires xlrd)."""
        try:
            import pandas as pd
            
            # Read XLS file
            df = pd.read_excel(input_path, engine='xlrd')
            
            # Save as XLSX
            df.to_excel(output_path, index=False, engine='openpyxl')
            
        except ImportError:
            raise Exception("pandas or xlrd not available for XLS conversion")
        except Exception as e:
            raise Exception(f"pandas XLS conversion failed: {e}")
    
    @staticmethod
    def cleanup_temp_file(file_path: str):
        """Clean up temporary converted file."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                # Also remove parent temp directory if empty
                parent_dir = os.path.dirname(file_path)
                if os.path.exists(parent_dir) and not os.listdir(parent_dir):
                    os.rmdir(parent_dir)
        except Exception:
            pass  # Ignore cleanup errors