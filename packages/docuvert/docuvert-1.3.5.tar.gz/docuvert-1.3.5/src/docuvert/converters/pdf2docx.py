
import sys
import os
# Add the parent directory to the path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from core.ast import ASTNode
from core.exceptions import ConversionError, DependencyError
from utils.ocr_utils import OCRExtractor
from pdf2docx import Converter

try:
    from docx import Document
    from docx.shared import Inches
except ImportError as e:
    raise DependencyError(
        "python-docx is required for DOCX generation",
        missing_dependency="python-docx"
    ) from e

class Pdf2DocxConverter:
    """
    Converts a PDF file to a DOCX file.
    """
    def parse_pdf2ast(self, input_path: str) -> ASTNode:
        """
        Parses a PDF file and converts it to an AST.

        Args:
            input_path (str): The path to the input PDF file.

        Returns:
            ASTNode: The root of the generated AST.
        """
        # TODO: Implement the logic to parse the PDF and build an AST.
        # This will involve using a library like pdfminer.six or PyMuPDF.
        print(f"Parsing PDF at {input_path} and converting to AST.")
        return ASTNode(type="root")

    def ast2docx(self, ast_root: ASTNode, output_path: str):
        """
        Converts an AST to a DOCX file.

        Args:
            ast_root (ASTNode): The root of the AST.
            output_path (str): The path to the output DOCX file.
        """
        # TODO: Implement the logic to convert the AST to a DOCX document.
        # This will involve using a library like python-docx.
        print(f"Converting AST to DOCX at {output_path}")

    def convert(self, input_path: str, output_path: str):
        """
        Converts a PDF file to a DOCX file with OCR fallback.

        Args:
            input_path (str): The path to the input PDF file.
            output_path (str): The path to the output DOCX file.
        """
        if not os.path.exists(input_path):
            raise ConversionError(f"Input file not found: {input_path}")

        try:
            print(f"🔄 Converting PDF to DOCX: {input_path}")

            # Try pdf2docx first (preserves formatting for text-based PDFs)
            try:
                cv = Converter(input_path)
                cv.convert(output_path)
                cv.close()
                print(f"✅ Successfully converted '{input_path}' to '{output_path}' using pdf2docx")
                return

            except Exception as e:
                print(f"⚠️ pdf2docx conversion failed: {e}")
                print("🔍 Falling back to OCR text extraction...")

                # Fallback: Extract text with OCR and create simple DOCX
                text = OCRExtractor.extract_text_from_pdf(input_path, use_ocr=True)

                if not text.strip():
                    text = "No text could be extracted from this PDF."

                # Create DOCX document with extracted text
                doc = Document()
                doc.add_heading('Extracted PDF Content', 0)

                # Split text into pages and add as separate sections
                pages = text.split('--- Page')
                for i, page_content in enumerate(pages):
                    if page_content.strip():
                        if i > 0:  # Skip empty first element before first page marker
                            doc.add_heading(f'Page {i}', level=1)

                        # Add page content as paragraphs
                        paragraphs = page_content.split('\n\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())

                doc.save(output_path)
                print(f"✅ Successfully converted '{input_path}' to '{output_path}' using OCR")

        except Exception as e:
            raise ConversionError(
                f"PDF to DOCX conversion failed: {e}",
                source_format="pdf",
                target_format="docx",
                suggestions=[
                    "For scanned PDFs, install OCR support: pip install pytesseract",
                    "Install Tesseract: brew install tesseract (macOS) or apt-get install tesseract-ocr (Linux)",
                    "Ensure PDF is not corrupted or password-protected"
                ]
            )
