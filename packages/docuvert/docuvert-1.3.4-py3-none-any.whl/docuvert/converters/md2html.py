"""
Markdown to HTML converter with comprehensive formatting support.
"""

import sys
import os
from typing import Optional, Dict, Any

# Add the parent directory to the path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from core.ast import ASTNode, NodeType
from core.exceptions import ConversionError, DependencyError


class Md2HtmlConverter:
    """Converts Markdown files to HTML format."""
    
    def parse_md2ast(self, input_path: str) -> ASTNode:
        """Parse Markdown to AST representation."""
        # This would need a proper Markdown parser implementation
        # For now, we'll use the existing txt2md logic as a starting point
        from .txt2md import Txt2MdConverter
        
        # Create a basic AST from markdown content
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return self._parse_markdown_content(content)
    
    def _parse_markdown_content(self, content: str) -> ASTNode:
        """Parse markdown content into AST nodes."""
        from core.ast import create_document, create_paragraph, create_heading
        
        doc = create_document()
        lines = content.split('\\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('# ').strip()
                doc.add_child(create_heading(header_text, level))
                i += 1
            
            # Code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                code_content = '\\n'.join(code_lines)
                doc.add_child(ASTNode(NodeType.CODE_BLOCK, content=code_content))
                i += 1  # Skip closing ```
            
            # Lists
            elif line.startswith(('- ', '* ', '+ ')):
                list_node = ASTNode(NodeType.LIST_UNORDERED)
                while i < len(lines) and lines[i].strip().startswith(('- ', '* ', '+ ')):
                    item_text = lines[i].strip()[2:].strip()
                    list_node.add_child(ASTNode(NodeType.LIST_ITEM, content=item_text))
                    i += 1
                doc.add_child(list_node)
            
            # Numbered lists
            elif line and line[0].isdigit() and '. ' in line:
                list_node = ASTNode(NodeType.LIST_ORDERED)
                while i < len(lines) and lines[i].strip() and lines[i].strip()[0].isdigit():
                    if '. ' in lines[i]:
                        item_text = lines[i].split('. ', 1)[1].strip()
                        list_node.add_child(ASTNode(NodeType.LIST_ITEM, content=item_text))
                    i += 1
                doc.add_child(list_node)
            
            # Regular paragraphs
            else:
                doc.add_child(create_paragraph(line))
                i += 1
        
        return doc
    
    def ast2html(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to HTML format."""
        from core.mapper import ast2html
        ast2html(ast_root, output_path)
    
    def convert(self, input_path: str, output_path: str) -> None:
        """Convert Markdown to HTML."""
        
        try:
            # Method 1: Use markdown library if available
            self._convert_with_markdown_lib(input_path, output_path)
            print(f"Successfully converted '{input_path}' to '{output_path}' (markdown library)")
            
        except Exception as lib_error:
            try:
                # Method 2: AST-based conversion
                ast_root = self.parse_md2ast(input_path)
                self.ast2html(ast_root, output_path)
                print(f"Successfully converted '{input_path}' to '{output_path}' (AST method)")
                
            except Exception as ast_error:
                try:
                    # Method 3: Pandoc fallback
                    self._convert_with_pandoc(input_path, output_path)
                    print(f"Successfully converted '{input_path}' to '{output_path}' (pandoc fallback)")
                    
                except Exception as pandoc_error:
                    raise ConversionError(
                        f"Markdown to HTML conversion failed. Lib: {lib_error}, AST: {ast_error}, Pandoc: {pandoc_error}",
                        source_format="md",
                        target_format="html"
                    )
    
    def _convert_with_markdown_lib(self, input_path: str, output_path: str) -> None:
        """Convert using the Python markdown library."""
        try:
            import markdown
            from markdown.extensions import codehilite, tables, toc
        except ImportError as e:
            raise DependencyError(
                "markdown library is required for Markdown conversion",
                missing_dependency="markdown"
            ) from e
        
        with open(input_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Configure markdown with extensions
        md = markdown.Markdown(extensions=[
            'codehilite',
            'tables', 
            'toc',
            'fenced_code',
            'attr_list',
            'def_list',
            'footnotes',
            'md_in_html'
        ])
        
        html_content = md.convert(markdown_content)
        
        # Wrap in complete HTML document
        full_html = f\"\"\"<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 2rem; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 1rem; border-radius: 5px; overflow-x: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        table, th, td {{ border: 1px solid #ddd; }}
        th, td {{ padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        blockquote {{ border-left: 4px solid #ccc; margin: 1rem 0; padding-left: 1rem; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>\"\"\"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
    
    def _convert_with_pandoc(self, input_path: str, output_path: str) -> None:
        """Fallback conversion using pandoc."""
        try:
            import pypandoc
        except ImportError as e:
            raise DependencyError(
                "pypandoc is required for pandoc fallback conversion",
                missing_dependency="pypandoc"
            ) from e
        
        pypandoc.convert_file(input_path, 'html', outputfile=output_path)


class Pdf2HtmlConverter:
    """Converts PDF files to HTML format."""
    
    def parse_pdf2ast(self, input_path: str) -> ASTNode:
        """Parse PDF to AST representation."""
        # This is a placeholder - would need proper PDF parsing
        from core.ast import create_document, create_paragraph
        
        doc = create_document()
        
        try:
            import PyMuPDF  # pymupdf
            
            pdf_doc = PyMuPDF.open(input_path)
            
            for page_num in range(pdf_doc.page_count):
                page = pdf_doc[page_num]
                text = page.get_text()
                
                # Split into paragraphs (simple approach)
                paragraphs = text.split('\\n\\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        doc.add_child(create_paragraph(para_text))
            
            pdf_doc.close()
            
        except ImportError:
            # Fallback without PDF parsing
            doc.add_child(create_paragraph("PDF parsing requires PyMuPDF library"))
        
        return doc
    
    def ast2html(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to HTML format."""
        from core.mapper import ast2html
        ast2html(ast_root, output_path)
    
    def convert(self, input_path: str, output_path: str) -> None:
        """Convert PDF to HTML."""
        
        try:
            # AST-based conversion
            ast_root = self.parse_pdf2ast(input_path)
            self.ast2html(ast_root, output_path)
            print(f"Successfully converted '{input_path}' to '{output_path}'")
            
        except Exception as e:
            raise ConversionError(
                f"PDF to HTML conversion failed: {e}",
                source_format="pdf",
                target_format="html",
                suggestions=[
                    "Install PyMuPDF: pip install pymupdf",
                    "Check that the PDF file is not corrupted or password-protected"
                ]
            )


class Docx2HtmlConverter:
    """Converts DOCX files to HTML format."""
    
    def parse_docx2ast(self, input_path: str) -> ASTNode:
        """Parse DOCX to AST representation."""
        try:
            from docx import Document
        except ImportError as e:
            raise DependencyError(
                "python-docx is required for DOCX conversion",
                missing_dependency="python-docx"
            ) from e
        
        from core.ast import create_document, create_paragraph, create_heading
        
        doc_ast = create_document()
        docx_doc = Document(input_path)
        
        for paragraph in docx_doc.paragraphs:
            if not paragraph.text.strip():
                continue
                
            # Check if it's a heading
            if paragraph.style.name.startswith('Heading'):
                try:
                    level = int(paragraph.style.name.split(' ')[1])
                    doc_ast.add_child(create_heading(paragraph.text, level))
                except (ValueError, IndexError):
                    doc_ast.add_child(create_paragraph(paragraph.text))
            else:
                doc_ast.add_child(create_paragraph(paragraph.text))
        
        # Process tables
        for table in docx_doc.tables:
            table_node = ASTNode(NodeType.TABLE)
            
            for row in table.rows:
                row_node = ASTNode(NodeType.TABLE_ROW)
                
                for cell in row.cells:
                    cell_node = ASTNode(NodeType.TABLE_CELL, content=cell.text)
                    row_node.add_child(cell_node)
                
                table_node.add_child(row_node)
            
            doc_ast.add_child(table_node)
        
        return doc_ast
    
    def ast2html(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to HTML format."""
        from core.mapper import ast2html
        ast2html(ast_root, output_path)
    
    def convert(self, input_path: str, output_path: str) -> None:
        """Convert DOCX to HTML."""
        
        try:
            # AST-based conversion
            ast_root = self.parse_docx2ast(input_path)
            self.ast2html(ast_root, output_path)
            print(f"Successfully converted '{input_path}' to '{output_path}'")
            
        except Exception as ast_error:
            try:
                # Pandoc fallback
                self._convert_with_pandoc(input_path, output_path)
                print(f"Successfully converted '{input_path}' to '{output_path}' (pandoc fallback)")
                
            except Exception as pandoc_error:
                raise ConversionError(
                    f"DOCX to HTML conversion failed. AST: {ast_error}, Pandoc: {pandoc_error}",
                    source_format="docx",
                    target_format="html"
                )
    
    def _convert_with_pandoc(self, input_path: str, output_path: str) -> None:
        """Fallback conversion using pandoc."""
        try:
            import pypandoc
        except ImportError as e:
            raise DependencyError(
                "pypandoc is required for pandoc fallback conversion",
                missing_dependency="pypandoc"
            ) from e
        
        pypandoc.convert_file(input_path, 'html', outputfile=output_path)