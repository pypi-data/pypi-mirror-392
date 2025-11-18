
"""
Universal AST mapping system for format conversion.

This module provides comprehensive AST-to-format conversion functions
that preserve document structure, formatting, and metadata across formats.
"""

from typing import Dict, Any, List, Optional
import os
from io import StringIO

from .ast import ASTNode, NodeType, StyleInfo, DocumentMetadata
from .exceptions import ConversionError, DependencyError


class ASTMapper:
    """Universal AST to format mapper."""
    
    def __init__(self):
        self.format_handlers = {
            'docx': self.ast_to_docx,
            'md': self.ast_to_markdown,
            'html': self.ast_to_html,
            'tex': self.ast_to_latex,
            'txt': self.ast_to_text,
            'rtf': self.ast_to_rtf,
            'json': self.ast_to_json,
        }
    
    def convert_ast(self, ast_root: ASTNode, output_path: str, format_type: str) -> None:
        """Convert AST to specified format."""
        
        if format_type not in self.format_handlers:
            raise ConversionError(
                f"Unsupported output format: {format_type}",
                error_code="UNSUPPORTED_OUTPUT_FORMAT"
            )
        
        handler = self.format_handlers[format_type]
        handler(ast_root, output_path)
    
    def ast_to_docx(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to DOCX format using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError as e:
            raise DependencyError(
                "python-docx is required for DOCX conversion",
                missing_dependency="python-docx"
            ) from e
        
        doc = Document()
        
        # Set document metadata if available
        if ast_root.metadata:
            core_props = doc.core_properties
            if ast_root.metadata.title:
                core_props.title = ast_root.metadata.title
            if ast_root.metadata.author:
                core_props.author = ast_root.metadata.author
            if ast_root.metadata.subject:
                core_props.subject = ast_root.metadata.subject
            if ast_root.metadata.keywords:
                core_props.keywords = ', '.join(ast_root.metadata.keywords)
        
        # Process AST nodes
        self._process_docx_node(doc, ast_root)
        
        doc.save(output_path)
    
    def _process_docx_node(self, doc, node: ASTNode, parent_element=None):
        """Process a single AST node for DOCX conversion."""
        
        if node.type == NodeType.DOCUMENT:
            for child in node.children:
                self._process_docx_node(doc, child)
                
        elif node.type == NodeType.HEADING:
            level = node.attributes.get('level', 1)
            heading = doc.add_heading(node.content or '', level=min(level, 9))
            
        elif node.type == NodeType.PARAGRAPH:
            para = doc.add_paragraph(node.content or '')
            if node.styles and node.styles.text_align:
                align_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER, 
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if node.styles.text_align in align_map:
                    para.alignment = align_map[node.styles.text_align]
            
            # Process inline formatting children
            for child in node.children:
                self._process_docx_inline(para, child)
                
        elif node.type == NodeType.TABLE:
            rows = len([child for child in node.children if child.type == NodeType.TABLE_ROW])
            if rows > 0:
                first_row = next(child for child in node.children if child.type == NodeType.TABLE_ROW)
                cols = len([cell for cell in first_row.children if cell.type in [NodeType.TABLE_CELL, NodeType.TABLE_HEADER]])
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for row_idx, row_node in enumerate([child for child in node.children if child.type == NodeType.TABLE_ROW]):
                    for col_idx, cell_node in enumerate([cell for cell in row_node.children if cell.type in [NodeType.TABLE_CELL, NodeType.TABLE_HEADER]]):
                        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_node.content or ''
        
        elif node.type == NodeType.LIST_ORDERED:
            for item in node.children:
                if item.type == NodeType.LIST_ITEM:
                    doc.add_paragraph(item.content or '', style='List Number')
                    
        elif node.type == NodeType.LIST_UNORDERED:
            for item in node.children:
                if item.type == NodeType.LIST_ITEM:
                    doc.add_paragraph(item.content or '', style='List Bullet')
                    
        elif node.type == NodeType.CODE_BLOCK:
            para = doc.add_paragraph(node.content or '')
            # Set monospace font for code
            for run in para.runs:
                run.font.name = 'Courier New'
                
        # Process remaining children
        if node.type not in [NodeType.PARAGRAPH, NodeType.TABLE, NodeType.LIST_ORDERED, NodeType.LIST_UNORDERED]:
            for child in node.children:
                self._process_docx_node(doc, child)
    
    def _process_docx_inline(self, para, node: ASTNode):
        """Process inline formatting nodes for DOCX."""
        run = para.add_run(node.content or '')
        
        if node.type == NodeType.BOLD:
            run.bold = True
        elif node.type == NodeType.ITALIC:
            run.italic = True
        elif node.type == NodeType.UNDERLINE:
            run.underline = True
        elif node.type == NodeType.CODE_INLINE:
            run.font.name = 'Courier New'
            
    def ast_to_markdown(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to Markdown format."""
        
        output = StringIO()
        self._process_markdown_node(output, ast_root)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output.getvalue())
    
    def _process_markdown_node(self, output: StringIO, node: ASTNode, level: int = 0):
        """Process a single AST node for Markdown conversion."""
        
        if node.type == NodeType.DOCUMENT:
            for child in node.children:
                self._process_markdown_node(output, child, level)
                
        elif node.type == NodeType.HEADING:
            level = node.attributes.get('level', 1)
            output.write('#' * level + ' ' + (node.content or '') + '\n\n')
            
        elif node.type == NodeType.PARAGRAPH:
            content = node.content or ''
            # Process inline formatting
            for child in node.children:
                content += self._process_markdown_inline(child)
            output.write(content + '\n\n')
            
        elif node.type == NodeType.TABLE:
            rows = [child for child in node.children if child.type == NodeType.TABLE_ROW]
            if not rows:
                return
                
            # Process table rows
            for row_idx, row_node in enumerate(rows):
                output.write('|')
                cells = [cell for cell in row_node.children if cell.type in [NodeType.TABLE_CELL, NodeType.TABLE_HEADER]]
                
                for cell in cells:
                    output.write(' ' + (cell.content or '') + ' |')
                output.write('\n')
                
                # Add header separator after first row
                if row_idx == 0:
                    output.write('|')
                    for _ in cells:
                        output.write('---|')
                    output.write('\n')
            output.write('\n')
            
        elif node.type == NodeType.LIST_ORDERED:
            for idx, item in enumerate([child for child in node.children if child.type == NodeType.LIST_ITEM], 1):
                output.write(f'{idx}. {item.content or ""}\n')
            output.write('\n')
            
        elif node.type == NodeType.LIST_UNORDERED:
            for item in [child for child in node.children if child.type == NodeType.LIST_ITEM]:
                output.write(f'- {item.content or ""}\n')
            output.write('\n')
            
        elif node.type == NodeType.CODE_BLOCK:
            output.write('```\n')
            output.write(node.content or '')
            if not (node.content or '').endswith('\n'):
                output.write('\n')
            output.write('```\n\n')
            
        elif node.type == NodeType.LINE_BREAK:
            output.write('\n')
            
        # Process remaining children
        for child in node.children:
            if child.type not in [NodeType.TEXT, NodeType.BOLD, NodeType.ITALIC, NodeType.CODE_INLINE]:
                self._process_markdown_node(output, child, level)
    
    def _process_markdown_inline(self, node: ASTNode) -> str:
        """Process inline formatting for Markdown."""
        content = node.content or ''
        
        if node.type == NodeType.BOLD:
            return f'**{content}**'
        elif node.type == NodeType.ITALIC:
            return f'*{content}*'
        elif node.type == NodeType.CODE_INLINE:
            return f'`{content}`'
        elif node.type == NodeType.LINK:
            url = node.attributes.get('href', '')
            return f'[{content}]({url})'
        else:
            return content
    
    def ast_to_html(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to HTML format."""
        
        output = StringIO()
        
        # HTML document structure
        output.write('<!DOCTYPE html>\n<html>\n<head>\n')
        
        # Add metadata
        if ast_root.metadata:
            if ast_root.metadata.title:
                output.write(f'<title>{self._html_escape(ast_root.metadata.title)}</title>\n')
            if ast_root.metadata.author:
                output.write(f'<meta name="author" content="{self._html_escape(ast_root.metadata.author)}">\n')
        
        output.write('<meta charset="utf-8">\n</head>\n<body>\n')
        
        # Process body content
        for child in ast_root.children:
            self._process_html_node(output, child)
            
        output.write('</body>\n</html>\n')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output.getvalue())
    
    def _process_html_node(self, output: StringIO, node: ASTNode):
        """Process a single AST node for HTML conversion."""
        
        if node.type == NodeType.HEADING:
            level = min(node.attributes.get('level', 1), 6)
            output.write(f'<h{level}>{self._html_escape(node.content or "")}</h{level}>\n')
            
        elif node.type == NodeType.PARAGRAPH:
            output.write('<p>')
            if node.content:
                output.write(self._html_escape(node.content))
            for child in node.children:
                self._process_html_inline(output, child)
            output.write('</p>\n')
            
        elif node.type == NodeType.TABLE:
            output.write('<table>\n')
            rows = [child for child in node.children if child.type == NodeType.TABLE_ROW]
            
            for row_idx, row_node in enumerate(rows):
                output.write('<tr>')
                cells = [cell for cell in row_node.children if cell.type in [NodeType.TABLE_CELL, NodeType.TABLE_HEADER]]
                
                for cell in cells:
                    tag = 'th' if cell.type == NodeType.TABLE_HEADER else 'td'
                    output.write(f'<{tag}>{self._html_escape(cell.content or "")}</{tag}>')
                output.write('</tr>\n')
                
            output.write('</table>\n')
            
        elif node.type in [NodeType.LIST_ORDERED, NodeType.LIST_UNORDERED]:
            tag = 'ol' if node.type == NodeType.LIST_ORDERED else 'ul'
            output.write(f'<{tag}>\n')
            
            for item in [child for child in node.children if child.type == NodeType.LIST_ITEM]:
                output.write(f'<li>{self._html_escape(item.content or "")}</li>\n')
                
            output.write(f'</{tag}>\n')
            
        elif node.type == NodeType.CODE_BLOCK:
            output.write('<pre><code>')
            output.write(self._html_escape(node.content or ''))
            output.write('</code></pre>\n')
            
        # Process remaining children
        for child in node.children:
            if child.type not in [NodeType.TEXT, NodeType.BOLD, NodeType.ITALIC, NodeType.CODE_INLINE, NodeType.LINK]:
                self._process_html_node(output, child)
    
    def _process_html_inline(self, output: StringIO, node: ASTNode):
        """Process inline formatting for HTML."""
        content = self._html_escape(node.content or '')
        
        if node.type == NodeType.BOLD:
            output.write(f'<strong>{content}</strong>')
        elif node.type == NodeType.ITALIC:
            output.write(f'<em>{content}</em>')
        elif node.type == NodeType.CODE_INLINE:
            output.write(f'<code>{content}</code>')
        elif node.type == NodeType.LINK:
            url = self._html_escape(node.attributes.get('href', ''))
            output.write(f'<a href="{url}">{content}</a>')
        else:
            output.write(content)
    
    def _html_escape(self, text: str) -> str:
        """Escape HTML special characters."""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#x27;')
    
    def ast_to_latex(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to LaTeX format."""
        
        output = StringIO()
        
        # LaTeX document structure
        output.write('\\documentclass{article}\n')
        output.write('\\usepackage[utf8]{inputenc}\n')
        output.write('\\usepackage{graphicx}\n')
        output.write('\\usepackage{booktabs}\n')
        output.write('\\usepackage{listings}\n')
        
        # Add metadata
        if ast_root.metadata:
            if ast_root.metadata.title:
                output.write(f'\\title{{{self._latex_escape(ast_root.metadata.title)}}}\n')
            if ast_root.metadata.author:
                output.write(f'\\author{{{self._latex_escape(ast_root.metadata.author)}}}\n')
                
        output.write('\\begin{document}\n')
        
        if ast_root.metadata and ast_root.metadata.title:
            output.write('\\maketitle\n')
            
        # Process content
        for child in ast_root.children:
            self._process_latex_node(output, child)
            
        output.write('\\end{document}\n')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output.getvalue())
    
    def _process_latex_node(self, output: StringIO, node: ASTNode):
        """Process a single AST node for LaTeX conversion."""
        
        if node.type == NodeType.HEADING:
            level = node.attributes.get('level', 1)
            commands = {1: 'section', 2: 'subsection', 3: 'subsubsection'}
            cmd = commands.get(level, 'paragraph')
            output.write(f'\\{cmd}{{{self._latex_escape(node.content or "")}}}\n\n')
            
        elif node.type == NodeType.PARAGRAPH:
            content = self._latex_escape(node.content or '')
            for child in node.children:
                content += self._process_latex_inline(child)
            output.write(content + '\n\n')
            
        elif node.type == NodeType.CODE_BLOCK:
            output.write('\\begin{lstlisting}\n')
            output.write(node.content or '')
            if not (node.content or '').endswith('\n'):
                output.write('\n')
            output.write('\\end{lstlisting}\n\n')
            
        # Process children
        for child in node.children:
            if child.type not in [NodeType.TEXT, NodeType.BOLD, NodeType.ITALIC, NodeType.CODE_INLINE]:
                self._process_latex_node(output, child)
    
    def _process_latex_inline(self, node: ASTNode) -> str:
        """Process inline formatting for LaTeX."""
        content = self._latex_escape(node.content or '')
        
        if node.type == NodeType.BOLD:
            return f'\\textbf{{{content}}}'
        elif node.type == NodeType.ITALIC:
            return f'\\textit{{{content}}}'
        elif node.type == NodeType.CODE_INLINE:
            return f'\\texttt{{{content}}}'
        else:
            return content
    
    def _latex_escape(self, text: str) -> str:
        """Escape LaTeX special characters."""
        escapes = {
            '&': '\\&', '%': '\\%', '$': '\\$', '#': '\\#', 
            '^': '\\textasciicircum{}', '_': '\\_', '{': '\\{', '}': '\\}',
            '~': '\\textasciitilde{}', '\\': '\\textbackslash{}'
        }
        for char, escape in escapes.items():
            text = text.replace(char, escape)
        return text
    
    def ast_to_text(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to plain text format."""
        
        output = StringIO()
        self._process_text_node(output, ast_root)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output.getvalue())
    
    def _process_text_node(self, output: StringIO, node: ASTNode):
        """Process a single AST node for text conversion."""
        
        if node.type == NodeType.HEADING:
            level = node.attributes.get('level', 1)
            output.write('\n' + '=' * (80 // level) + '\n')
            output.write((node.content or '').upper() + '\n')
            output.write('=' * (80 // level) + '\n\n')
            
        elif node.type == NodeType.PARAGRAPH:
            content = node.content or ''
            for child in node.children:
                content += child.get_text_content()
            output.write(content + '\n\n')
            
        elif node.type == NodeType.CODE_BLOCK:
            output.write('Code:\n')
            output.write('-' * 40 + '\n')
            output.write(node.content or '')
            if not (node.content or '').endswith('\n'):
                output.write('\n')
            output.write('-' * 40 + '\n\n')
            
        # Process children
        for child in node.children:
            self._process_text_node(output, child)
    
    def ast_to_rtf(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to RTF format (basic implementation)."""
        
        output = StringIO()
        
        # RTF header
        output.write('{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\n')
        
        # Process content
        self._process_rtf_node(output, ast_root)
        
        # RTF footer
        output.write('}\n')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output.getvalue())
    
    def _process_rtf_node(self, output: StringIO, node: ASTNode):
        """Process a single AST node for RTF conversion."""
        
        if node.type == NodeType.HEADING:
            level = node.attributes.get('level', 1)
            size = max(24 - (level * 2), 12)  # Larger font for headings
            output.write(f'\\fs{size}\\b {node.content or ""}\\b0\\fs24\\par\n')
            
        elif node.type == NodeType.PARAGRAPH:
            content = node.content or ''
            for child in node.children:
                content += self._process_rtf_inline(child)
            output.write(content + '\\par\n')
            
        # Process children
        for child in node.children:
            if child.type not in [NodeType.TEXT, NodeType.BOLD, NodeType.ITALIC]:
                self._process_rtf_node(output, child)
    
    def _process_rtf_inline(self, node: ASTNode) -> str:
        """Process inline formatting for RTF."""
        content = node.content or ''
        
        if node.type == NodeType.BOLD:
            return f'\\b {content}\\b0 '
        elif node.type == NodeType.ITALIC:
            return f'\\i {content}\\i0 '
        else:
            return content
    
    def ast_to_json(self, ast_root: ASTNode, output_path: str) -> None:
        """Convert AST to JSON format (for debugging/interchange)."""
        import json
        
        def node_to_dict(node: ASTNode) -> Dict[str, Any]:
            return {
                'type': node.type.value,
                'content': node.content,
                'children': [node_to_dict(child) for child in node.children],
                'styles': node.styles.__dict__ if node.styles else {},
                'attributes': node.attributes,
                'metadata': node.metadata.__dict__ if node.metadata else None
            }
        
        data = node_to_dict(ast_root)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)


# Global mapper instance
_mapper = ASTMapper()

# Legacy function wrappers for backward compatibility
def ast2docx(ast_root: ASTNode, output_path: str) -> None:
    """Convert AST to DOCX format."""
    _mapper.ast_to_docx(ast_root, output_path)

def ast2md(ast_root: ASTNode, output_path: str) -> None:
    """Convert AST to Markdown format."""
    _mapper.ast_to_markdown(ast_root, output_path)

def ast2tex(ast_root: ASTNode, output_path: str) -> None:
    """Convert AST to LaTeX format."""
    _mapper.ast_to_latex(ast_root, output_path)

def ast2html(ast_root: ASTNode, output_path: str) -> None:
    """Convert AST to HTML format."""
    _mapper.ast_to_html(ast_root, output_path)

def ast2txt(ast_root: ASTNode, output_path: str) -> None:
    """Convert AST to plain text format."""
    _mapper.ast_to_text(ast_root, output_path)

def get_mapper() -> ASTMapper:
    """Get the global AST mapper instance."""
    return _mapper
