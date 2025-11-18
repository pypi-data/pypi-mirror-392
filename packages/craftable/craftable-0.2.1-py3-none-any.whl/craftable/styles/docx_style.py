"""DOCX style: writes a Word document table using python-docx.

Optional dependency: python-docx (group: docx)
"""

from typing import Any, IO, TYPE_CHECKING
from pathlib import Path

from .table_style import TableStyle

if TYPE_CHECKING:  # avoid circular import at runtime
    from ..craftable import ColDefList


class DocxStyle(TableStyle):
    def __init__(self, table_style: str | None = "Light List"):
        super().__init__()
        self.terminal_style = False
        self.string_output = False  # Binary format
        self.table_style_name = table_style

    ###############################################################################
    # write_table
    ###############################################################################

    def write_table(
        self,
        value_rows: list[list[Any]],
        header_row: list[str] | None,
        col_defs: "ColDefList",
        header_defs: "ColDefList | None",
        file: str | Path | IO[bytes],
    ) -> None:
        """Write table to DOCX file format."""
        try:
            from docx import Document  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.shared import Inches  # type: ignore
        except Exception as e:  # pragma: no cover - only hit without optional dep
            raise ImportError(
                "python-docx is required for DocxStyle; install group 'docx'"
            ) from e

        doc = Document()

        # Create table: rows = header + data
        n_rows = len(value_rows) + (1 if header_row else 0)
        n_cols = len(col_defs)
        table = doc.add_table(
            rows=n_rows if n_rows else 1, cols=n_cols if n_cols else 1
        )
        if self.table_style_name:
            try:
                table.style = self.table_style_name
            except Exception:
                # Silently ignore style setting if not available
                pass

        # Set column widths based on ColDef widths
        # Approximate: 1 character ~ 0.12 inches at 11pt
        for c, col_def in enumerate(col_defs):
            if c < len(table.columns):
                char_width = max(col_def.width, 1)
                table.columns[c].width = Inches(char_width * 0.12)

        # Write header row
        row_idx = 0
        if header_row:
            hdr_cells = table.rows[row_idx].cells
            for j, text in enumerate(header_row):
                if j >= len(hdr_cells):
                    break
                p = hdr_cells[j].paragraphs[0]
                # Clear any existing content and add a bold run
                p.clear()
                run = p.add_run(str(text))
                run.bold = True
                # Use alignment from header_defs if available, else center
                if header_defs and j < len(header_defs):
                    hdr_def = header_defs[j]
                    if ">" in hdr_def.align:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif "^" in hdr_def.align:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_idx += 1

        # Helper for formatting with preprocess/postprocess
        def _format_for_export(cd, val, row: list[Any], col_idx: int) -> str:
            v2 = cd.preprocess(val, row, col_idx)
            try:
                if v2 is None:
                    text = cd.prefix + (cd.none_text or "") + cd.suffix
                elif cd.format_spec:
                    fmt = f"{{:{cd.format_spec}}}"
                    text = cd.prefix + fmt.format(v2) + cd.suffix
                else:
                    text = cd.prefix + str(v2) + cd.suffix
            except Exception:
                text = cd.prefix + (cd.none_text if v2 is None else str(v2)) + cd.suffix
            return cd.postprocess(val, text, row, col_idx)

        # Write data rows
        for r, row in enumerate(value_rows):
            cells = table.rows[row_idx].cells
            for c, (val, col_def) in enumerate(zip(row, col_defs)):
                if c >= len(cells):
                    break
                p = cells[c].paragraphs[0]
                p.clear()
                run = p.add_run(_format_for_export(col_def, val, row, c))
                # Align mapping
                if ">" in col_def.align:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif "^" in col_def.align:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            row_idx += 1

        # Save (coerce Path to str for python-docx type expectations)
        save_target = str(file) if isinstance(file, Path) else file
        doc.save(save_target)
