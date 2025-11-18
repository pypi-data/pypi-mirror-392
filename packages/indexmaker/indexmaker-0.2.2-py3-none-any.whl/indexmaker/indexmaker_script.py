"""
This script processes a PDF or Word document to create an index of words
(in docx format) along with the pages they appear on.
"""

import re
import pdfplumber
import argparse
from collections import defaultdict
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.tokenize import word_tokenize
#import docx2pdf
from docx import Document
import nltk
import sys, os


# NLTK setup
for resource in ('tokenizers/punkt','tokenizers/punkt_tab', 'corpora/stopwords'):
    try:
        nltk.data.find(resource)
    except LookupError:
        nltk.download(resource.split('/')[1])

# Constants
MIN_NUM_LETTERS = 2


class IndexMaker():

    def __init__(self, input_file_path:str,  start_from_page:int=1, output_docx:str=None, own_index:str=None, gui_obj=None):
        """
        Initialize the IndexMaker with the path to the PDF and output DOCX file.
        input_file_path: Path to the input PDF file.
        output_docx: Path to the output Word document.
        own_index: Optional pre-defined index dictionary
        start_from_page: Page number to start processing from 
         (only those words will be considered).
        """
        self.input_file_path = input_file_path
        if not os.path.isfile(input_file_path):
                    raise FileNotFoundError(f"The file {input_file_path} does not exist.")

        if not output_docx:
            output_docx = os.path.splitext(os.path.basename(input_file_path))[0] + "_index.docx"
        self.output_docx = output_docx if output_docx else "word_index.docx"
        self.pages = {}
        self.index = None 
        self.own_index = own_index
        self.text_extracted = False
        self.own_index_set = set()
        self._start_from_page = start_from_page
        self.gui_obj = gui_obj

        # If any, extract the list of words to consider
        self.process_own_index()


    @property
    def input_format(self):
        extension = os.path.splitext(self.input_file_path)[1].lower()
        if extension == ".pdf":
            return "pdf"
        elif extension in [".docx", ".doc"]:
            return "word"
        else:
            raise ValueError("Unsupported file format. Please provide a PDF or Word document.")

        send_msg_to_gui(f"Detected input format: {extension}")

    @property
    def start_from_page(self): 

        if isinstance(self._start_from_page, str):
            try:
                return int(max(1, int(self._start_from_page)))
            except:
                send_msg_to_gui("Warning: start_from_page should be an integer. Defaulting to 1.")
                return 1

        elif isinstance(self._start_from_page, int):
            return max(1, _start_from_page)

        elif isinstance(self._start_from_page, float):
            return int(max(1, int(self._start_from_page)))
        else:
            send_msg_to_gui("Warning: start_from_page should be an integer. Defaulting to 1.")
            return 1


    def process_pdf(self, input_file):
        """
        Extract text page by page
        """
        with pdfplumber.open(input_file) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                extract = page.extract_text() or ""
                if page_number >= self.start_from_page:
                    # If we want for example to start at the fourth page and consider it as page 1
                    shift = self.start_from_page - 1
                    self.pages[page_number - shift] = extract.lower()


    def process_microsoft_word(self):
        """
        If the input file is in Microsoft Word format,
        extract text from it page by page after converting it to PDF.
        """
        # temp_pdf = os.path.splitext(os.path.basename(self.input_file_path))[0] + "_temp.pdf"
        # p = docx2pdf.convert(self.input_file_path,temp_pdf)
        # process_pdf(temp_pdf)

        # docx2pdf doesn't work for Linux
        pass

    def process_own_index(self):
        """
        Load a pre-defined index from a file.
        Only these words (as a set) will be considered afterwards
        """
        if self.own_index:

            # Implementation depends on the format of own_index file
            own_index_extension = os.path.splitext(self.own_index)[1].lower()

            # extract all the words in the document and put them in a set
            # pdf file
            if own_index_extension == ".pdf":
                #Extract all the words from the pdf file
                with pdfplumber.open(self.input_file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text() or ""
                        words_set = set(word.lower() for word in word_tokenize(text) if word.isalpha())
                        self.own_index_set.update(words_set)

                        #self.pages[page_number] = extract.lower()
                
            elif own_index_extension in [".docx", ".doc"]:
                for paragraph in Document(self.own_index).paragraphs:
                    text = paragraph.text or ""
                    # filter so that only words matching the pattern "\w+\:" are considered"
                    pattern = r"\w+\:"
                    filtered_words_list = re.findall(pattern, text)
                    filtered_words_str = ' '.join(filtered_words_list)

                    words_set = set(word.lower() for word in word_tokenize(filtered_words_str) if word.isalpha())
                    self.own_index_set.update(words_set)

            elif own_index_extension == ".txt":
                with open(self.own_index, 'r', encoding='utf-8') as f:
                    text = f.read()
                    words_set = set(word.lower() for word in word_tokenize(text) if word.isalpha())
                    self.own_index_set.update(words_set)
            else:

                raise ValueError("Unsupported file format for own_index. Please provide a PDF or Word document.")

            # Now remove all numbers and punctuation signs from the set, 
            self.own_index_set = {word for word in self.own_index_set if len(word) > MIN_NUM_LETTERS}
        else:
            pass

    def extract_all_text(self):
        """
        Extract text from the input file based on its format.
        Put the words in a dict with page numbers as keys.
        """
        if self.input_format == "pdf":
            self.process_pdf(self.input_file_path)
        elif self.input_format == "word":
            #self.process_microsoft_word()
            raise NotImplementedError("Word document processing is not implemented yet.")
        else:
            raise ValueError("Unsupported file format. Please provide a PDF or Word document.")

        self.text_extracted = True

    def create_index(self):

        if not self.text_extracted:
            self.extract_text()

        # Initialize stemmer and stop words
        stemmer = PorterStemmer()
        stop_words = set(stopwords.words("english"))
        word_pages = defaultdict(set)
        stem_to_original = {}

        # Process each page
        for page_number, text in self.pages.items():
            tokens = word_tokenize(text)
            for word in tokens:
                # Skip single-letter words and stop words
                if len(word) > 1 and word.isalpha() and word.lower() not in stop_words:
                    stemmed = stemmer.stem(word.lower())
                    if stemmed not in stem_to_original:
                        stem_to_original[stemmed] = word.lower()
                    word_pages[stem_to_original[stemmed]].add(page_number)

        # Convert to a sorted dictionary
        if self.own_index_set:
            # Filter the word_pages to include only those in own_index_set
            filtered_word_pages = {word: pages for word, pages in word_pages.items() if word in self.own_index_set}
            self.index = {word: sorted(pages) for word, pages in filtered_word_pages.items()}
        else:
            self.index = {word: sorted(pages) for word, pages in word_pages.items()}

    def save_index_to_docx(self):
        # Create a Word document
        doc = Document()
        doc.add_heading("Word Index", level=1)

        if self.index:
            # Sort words alphabetically
            sorted_words = sorted(self.index.keys())
            current_letter = None

            for word in sorted_words:
                # Group by the first letter
                first_letter = word[0].upper()
                if first_letter != current_letter:
                    current_letter = first_letter
                    doc.add_heading(current_letter, level=2)

                # Add the word and its pages in the same paragraph
                paragraph = doc.add_paragraph(f"{word}: {', '.join(map(str, self.index[word]))}")
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = 0
                paragraph_format.space_after = 0
                paragraph_format.line_spacing = 1  # Optional: Set line spacing to single

            # Save the document
            doc.save(self.output_docx)
        else:
            send_msg_to_gui("Empty index!")

    def send_msg_to_gui(self, msg, clearscreen=False):
        """
        Send a message to the GUI if available, otherwise print to console.
        """
        if self.gui_obj:
            self.gui_obj.display_msg(msg, clearscreen)
        else:
            print(msg)

    def main(self):
        self.send_msg_to_gui("Starting text extraction...")
        self.extract_all_text()

        self.send_msg_to_gui("Creating index...")
        self.create_index()

        self.send_msg_to_gui("Saving index to Word document...")
        self.save_index_to_docx()


        return self.output_docx


#----------MAIN

if __name__ == "__main__":

    parser = argparse.ArgumentParser(description="Generate a word index from a PDF file.")
    parser.add_argument("input_file", help="Path to the input file.")
    parser.add_argument("--output", help="Path to the output Word document.", default=None)
    parser.add_argument("--start_from_page", type=int, help="Page number to start processing from.", default=1)
    parser.add_argument("--own_index", help="Path to a pre-defined index file (optional).", default=None)

    args = parser.parse_args()


    # Example usage
    index_maker = IndexMaker(input_file_path=args.input_file, 
                             output_docx=args.output, 
                             start_from_page=args.start_from_page, 
                             own_index=args.own_index)
    index_maker.main()

