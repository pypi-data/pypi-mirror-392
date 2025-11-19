"""Data loading utilities for various file formats"""

import PyPDF2
import docx
import pandas as pd
from pathlib import Path
from typing import Optional

class UnstructuredDataLoader:
    """Load unstructured data from various sources"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.load_pdf,
            '.docx': self.load_word,
            '.doc': self.load_word,
            '.xlsx': self.load_excel,
            '.xls': self.load_excel,
            '.txt': self.load_text,
            '.csv': self.load_csv
        }
    
    def load_pdf(self, file_path: Path) -> str:
        """Load PDF file"""
        try:
            print(f"📕 Loading PDF: {file_path.name}")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text += page.extract_text() + "\n"
                    print(f"   → Page {page_num} extracted")
            
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from PDF")
            return text.strip()
        except Exception as e:
            print(f"❌ Error reading PDF: {e}")
            return ""
    
    def load_word(self, file_path: Path) -> str:
        """Load Word document"""
        try:
            print(f"📘 Loading Word doc: {file_path.name}")
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from Word doc")
            return text.strip()
        except Exception as e:
            print(f"❌ Error reading Word document: {e}")
            return ""
    
    def load_excel(self, file_path: Path) -> str:
        """Load Excel file"""
        try:
            print(f"📊 Loading Excel: {file_path.name}")
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"\nSheet: {sheet_name}\n"
                text += "-" * 40 + "\n"
                
                for index, row in df.iterrows():
                    row_text = []
                    for col in df.columns:
                        if pd.notna(row[col]):
                            row_text.append(f"{col}: {row[col]}")
                    if row_text:
                        text += ", ".join(row_text) + "\n"
            
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from Excel")
            return text.strip()
        except Exception as e:
            print(f"❌ Error reading Excel file: {e}")
            return ""
    
    def load_csv(self, file_path: Path) -> str:
        """Load CSV file"""
        try:
            print(f"📊 Loading CSV: {file_path.name}")
            df = pd.read_csv(file_path)
            text = df.to_string()
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from CSV")
            return text
        except Exception as e:
            print(f"❌ Error reading CSV file: {e}")
            return ""
    
    def load_text(self, file_path: Path) -> str:
        """Load text file"""
        try:
            print(f"📄 Loading text file: {file_path.name}")
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from text file")
            return text.strip()
        except Exception as e:
            print(f"❌ Error reading text file: {e}")
            return ""
    
    def auto_load(self, file_path: str) -> Optional[str]:
        """Automatically detect file type and load"""
        path = Path(file_path)
        if not path.exists():
            print(f"❌ File not found: {file_path}")
            return None
        
        extension = path.suffix.lower()
        if extension in self.supported_formats:
            return self.supported_formats[extension](path)
        else:
            print(f"❌ Unsupported file format: {extension}")
            return None
    
    def load_sample(self) -> str:
        """Load sample data for testing"""
        print("📝 Loading sample data...")
        sample_data = """
John Smith bought 3 laptops on March 15, 2024 for $2,400 total. His email is john.smith@email.com and phone is 555-0123.
Sarah Johnson (sarah.j@company.com) ordered 5 monitors on March 16, 2024. Each monitor cost $300. She can be reached at 555-0456.
Mike Davis purchased 2 keyboards and 3 mice on March 17, 2024. Total was $180. Contact: m.davis@email.com, phone: 555-0789.
Emma Wilson bought 1 laptop and 2 monitors on March 18, 2024 for $1,400. Email: emma.w@tech.com, tel: 555-0234.
The same John Smith made another purchase on March 20, 2024 - 1 printer for $450.
Robert Brown ordered 10 keyboards on March 21, 2024 at $50 each. Contact: r.brown@office.com, 555-0567.
"""
        word_count = len(sample_data.split())
        print(f"✅ Loaded sample data ({word_count} words)")
        return sample_data.strip()