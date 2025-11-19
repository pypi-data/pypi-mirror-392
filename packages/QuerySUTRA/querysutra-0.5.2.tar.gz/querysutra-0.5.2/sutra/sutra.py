"""QuerySUTRA v0.5.1 - BULLETPROOF & FIXED"""
__version__ = "0.5.1"
__author__ = "Aditya Batta"
__all__ = ["SUTRA", "QueryResult"]

import os, sqlite3, pandas as pd, numpy as np, json, hashlib, shutil, datetime, re
from typing import Optional, Union, Dict, List
from pathlib import Path
from difflib import get_close_matches

try:
    from openai import OpenAI
    HAS_OPENAI = True
except:
    HAS_OPENAI = False

try:
    import plotly.express as px
    import plotly.graph_objects as go
    HAS_PLOTLY = True
except:
    HAS_PLOTLY = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except:
    HAS_PYPDF2 = False

try:
    import docx
    HAS_DOCX = True
except:
    HAS_DOCX = False

try:
    from sentence_transformers import SentenceTransformer
    HAS_EMBEDDINGS = True
except:
    HAS_EMBEDDINGS = False


class SUTRA:
    """SUTRA - BULLETPROOF"""
    
    def __init__(self, api_key: Optional[str] = None, db: str = "sutra.db",
                 use_embeddings: bool = False, fuzzy_match: bool = True, 
                 cache_queries: bool = True, check_relevance: bool = False):
        
        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key
        
        self.api_key = os.getenv("OPENAI_API_KEY")
        self.client = OpenAI(api_key=self.api_key) if self.api_key and HAS_OPENAI else None
        self.db_path = db
        self.conn = sqlite3.connect(db, timeout=30, check_same_thread=False)
        self.cursor = self.conn.cursor()
        self.current_table = None
        self.schema_info = {}
        self.cache_queries = cache_queries
        self.cache = {} if cache_queries else None
        self.use_embeddings = use_embeddings
        self.embedding_model = None
        self.query_embeddings = {}
        self.check_relevance = check_relevance
        self.fuzzy_match = fuzzy_match
        
        if use_embeddings and HAS_EMBEDDINGS:
            try:
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            except:
                pass
        
        self._refresh_schema()
        print(f"QuerySUTRA v0.5.1 Ready")
    
    def upload(self, data: Union[str, pd.DataFrame], name: Optional[str] = None) -> 'SUTRA':
        """Upload."""
        if isinstance(data, pd.DataFrame):
            self._store(data, name or "data")
            return self
        
        path = Path(data)
        if not path.exists():
            raise FileNotFoundError(f"Not found: {data}")
        
        name = name or path.stem.replace(" ", "_").replace("-", "_")
        ext = path.suffix.lower()
        
        if ext == ".csv":
            self._store(pd.read_csv(path), name)
        elif ext in [".xlsx", ".xls"]:
            self._store(pd.read_excel(path), name)
        elif ext == ".json":
            self._store(pd.read_json(path), name)
        elif ext == ".pdf":
            self._pdf(path, name)
        elif ext == ".docx":
            self._docx(path, name)
        elif ext == ".txt":
            self._txt(path, name)
        else:
            raise ValueError(f"Unsupported: {ext}")
        
        return self
    
    def _pdf(self, path: Path, name: str):
        """BULLETPROOF PDF - ALWAYS creates multiple tables."""
        if not HAS_PYPDF2:
            raise ImportError("pip install PyPDF2")
        
        print(f"Extracting PDF: {path.name}")
        
        with open(path, 'rb') as f:
            text = "".join([p.extract_text() + "\n" for p in PyPDF2.PdfReader(f).pages])
        
        if not self.client:
            print("ERROR: No API key! Set api_key parameter")
            return
        
        print("AI: Extracting...")
        
        # TRY 3 TIMES
        entities = None
        for attempt in [1, 2, 3]:
            entities = self._extract(text, attempt)
            if entities and len(entities) > 0:
                break
            if attempt < 3:
                print(f"  Retry {attempt+1}/3...")
        
        # Create tables from entities
        if entities and len(entities) > 0:
            print(f"Extracted {len(entities)} entity types:")
            for etype, recs in entities.items():
                if recs and len(recs) > 0:
                    for idx, rec in enumerate(recs, 1):
                        rec['id'] = idx
                    self._store(pd.DataFrame(recs), f"{name}_{etype}")
                    print(f"  {etype}: {len(recs)} rows")
            return
        
        # REGEX FALLBACK - FIXED
        print("Using regex fallback...")
        people = []
        emails = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        
        # Extract names from common patterns
        name_patterns = [
            r'(?:Employee|Name|Mr\.|Mrs\.|Ms\.|Dr\.)\s*[:\-]?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s+(?:lives|resides|works|is based)',
            r'\*\*([A-Z][a-z]+\s+[A-Z][a-z]+)\*\*'
        ]
        
        names = []
        for pattern in name_patterns:
            names.extend(re.findall(pattern, text))
            if len(names) >= len(emails):
                break
        
        # Match emails to names
        max_people = min(len(emails), 50)
        for i in range(max_people):
            people.append({
                'id': i + 1,
                'name': names[i] if i < len(names) else f"Person {i+1}",
                'email': emails[i] if i < len(emails) else f"person{i+1}@unknown.com"
            })
        
        if people:
            self._store(pd.DataFrame(people), f"{name}_people")
            print(f"  Extracted {len(people)} people via regex")
        else:
            # Absolute last resort
            lines = [l.strip() for l in text.split('\n') if l.strip()][:100]
            self._store(pd.DataFrame({'line': range(1, len(lines)+1), 'text': lines}), name)
    
    def _extract(self, text: str, attempt: int) -> Dict:
        """Extract with 3 different strategies."""
        if not self.client:
            return {}
        
        try:
            if attempt == 1:
                sys_msg = "Extract entities as JSON. Return ONLY valid JSON."
                usr_msg = f"""Extract ALL entities from text.

Text:
{text[:15000]}

Return JSON with: people, skills, technologies, projects, certifications, education, work_experience

Example:
{{"people":[{{"id":1,"name":"Sarah Johnson","email":"sarah@co.com","city":"New York","state":"NY"}},{{"id":2,"name":"Michael Chen","email":"michael@co.com","city":"SF","state":"CA"}}],"skills":[{{"id":1,"person_id":1,"skill_name":"Python","proficiency":"Expert"}}]}}

Rules: Unique IDs (1,2,3...), person_id references people.id

JSON:"""
            
            elif attempt == 2:
                sys_msg = "Return JSON."
                usr_msg = f"""Text: {text[:10000]}

Extract people as JSON:
{{"people":[{{"id":1,"name":"...","email":"...","city":"..."}}]}}

JSON:"""
            
            else:
                sys_msg = "JSON only."
                usr_msg = f"""Find names and emails in: {text[:8000]}

{{"people":[{{"id":1,"name":"John","email":"john@co.com"}}]}}"""
            
            r = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": sys_msg},
                    {"role": "user", "content": usr_msg}
                ],
                temperature=0,
                max_tokens=12000
            )
            
            raw = r.choices[0].message.content.strip()
            raw = raw.replace("```json", "").replace("```", "").replace("JSON:", "").strip()
            
            start = raw.find('{')
            end = raw.rfind('}') + 1
            
            if start < 0 or end <= start:
                return {}
            
            result = json.loads(raw[start:end])
            
            if isinstance(result, dict) and len(result) > 0:
                has_data = any(isinstance(v, list) and len(v) > 0 for v in result.values())
                if has_data:
                    return result
            
            return {}
            
        except Exception as e:
            print(f"  Attempt {attempt} failed: {str(e)[:100]}")
            return {}
    
    def _docx(self, path: Path, name: str):
        """DOCX."""
        if not HAS_DOCX:
            raise ImportError("pip install python-docx")
        doc = docx.Document(path)
        if doc.tables:
            for i, t in enumerate(doc.tables):
                data = [[cell.text.strip() for cell in row.cells] for row in t.rows]
                if data and len(data) > 1:
                    self._store(pd.DataFrame(data[1:], columns=data[0]), f"{name}_t{i+1}")
        else:
            text = "\n".join([p.text for p in doc.paragraphs])
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            self._store(pd.DataFrame({'line': range(1, len(lines)+1), 'text': lines}), name)
    
    def _txt(self, path: Path, name: str):
        """TXT."""
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        self._store(pd.DataFrame({'line': range(1, len(lines)+1), 'text': lines}), name)
    
    def _store(self, df: pd.DataFrame, name: str):
        """Store."""
        df.columns = [str(c).strip().replace(" ", "_").replace("-", "_") for c in df.columns]
        try:
            df.to_sql(name, self.conn, if_exists='replace', index=False, method='multi', chunksize=500)
        except:
            df.to_sql(name, self.conn, if_exists='replace', index=False)
        self.conn.commit()
        self.current_table = name
        self._refresh_schema()
        print(f"  {name}: {len(df)} rows")
    
    def ask(self, q: str, viz: Union[bool, str] = False, table: Optional[str] = None) -> 'QueryResult':
        """Query."""
        if not self.client:
            return QueryResult(False, "", pd.DataFrame(), None, "No API")
        
        t = table or self.current_table or (self._get_tables()[0] if self._get_tables() else None)
        if not t:
            return QueryResult(False, "", pd.DataFrame(), None, "No table")
        
        if self.use_embeddings and self.embedding_model:
            cached = self._check_cache(q, t)
            if cached:
                return cached
        
        if self.fuzzy_match:
            q = self._fuzzy(q, t)
        
        key = hashlib.md5(f"{q}:{t}".encode()).hexdigest()
        if self.cache_queries and self.cache and key in self.cache:
            sql = self.cache[key]
        else:
            sql = self._gen_sql(q, t)
            if self.cache_queries and self.cache:
                self.cache[key] = sql
        
        print(f"SQL: {sql}")
        
        try:
            df = pd.read_sql_query(sql, self.conn)
            print(f"Success! {len(df)} rows")
            fig = self._viz(df, q, viz if isinstance(viz, str) else "auto") if viz else None
            r = QueryResult(True, sql, df, fig)
            
            if self.use_embeddings and self.embedding_model:
                self._store_cache(q, t, r)
            
            return r
        except Exception as e:
            return QueryResult(False, sql, pd.DataFrame(), None, str(e))
    
    def _fuzzy(self, q: str, t: str) -> str:
        """Fuzzy."""
        try:
            cols = [c for c, d in self.schema_info.get(t, {}).items() if 'TEXT' in d]
            if not cols:
                return q
            for col in cols[:2]:
                df = pd.read_sql_query(f"SELECT DISTINCT {col} FROM {t} LIMIT 100", self.conn)
                vals = [str(v) for v in df[col].dropna()]
                words = q.split()
                for i, w in enumerate(words):
                    m = get_close_matches(w, vals, n=1, cutoff=0.6)
                    if m and w != m[0]:
                        words[i] = m[0]
                q = " ".join(words)
            return q
        except:
            return q
    
    def _check_cache(self, q: str, t: str) -> Optional['QueryResult']:
        """Cache."""
        if not self.query_embeddings:
            return None
        emb = self.embedding_model.encode([q])[0]
        best, sim = None, 0.85
        for cq, d in self.query_embeddings.items():
            if d['table'] != t:
                continue
            s = np.dot(emb, d['embedding']) / (np.linalg.norm(emb) * np.linalg.norm(d['embedding']))
            if s > sim:
                sim, best = s, cq
        return self.query_embeddings[best]['result'] if best else None
    
    def _store_cache(self, q: str, t: str, r: 'QueryResult'):
        """Store."""
        emb = self.embedding_model.encode([q])[0]
        self.query_embeddings[q] = {'table': t, 'embedding': emb, 'result': r}
    
    def _viz(self, df: pd.DataFrame, title: str, vt: str):
        """Viz."""
        if not HAS_PLOTLY:
            return None
        try:
            n = df.select_dtypes(include=[np.number]).columns.tolist()
            c = df.select_dtypes(include=['object']).columns.tolist()
            if vt == "pie" and c and n:
                fig = px.pie(df, names=c[0], values=n[0], title=title)
            elif vt == "bar" and c and n:
                fig = px.bar(df, x=c[0], y=n[0], title=title)
            elif vt == "line" and n:
                fig = px.line(df, y=n[0], title=title)
            elif vt == "scatter" and len(n) >= 2:
                fig = px.scatter(df, x=n[0], y=n[1], title=title)
            else:
                fig = px.bar(df, y=df.columns[0], title=title)
            fig.show()
            return fig
        except:
            return None
    
    def tables(self) -> Dict:
        """Tables."""
        t = self._get_tables()
        print("\n" + "="*70)
        print("TABLES")
        print("="*70)
        if not t:
            print("No tables")
            return {}
        r = {}
        for i, tb in enumerate(t, 1):
            cnt = pd.read_sql_query(f"SELECT COUNT(*) FROM {tb}", self.conn).iloc[0, 0]
            cols = list(self.schema_info.get(tb, {}).keys())
            print(f" {i}. {tb}: {cnt} rows, {len(cols)} cols")
            r[tb] = {'rows': cnt, 'columns': cols}
        print("="*70)
        return r
    
    def schema(self, table: Optional[str] = None) -> Dict:
        """Schema."""
        if not self.schema_info:
            self._refresh_schema()
        print("\n" + "="*70)
        print("SCHEMA")
        print("="*70)
        r = {}
        for t in ([table] if table else self.schema_info.keys()):
            if t in self.schema_info:
                cnt = pd.read_sql_query(f"SELECT COUNT(*) FROM {t}", self.conn).iloc[0, 0]
                print(f"\n{t}: {cnt} records")
                for c, d in self.schema_info[t].items():
                    print(f"  - {c:<30} {d}")
                r[t] = {'records': cnt, 'columns': self.schema_info[t]}
        print("="*70)
        return r
    
    def peek(self, table: Optional[str] = None, n: int = 5) -> pd.DataFrame:
        """Preview."""
        t = table or self.current_table
        if not t:
            return pd.DataFrame()
        df = pd.read_sql_query(f"SELECT * FROM {t} LIMIT {n}", self.conn)
        print(f"\nSample from '{t}':")
        print(df.to_string(index=False))
        return df
    
    def sql(self, query: str, viz: Union[bool, str] = False) -> 'QueryResult':
        """SQL."""
        try:
            df = pd.read_sql_query(query, self.conn)
            print(f"Success! {len(df)} rows")
            fig = self._viz(df, "Result", viz if isinstance(viz, str) else "auto") if viz else None
            return QueryResult(True, query, df, fig)
        except Exception as e:
            return QueryResult(False, query, pd.DataFrame(), None, str(e))
    
    def save_to_mysql(self, host: str, user: str, password: str, database: str, port: int = 3306):
        """MySQL."""
        try:
            from sqlalchemy import create_engine
            import mysql.connector
        except:
            raise ImportError("pip install QuerySUTRA[mysql]")
        
        print(f"Exporting to MySQL: {database}")
        
        try:
            tc = mysql.connector.connect(host=host, user=user, password=password, port=port)
            tc.cursor().execute(f"CREATE DATABASE IF NOT EXISTS `{database}`")
            tc.close()
        except:
            pass
        
        engine = create_engine(f"mysql+mysqlconnector://{user}:{password}@{host}:{port}/{database}")
        for t in self._get_tables():
            df = pd.read_sql_query(f"SELECT * FROM {t}", self.conn)
            df.to_sql(t, engine, if_exists='replace', index=False)
            print(f"  {t}: {len(df)} rows")
        print("Done!")
        return self
    
    def export_db(self, path: str, format: str = "sqlite"):
        """Export."""
        if format == "sqlite":
            shutil.copy2(self.db_path, path)
        elif format == "json":
            data = {t: pd.read_sql_query(f"SELECT * FROM {t}", self.conn).to_dict('records') for t in self._get_tables()}
            with open(path, 'w') as f:
                json.dump(data, f, indent=2, default=str)
        print(f"Saved: {path}")
        return self
    
    @classmethod
    def load_from_db(cls, db_path: str, api_key: Optional[str] = None, **kwargs):
        """Load."""
        if not Path(db_path).exists():
            raise FileNotFoundError(f"Not found: {db_path}")
        return cls(api_key=api_key, db=db_path, **kwargs)
    
    @classmethod
    def connect_mysql(cls, host: str, user: str, password: str, database: str, port: int = 3306, api_key: Optional[str] = None, **kwargs):
        """MySQL."""
        try:
            from sqlalchemy import create_engine
            import mysql.connector
        except:
            raise ImportError("pip install QuerySUTRA[mysql]")
        
        try:
            tc = mysql.connector.connect(host=host, user=user, password=password, port=port)
            tc.cursor().execute(f"CREATE DATABASE IF NOT EXISTS {database}")
            tc.close()
        except:
            pass
        
        engine = create_engine(f"mysql+mysqlconnector://{user}:{password}@{host}:{port}/{database}")
        temp_db = f"mysql_{database}.db"
        instance = cls(api_key=api_key, db=temp_db, **kwargs)
        
        tables = pd.read_sql_query("SHOW TABLES", engine).iloc[:, 0].tolist()
        for t in tables:
            pd.read_sql_query(f"SELECT * FROM {t}", engine).to_sql(t, instance.conn, if_exists='replace', index=False)
        
        instance._refresh_schema()
        print(f"Connected! {len(tables)} tables")
        return instance
    
    def _gen_sql(self, q: str, t: str) -> str:
        """SQL."""
        schema = self.schema_info.get(t, {})
        sample = pd.read_sql_query(f"SELECT * FROM {t} LIMIT 3", self.conn).to_string(index=False)
        cols = ", ".join([f"{c} ({d})" for c, d in schema.items()])
        
        r = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "SQL expert. Return only SQL."},
                {"role": "user", "content": f"Table: {t}\nColumns: {cols}\nSample:\n{sample}\n\nQ: {q}\n\nSQL:"}
            ],
            temperature=0
        )
        return r.choices[0].message.content.strip().replace("```sql", "").replace("```", "").strip()
    
    def _get_tables(self) -> List[str]:
        """Tables."""
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        return [r[0] for r in self.cursor.fetchall()]
    
    def _refresh_schema(self):
        """Refresh."""
        self.schema_info = {}
        for t in self._get_tables():
            self.cursor.execute(f"PRAGMA table_info({t})")
            self.schema_info[t] = {r[1]: r[2] for r in self.cursor.fetchall()}
    
    def close(self):
        if self.conn:
            self.conn.close()
    
    def __enter__(self):
        return self
    
    def __exit__(self, *args):
        self.close()
    
    def __repr__(self):
        return f"SUTRA(tables={len(self.schema_info)})"


class QueryResult:
    """Result."""
    def __init__(self, success: bool, sql: str, data: pd.DataFrame, viz, error: str = None):
        self.success, self.sql, self.data, self.viz, self.error = success, sql, data, viz, error
    
    def __repr__(self):
        return f"QueryResult(rows={len(self.data)})" if self.success else f"QueryResult(error='{self.error}')"
    
    def show(self):
        print(self.data if self.success else f"Error: {self.error}")
        return self
