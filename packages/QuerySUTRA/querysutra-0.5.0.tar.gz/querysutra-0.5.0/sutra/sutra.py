"""
QuerySUTRA v0.4.5 - FIXED AI EXTRACTION
Debug mode to see why extraction fails
"""

__version__ = "0.4.5"
__author__ = "Aditya Batta"
__all__ = ["SUTRA", "QueryResult", "quick_start"]

import os
import sqlite3
import pandas as pd
import numpy as np
from typing import Optional, Union, Dict, List
from pathlib import Path
import json
import hashlib
import warnings
import shutil
import datetime
from io import StringIO
from difflib import get_close_matches
warnings.filterwarnings('ignore')

try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    import plotly.express as px
    import plotly.graph_objects as go
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

try:
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from sentence_transformers import SentenceTransformer
    HAS_EMBEDDINGS = True
except ImportError:
    HAS_EMBEDDINGS = False


class SUTRA:
    """SUTRA: Structured-Unstructured-Text-Retrieval-Architecture"""
    
    def __init__(self, api_key: Optional[str] = None, db: str = "sutra.db",
                 use_embeddings: bool = False, check_relevance: bool = False,
                 fuzzy_match: bool = True, cache_queries: bool = True, debug: bool = False):
        """Initialize."""
        print("Initializing QuerySUTRA v0.4.5")
        
        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key
        
        self.api_key = os.getenv("OPENAI_API_KEY")
        self.client = OpenAI(api_key=self.api_key) if self.api_key and HAS_OPENAI else None
        
        self.db_path = db
        self.debug = debug
        
        try:
            self.conn = sqlite3.connect(db, timeout=30, check_same_thread=False)
            self.conn.execute("PRAGMA journal_mode=WAL")
            self.conn.execute("PRAGMA synchronous=NORMAL")
        except:
            self.conn = sqlite3.connect(db, check_same_thread=False)
        
        self.cursor = self.conn.cursor()
        self.current_table = None
        self.schema_info = {}
        
        self.cache_queries = cache_queries
        self.cache = {} if cache_queries else None
        self.use_embeddings = use_embeddings
        self.embedding_model = None
        self.query_embeddings = {}
        self.check_relevance = check_relevance
        self.fuzzy_match = fuzzy_match
        
        if use_embeddings and HAS_EMBEDDINGS:
            try:
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            except:
                self.use_embeddings = False
        
        self._refresh_schema()
        print(f"Ready! Database: {db}")
    
    @classmethod
    def load_from_db(cls, db_path: str, api_key: Optional[str] = None, **kwargs):
        """Load existing database."""
        if not Path(db_path).exists():
            raise FileNotFoundError(f"Not found: {db_path}")
        return cls(api_key=api_key, db=db_path, **kwargs)
    
    @classmethod
    def connect_mysql(cls, host: str, user: str, password: str, database: str,
                     port: int = 3306, api_key: Optional[str] = None, **kwargs):
        """Connect to MySQL."""
        try:
            from sqlalchemy import create_engine
            import mysql.connector
        except ImportError:
            raise ImportError("Run: pip install QuerySUTRA[mysql]")
        
        print(f"Connecting to MySQL...")
        
        try:
            temp_conn = mysql.connector.connect(host=host, user=user, password=password, port=port)
            temp_cursor = temp_conn.cursor()
            temp_cursor.execute(f"CREATE DATABASE IF NOT EXISTS {database}")
            temp_cursor.close()
            temp_conn.close()
        except:
            pass
        
        engine = create_engine(f"mysql+mysqlconnector://{user}:{password}@{host}:{port}/{database}")
        temp_db = f"sutra_mysql_{database}.db"
        instance = cls(api_key=api_key, db=temp_db, **kwargs)
        
        tables = pd.read_sql_query("SHOW TABLES", engine).iloc[:, 0].tolist()
        
        for table in tables:
            df = pd.read_sql_query(f"SELECT * FROM {table}", engine)
            df.to_sql(table, instance.conn, if_exists='replace', index=False)
        
        instance._refresh_schema()
        print(f"Connected! {len(tables)} tables")
        return instance
    
    @classmethod
    def connect_postgres(cls, host: str, user: str, password: str, database: str,
                        port: int = 5432, api_key: Optional[str] = None, **kwargs):
        """Connect to PostgreSQL."""
        try:
            from sqlalchemy import create_engine
        except ImportError:
            raise ImportError("Run: pip install QuerySUTRA[postgres]")
        
        print(f"Connecting to PostgreSQL...")
        
        engine = create_engine(f"postgresql://{user}:{password}@{host}:{port}/{database}")
        temp_db = f"sutra_postgres_{database}.db"
        instance = cls(api_key=api_key, db=temp_db, **kwargs)
        
        tables = pd.read_sql_query("SELECT tablename FROM pg_tables WHERE schemaname='public'", engine)['tablename'].tolist()
        
        for table in tables:
            df = pd.read_sql_query(f"SELECT * FROM {table}", engine)
            df.to_sql(table, instance.conn, if_exists='replace', index=False)
        
        instance._refresh_schema()
        print(f"Connected! {len(tables)} tables")
        return instance
    
    def upload(self, data: Union[str, pd.DataFrame], name: Optional[str] = None, 
               extract_entities: Optional[List[str]] = None,
               auto_export_mysql: Optional[Dict[str, str]] = None) -> 'SUTRA':
        """Upload data."""
        print("\nUploading...")
        
        if isinstance(data, pd.DataFrame):
            name = name or "data"
            self._store_dataframe(data, name)
        else:
            path = Path(data)
            if not path.exists():
                raise FileNotFoundError(f"Not found: {data}")
            
            name = name or path.stem.replace(" ", "_").replace("-", "_")
            ext = path.suffix.lower()
            
            print(f"File: {path.name}")
            
            if ext == ".csv":
                self._store_dataframe(pd.read_csv(path), name)
            elif ext in [".xlsx", ".xls"]:
                self._store_dataframe(pd.read_excel(path), name)
            elif ext == ".json":
                self._store_dataframe(pd.read_json(path), name)
            elif ext == ".sql":
                with open(path) as f:
                    self.cursor.executescript(f.read())
                self.conn.commit()
                self._refresh_schema()
            elif ext == ".pdf":
                self._smart_upload_pdf(path, name, extract_entities)
            elif ext == ".docx":
                self._smart_upload_docx(path, name, extract_entities)
            elif ext == ".txt":
                self._smart_upload_txt(path, name, extract_entities)
            else:
                raise ValueError(f"Unsupported: {ext}")
        
        if auto_export_mysql:
            print("\nAuto-exporting to MySQL...")
            self.save_to_mysql(
                host=auto_export_mysql.get('host', 'localhost'),
                user=auto_export_mysql.get('user', 'root'),
                password=auto_export_mysql['password'],
                database=auto_export_mysql['database'],
                port=auto_export_mysql.get('port', 3306)
            )
        
        return self
    
    def _smart_upload_pdf(self, path: Path, base_name: str, extract_entities: Optional[List[str]] = None):
        """Parse PDF."""
        if not HAS_PYPDF2:
            raise ImportError("Run: pip install PyPDF2")
        
        print("Extracting PDF...")
        
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            full_text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                full_text += page.extract_text() + "\n"
                print(f"  Page {page_num}/{len(pdf_reader.pages)}")
        
        if self.client:
            print("AI: Extracting entities...")
            
            chunk_size = 10000
            all_entities = {}
            
            for i in range(0, len(full_text), chunk_size):
                chunk = full_text[i:i+chunk_size]
                chunk_num = (i // chunk_size) + 1
                total_chunks = (len(full_text) // chunk_size) + 1
                
                if total_chunks > 1:
                    print(f"  Chunk {chunk_num}/{total_chunks}...")
                
                entities = self._extract_chunk(chunk, extract_entities)
                
                if self.debug:
                    print(f"    DEBUG: Chunk {chunk_num} returned {len(entities)} entity types")
                
                for entity_type, records in entities.items():
                    if entity_type not in all_entities:
                        all_entities[entity_type] = []
                    all_entities[entity_type].extend(records)
            
            if self.debug:
                print(f"  DEBUG: Total entities collected: {len(all_entities)}")
                for k, v in all_entities.items():
                    print(f"    - {k}: {len(v)} records")
            
            # Renumber IDs
            for entity_type, records in all_entities.items():
                for idx, record in enumerate(records, 1):
                    record['id'] = idx
            
            # Create tables
            if all_entities:
                print(f"\nCreated {len(all_entities)} tables:")
                for entity_type, records in all_entities.items():
                    if records:
                        table_name = f"{base_name}_{entity_type}"
                        df = pd.DataFrame(records)
                        self._store_dataframe_safe(df, table_name)
                        print(f"  {entity_type}: {len(df)} records")
                return
        
        print("Creating simple table")
        self._store_dataframe(self._parse_text_simple(full_text), base_name)
    
    def _extract_chunk(self, text: str, custom_entities: Optional[List[str]] = None) -> Dict:
        """Extract entities - WITH BETTER ERROR HANDLING."""
        if not self.client:
            return {}
        
        try:
            prompt = f"""Extract ALL structured entities from this text.

Text:
{text[:8000]}

Extract: people, skills, technologies, projects, certifications, education, work_experience, events, organizations, or ANY structured data.

Return JSON with arrays. Sequential IDs. Foreign keys reference primary keys.

{{
  "people": [{{"id": 1, "name": "John", "email": "john@co.com", "city": "Dallas"}}, ...],
  "skills": [{{"id": 1, "person_id": 1, "skill_name": "Python"}}, ...]
}}

ONLY valid JSON. No explanations."""

            resp = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Extract ALL entities with unique IDs. Return ONLY valid JSON, nothing else."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0,
                max_tokens=8000
            )
            
            json_text = resp.choices[0].message.content.strip()
            
            if self.debug:
                print(f"    DEBUG: AI response length: {len(json_text)} chars")
                print(f"    DEBUG: First 200 chars: {json_text[:200]}")
            
            json_text = json_text.replace("```json", "").replace("```", "").strip()
            
            result = json.loads(json_text)
            
            if self.debug:
                print(f"    DEBUG: Parsed {len(result)} entity types")
            
            return result
            
        except json.JSONDecodeError as e:
            if self.debug:
                print(f"    DEBUG: JSON parse error: {e}")
                print(f"    DEBUG: Response was: {json_text[:500]}")
            return {}
        except Exception as e:
            if self.debug:
                print(f"    DEBUG: Extraction error: {e}")
            return {}
    
    def _smart_upload_docx(self, path: Path, base_name: str, extract_entities: Optional[List[str]] = None):
        """Parse DOCX."""
        if not HAS_DOCX:
            raise ImportError("Run: pip install python-docx")
        
        doc = docx.Document(path)
        
        if doc.tables:
            for i, table in enumerate(doc.tables):
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if data and len(data) > 1:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    self._store_dataframe(df, f"{base_name}_table_{i+1}" if len(doc.tables) > 1 else base_name)
            return
        
        text = "\n".join([para.text for para in doc.paragraphs])
        
        if self.client and len(text) > 0:
            entities = self._extract_chunk(text, extract_entities)
            if entities:
                for entity_type, records in entities.items():
                    if records:
                        df = pd.DataFrame(records)
                        self._store_dataframe_safe(df, f"{base_name}_{entity_type}")
                return
        
        self._store_dataframe(self._parse_text_simple(text), base_name)
    
    def _smart_upload_txt(self, path: Path, base_name: str, extract_entities: Optional[List[str]] = None):
        """Parse TXT."""
        with open(path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        if self.client and len(text) > 0:
            entities = self._extract_chunk(text, extract_entities)
            if entities:
                for entity_type, records in entities.items():
                    if records:
                        df = pd.DataFrame(records)
                        self._store_dataframe_safe(df, f"{base_name}_{entity_type}")
                return
        
        self._store_dataframe(self._parse_text_simple(text), base_name)
    
    def _store_dataframe_safe(self, df: pd.DataFrame, name: str):
        """Store."""
        try:
            df.columns = [str(c).strip().replace(" ", "_").replace("-", "_") for c in df.columns]
            df.to_sql(name, self.conn, if_exists='replace', index=False, method='multi', chunksize=500)
            self.conn.commit()
            self.current_table = name
            self._refresh_schema()
        except:
            df.to_sql(name, self.conn, if_exists='replace', index=False)
            self.conn.commit()
            self.current_table = name
            self._refresh_schema()
    
    def _parse_text_simple(self, text: str) -> pd.DataFrame:
        """Simple parsing."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines:
            return pd.DataFrame({'content': ['No content']})
        return pd.DataFrame({'line_number': range(1, len(lines) + 1), 'content': lines})
    
    def _store_dataframe(self, df: pd.DataFrame, name: str):
        """Store."""
        self._store_dataframe_safe(df, name)
        print(f"Uploaded: {name} ({len(df)} rows)")
    
    def ask(self, question: str, viz: Union[bool, str] = False, table: Optional[str] = None) -> 'QueryResult':
        """Query."""
        if not self.client:
            return QueryResult(False, "", pd.DataFrame(), None, "No API key")
        
        print(f"\nQuestion: {question}")
        
        if self.check_relevance and not self._is_relevant_query(question):
            print("Warning: Irrelevant")
            choice = input("Continue? (yes/no): ").strip().lower()
            if choice not in ['yes', 'y']:
                return QueryResult(False, "", pd.DataFrame(), None, "Irrelevant")
        
        tbl = table or self.current_table or (self._get_table_names()[0] if self._get_table_names() else None)
        if not tbl:
            return QueryResult(False, "", pd.DataFrame(), None, "No table")
        
        if self.use_embeddings and self.embedding_model:
            cached = self._check_embedding_cache(question, tbl)
            if cached:
                print("  Cached")
                return cached
        
        if self.fuzzy_match:
            question = self._apply_fuzzy_matching(question, tbl)
        
        cache_key = hashlib.md5(f"{question}:{tbl}".encode()).hexdigest()
        if self.cache_queries and self.cache and cache_key in self.cache:
            sql_query = self.cache[cache_key]
            print("  From cache")
        else:
            sql_query = self._generate_sql(question, tbl)
            if self.cache_queries and self.cache is not None:
                self.cache[cache_key] = sql_query
        
        print(f"SQL: {sql_query}")
        
        try:
            df = pd.read_sql_query(sql_query, self.conn)
            print(f"Success! {len(df)} rows")
            
            fig = None
            if viz:
                viz_type = viz if isinstance(viz, str) else "auto"
                fig = self._visualize(df, question, viz_type)
            
            result = QueryResult(True, sql_query, df, fig)
            
            if self.use_embeddings and self.embedding_model:
                self._store_in_embedding_cache(question, tbl, result)
            
            return result
        except Exception as e:
            print(f"Error: {e}")
            return QueryResult(False, sql_query, pd.DataFrame(), None, str(e))
    
    def _is_relevant_query(self, question: str) -> bool:
        """Check relevance."""
        if not self.client:
            return True
        try:
            tables = self._get_table_names()[:3]
            resp = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Return 'yes' or 'no'."},
                    {"role": "user", "content": f"Relevant to DB with tables {', '.join(tables)}?\n\nQ: {question}\n\nyes/no:"}
                ],
                temperature=0,
                max_tokens=5
            )
            return 'yes' in resp.choices[0].message.content.lower()
        except:
            return True
    
    def _apply_fuzzy_matching(self, question: str, table: str) -> str:
        """Fuzzy."""
        if not self.schema_info.get(table):
            return question
        
        try:
            string_cols = [col for col, dtype in self.schema_info[table].items() if 'TEXT' in dtype]
            if not string_cols:
                return question
            
            for col in string_cols[:2]:
                df = pd.read_sql_query(f"SELECT DISTINCT {col} FROM {table} LIMIT 100", self.conn)
                values = [str(v) for v in df[col].dropna().tolist()]
                
                words = question.split()
                for i, word in enumerate(words):
                    matches = get_close_matches(word, values, n=1, cutoff=0.6)
                    if matches and word != matches[0]:
                        words[i] = matches[0]
                        print(f"  Fuzzy: '{word}' -> '{matches[0]}'")
                question = " ".join(words)
            return question
        except:
            return question
    
    def _check_embedding_cache(self, question: str, table: str) -> Optional['QueryResult']:
        """Check cache."""
        if not self.query_embeddings:
            return None
        
        q_emb = self.embedding_model.encode([question])[0]
        best_match, best_sim = None, 0.85
        
        for cached_q, data in self.query_embeddings.items():
            if data['table'] != table:
                continue
            sim = np.dot(q_emb, data['embedding']) / (np.linalg.norm(q_emb) * np.linalg.norm(data['embedding']))
            if sim > best_sim:
                best_sim = sim
                best_match = cached_q
        
        if best_match:
            print(f"  Similar ({best_sim:.0%})")
            return self.query_embeddings[best_match]['result']
        return None
    
    def _store_in_embedding_cache(self, question: str, table: str, result: 'QueryResult'):
        """Store."""
        q_emb = self.embedding_model.encode([question])[0]
        self.query_embeddings[question] = {'table': table, 'embedding': q_emb, 'result': result}
    
    def _visualize(self, df: pd.DataFrame, title: str, viz_type: str = "auto"):
        """Viz."""
        if not HAS_PLOTLY and not HAS_MATPLOTLIB:
            return None
        print(f"Creating {viz_type} chart...")
        return self._plotly_viz(df, title, viz_type) if HAS_PLOTLY else self._matplotlib_viz(df, title, viz_type)
    
    def _plotly_viz(self, df: pd.DataFrame, title: str, viz_type: str):
        """Plotly."""
        try:
            num = df.select_dtypes(include=[np.number]).columns.tolist()
            cat = df.select_dtypes(include=['object']).columns.tolist()
            
            if viz_type == "table":
                fig = go.Figure(data=[go.Table(header=dict(values=list(df.columns)), cells=dict(values=[df[c] for c in df.columns]))])
            elif viz_type == "pie" and cat and num:
                fig = px.pie(df, names=cat[0], values=num[0], title=title)
            elif viz_type == "bar" and cat and num:
                fig = px.bar(df, x=cat[0], y=num[0], title=title)
            elif viz_type == "line" and num:
                fig = px.line(df, y=num[0], title=title)
            elif viz_type == "scatter" and len(num) >= 2:
                fig = px.scatter(df, x=num[0], y=num[1], title=title)
            elif viz_type == "heatmap" and len(num) >= 2:
                corr = df[num].corr()
                fig = go.Figure(data=go.Heatmap(z=corr.values, x=corr.columns, y=corr.columns))
                fig.update_layout(title=title)
            else:
                if cat and num:
                    fig = px.pie(df, names=cat[0], values=num[0], title=title) if len(df) <= 10 else px.bar(df, x=cat[0], y=num[0], title=title)
                else:
                    fig = px.bar(df, y=df.columns[0], title=title)
            fig.show()
            return fig
        except:
            return None
    
    def _matplotlib_viz(self, df: pd.DataFrame, title: str, viz_type: str):
        """Matplotlib."""
        try:
            plt.figure(figsize=(10, 6))
            num = df.select_dtypes(include=[np.number]).columns
            if viz_type == "pie":
                df[df.columns[0]].value_counts().plot(kind='pie')
            elif viz_type == "line" and len(num) > 0:
                df[num[0]].plot(kind='line')
            else:
                (df[num[0]] if len(num) > 0 else df.iloc[:, 0].value_counts()).plot(kind='bar')
            plt.title(title)
            plt.tight_layout()
            plt.show()
            return plt.gcf()
        except:
            return None
    
    def tables(self) -> Dict[str, dict]:
        """List."""
        print("\n" + "="*70)
        print("TABLES")
        print("="*70)
        
        all_tables = self._get_table_names()
        if not all_tables:
            print("No tables")
            return {}
        
        result = {}
        for i, tbl in enumerate(all_tables, 1):
            cnt = pd.read_sql_query(f"SELECT COUNT(*) FROM {tbl}", self.conn).iloc[0, 0]
            cols = list(self.schema_info.get(tbl, {}).keys())
            print(f" {i}. {tbl}: {cnt} rows, {len(cols)} columns")
            result[tbl] = {'rows': cnt, 'columns': cols}
        
        print("="*70)
        return result
    
    def schema(self, table: Optional[str] = None) -> dict:
        """Schema."""
        if not self.schema_info:
            self._refresh_schema()
        
        print("\n" + "="*70)
        print("SCHEMA")
        print("="*70)
        
        result = {}
        for tbl in ([table] if table else self.schema_info.keys()):
            if tbl in self.schema_info:
                cnt = pd.read_sql_query(f"SELECT COUNT(*) FROM {tbl}", self.conn).iloc[0, 0]
                print(f"\n{tbl}: {cnt} records")
                for col, dtype in self.schema_info[tbl].items():
                    print(f"  - {col:<30} {dtype}")
                result[tbl] = {'records': cnt, 'columns': self.schema_info[tbl]}
        
        print("="*70)
        return result
    
    def peek(self, table: Optional[str] = None, n: int = 5) -> pd.DataFrame:
        """Preview."""
        tbl = table or self.current_table
        if not tbl:
            return pd.DataFrame()
        df = pd.read_sql_query(f"SELECT * FROM {tbl} LIMIT {n}", self.conn)
        print(f"\nSample from '{tbl}':")
        print(df.to_string(index=False))
        return df
    
    def info(self):
        """Info."""
        return self.tables()
    
    def sql(self, query: str, viz: Union[bool, str] = False) -> 'QueryResult':
        """SQL."""
        try:
            df = pd.read_sql_query(query, self.conn)
            print(f"Success! {len(df)} rows")
            fig = self._visualize(df, "Result", viz if isinstance(viz, str) else "auto") if viz else None
            return QueryResult(True, query, df, fig)
        except Exception as e:
            print(f"Error: {e}")
            return QueryResult(False, query, pd.DataFrame(), None, str(e))
    
    def interactive(self, question: str) -> 'QueryResult':
        """Interactive."""
        choice = input("Visualize? (yes/no/pie/bar/line/scatter): ").strip().lower()
        viz = choice if choice in ['pie', 'bar', 'line', 'scatter', 'table', 'heatmap'] else (True if choice in ['yes', 'y'] else False)
        return self.ask(question, viz=viz)
    
    def export_db(self, path: str, format: str = "sqlite"):
        """Export."""
        if format == "sqlite":
            shutil.copy2(self.db_path, path)
        elif format == "sql":
            with open(path, 'w', encoding='utf-8') as f:
                for line in self.conn.iterdump():
                    f.write(f'{line}\n')
        elif format == "json":
            data = {t: pd.read_sql_query(f"SELECT * FROM {t}", self.conn).to_dict('records') for t in self._get_table_names()}
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, default=str)
        elif format == "excel":
            with pd.ExcelWriter(path, engine='openpyxl') as writer:
                for t in self._get_table_names():
                    pd.read_sql_query(f"SELECT * FROM {t}", self.conn).to_excel(writer, sheet_name=t[:31], index=False)
        else:
            raise ValueError(f"Unsupported: {format}")
        print(f"Saved: {path}")
        return self
    
    def save_to_mysql(self, host: str, user: str, password: str, database: str, 
                      port: int = 3306, tables: Optional[List[str]] = None, auto_create: bool = True):
        """Export to MySQL."""
        try:
            from sqlalchemy import create_engine
            import mysql.connector
        except ImportError:
            raise ImportError("Run: pip install QuerySUTRA[mysql]")
        
        print(f"Exporting to MySQL: {host}/{database}")
        
        if auto_create:
            try:
                temp_conn = mysql.connector.connect(host=host, user=user, password=password, port=port)
                temp_cursor = temp_conn.cursor()
                temp_cursor.execute(f"CREATE DATABASE IF NOT EXISTS `{database}`")
                temp_cursor.close()
                temp_conn.close()
                print(f"  Database '{database}' ready")
            except Exception as e:
                print(f"  Warning: {e}")
        
        engine = create_engine(f"mysql+mysqlconnector://{user}:{password}@{host}:{port}/{database}")
        
        for t in (tables or self._get_table_names()):
            df = pd.read_sql_query(f"SELECT * FROM {t}", self.conn)
            df.to_sql(t, engine, if_exists='replace', index=False)
            print(f"  {t}: {len(df)} rows")
        
        print("Complete!")
        return self
    
    def save_to_postgres(self, host: str, user: str, password: str, database: str, port: int = 5432, tables: Optional[List[str]] = None):
        """PostgreSQL."""
        try:
            from sqlalchemy import create_engine
            engine = create_engine(f"postgresql://{user}:{password}@{host}:{port}/{database}")
            print(f"Exporting to PostgreSQL...")
            for t in (tables or self._get_table_names()):
                df = pd.read_sql_query(f"SELECT * FROM {t}", self.conn)
                df.to_sql(t, engine, if_exists='replace', index=False)
                print(f"  {t}: {len(df)} rows")
            print("Complete!")
            return self
        except ImportError:
            raise ImportError("Run: pip install QuerySUTRA[postgres]")
    
    def backup(self, path: str = None):
        """Backup."""
        dir = Path(path) if path else Path(".")
        dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.export_db(str(dir / f"sutra_{ts}.db"), "sqlite")
        self.export_db(str(dir / f"sutra_{ts}.json"), "json")
        print("Backup complete!")
        return self
    
    def export(self, data: pd.DataFrame, path: str, format: str = "csv"):
        """Export."""
        if format == "csv":
            data.to_csv(path, index=False)
        elif format in ["excel", "xlsx"]:
            data.to_excel(path, index=False)
        elif format == "json":
            data.to_json(path, orient="records", indent=2)
        print(f"Exported: {path}")
        return self
    
    def close(self):
        """Close."""
        if self.conn:
            self.conn.close()
    
    def _get_table_names(self) -> List[str]:
        """Tables."""
        self.cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        return [r[0] for r in self.cursor.fetchall()]
    
    def _refresh_schema(self):
        """Refresh."""
        self.schema_info = {}
        for tbl in self._get_table_names():
            self.cursor.execute(f"PRAGMA table_info({tbl})")
            self.schema_info[tbl] = {r[1]: r[2] for r in self.cursor.fetchall()}
    
    def _generate_sql(self, question: str, table: str) -> str:
        """SQL."""
        schema = self.schema_info.get(table, {})
        sample = pd.read_sql_query(f"SELECT * FROM {table} LIMIT 3", self.conn).to_string(index=False)
        schema_str = ", ".join([f"{col} ({dtype})" for col, dtype in schema.items()])
        
        resp = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "SQL expert. Return only SQL."},
                {"role": "user", "content": f"Table: {table}\nColumns: {schema_str}\nSample:\n{sample}\n\nQ: {question}\n\nSQL:"}
            ],
            temperature=0
        )
        return resp.choices[0].message.content.strip().replace("```sql", "").replace("```", "").strip()
    
    def __enter__(self):
        return self
    
    def __exit__(self, *args):
        self.close()
    
    def __repr__(self):
        return f"SUTRA(tables={len(self.schema_info)})"


class QueryResult:
    """Result."""
    def __init__(self, success: bool, sql: str, data: pd.DataFrame, viz, error: str = None):
        self.success, self.sql, self.data, self.viz, self.error = success, sql, data, viz, error
    
    def __repr__(self):
        return f"QueryResult(rows={len(self.data)})" if self.success else f"QueryResult(error='{self.error}')"
    
    def show(self):
        print(self.data if self.success else f"Error: {self.error}")
        return self


def quick_start(api_key: str, data_path: str, question: str, viz: Union[bool, str] = False):
    """Quick."""
    with SUTRA(api_key=api_key) as sutra:
        sutra.upload(data_path)
        return sutra.ask(question, viz=viz)
