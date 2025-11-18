import time
from pathlib import Path
import requests
import json
import re
import os
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import logging

session = requests.Session()


def create_path(path: str) -> None:
    if not os.path.exists(path):
        os.makedirs(path)


def get_article_html(cookie: str, article_id: str):
    """
    get article html
    :param cookie: website's cookie information
    :param article_id: article's url such as 1097430290934005800 (https://www.bilibili.com/opus/1097430290934005800?spm_id_from=333.1387.0.0)
    :return: html
    """
    url = f"https://www.bilibili.com/opus/{article_id}"
    headers = {
        'cookie': cookie,
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    response = session.get(url, headers=headers)
    return response.text


def get_p_tag_text(html) -> tuple:
    result = html.xpath('./span | ./strong')
    color = result[0].get('style')
    if color is not None:
        color = re.findall(r'color:(.*?);', result[0].get('style'))[0]
    is_strong = result[0].tag == 'strong'
    if len(result) != 0:
        return ''.join(result[i].text for i in range(0, len(result))), color, is_strong
    else:
        return '\n', color, is_strong


def get_img(cookie, article_id: str, url_list: list, path: Path) -> list:
    """
    get img.
    :param cookie: website's cookie information
    :param article_id: article's ID
    :param url_list: each item under the opus-module-content class
    :param path: img path
    :return: image path list
    """
    result = list()

    headers = {
        'cookie': cookie,
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36'
    }
    for url in url_list:
        response = session.get(url, headers=headers)
        time.sleep(0.5)

        name = url.split('/')[-1]

        Path(path/article_id).mkdir(parents=True, exist_ok=True)

        with open(path/article_id/name, 'wb') as f:
            f.write(response.content)

        logging.log(logging.INFO,'get' + str(path/article_id/name))

        result.append(path/article_id/name)

    return result


def get_rgb(node: dict) -> tuple:
    """
    get rgb color
    :param node: node
    :return: r, g, b
    """
    if 'color' in node['word'].keys() and node['word']['color'] is not None:
        color = node['word']['color'].strip('#')
        return int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    else:
        return 47, 50, 56


def add_hyperlink(paragraph, text, url, color="00699D", underline=True):
    """
    add hyperlink
    :param paragraph: paragraph
    :param text: text to be added
    :param url: link url
    :param color: color for hyperlink,default 00699D
    :param underline: text underline for hyperlink,default True
    :return: underline element
    """
    # get document part
    part = paragraph.part
    # create a relationship
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)

    # create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # create text run
    new_run = OxmlElement('w:r')

    # create runtime attributes
    rPr = OxmlElement('w:rPr')

    # set color
    color_elem = OxmlElement('w:color')
    color_elem.set(qn('w:val'), color)
    rPr.append(color_elem)

    # set underline
    if underline:
        u_elem = OxmlElement('w:u')
        u_elem.set(qn('w:val'), 'single')  # 单下划线
        rPr.append(u_elem)

    new_run.append(rPr)

    # create text element
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    text_elem.set(qn('xml:space'), 'preserve')

    new_run.append(text_elem)

    # add run to hyperlink
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def add_text(doc, item: dict, align: int = 0) -> None:
    """
    Add text to document.
    :param doc: docx document
    :param item: item's text
    :return: None
    """
    paragraph = doc.paragraphs
    is_first = True
    for node in item:
        if node['type'] == 'TEXT_NODE_TYPE_WORD':
            # font size
            fontSize = node['word']['font_size']

            # get bold
            if 'bold' in node['word']['style'].keys():
                is_blod = node['word']['style']['bold']
            else:
                is_blod = False

            words = node['word']['words']

            # get color
            r, g, b = get_rgb(node)

            if len(item) > 1 and is_first is False:
                run = paragraph.add_run(words)
                run.font.color.rgb = RGBColor(r, g, b)
                run.font.size = Pt(fontSize / 1.5)
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.bold = is_blod

            else:
                paragraph = doc.add_paragraph(words)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(r, g, b)
                    run.font.size = Pt(fontSize / 1.5)
                    run.font.name = "黑体"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.bold = is_blod

            is_first = False
        elif node['type'] == 'TEXT_NODE_TYPE_RICH':
            text = node['rich']['text']
            if 'jump_url' in node['rich'].keys():
                url = 'https://' + node['rich']['jump_url'].strip('http:').strip('https:').strip('//')
                add_hyperlink(paragraph, text, url)
            else:
                run = paragraph.add_run(text)
                run.font.color.rgb = RGBColor(47, 50, 56)
                run.font.size = Pt(17 / 1.5)
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    if align == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_image(doc, image_path_list) -> None:
    """
    Add img to document.
    :param doc: docx document
    :param image_path_list: image path list
    :return: None
    """
    section = doc.sections[0]

    available_width = (section.page_width - section.left_margin - section.right_margin)

    for image_path in image_path_list:
        img = Image.open(image_path)
        img.save(image_path)
        img.close()
        doc.add_picture(str(image_path), width=available_width)


def get_module(INITIAL_STATE: dict, module_name: str):
    INITIAL_STATE = INITIAL_STATE['detail']['modules']
    for item in INITIAL_STATE:
        if module_name in ''.join(list(item.keys())):
            return item
    return None


def add_auto_numbered_data(doc, items: list) -> None:
    """
    add auto numbered data.
    :param doc: docx document.
    :param items: items list
    :return: None
    """
    for nodes in items:
        order = nodes['order']
        nodes['nodes'][0]['word']['words'] = str(order) + '. ' + nodes['nodes'][0]['word']['words']
        add_text(doc, nodes['nodes'])


def add_code_text(doc, item: dict) -> None:
    content = item['code']['content']
    paragraph = doc.add_paragraph(content)
    for run in paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(14 / 2)
        run.font.color.rgb = RGBColor(47, 50, 56)


def add_link_card(doc, item: dict) -> None:
    """
    add link card.
    :param doc: docx document.
    :param item: items list
    :return: None
    """
    paragraph = doc.add_paragraph('\n')
    if 'opus' in item['link_card']['card']:
        title_and_url = item['link_card']['card']['opus']
    elif 'ugc' in item['link_card']['card']:
        title_and_url = item['link_card']['card']['ugc']

    add_hyperlink(paragraph, title_and_url['title'], title_and_url['jump_url'])
    doc.add_paragraph('\n')


def get_article(cookie: str, article_id: str, doc_storage_location: str = None, document_name: str = 'Document.doc',
                img_path: str = None) -> str:
    """
    get_article
    :param cookie: website's cookie information.
    :param article_id: article's url such as 1097430290934005800 (https://www.bilibili.com/opus/1097430290934005800?spm_id_from=333.1387.0.0)
    :param doc_storage_location: doc local storage in word.
    :param document_name: document name.
    :param img_path: img path
    :return: document text
    """
    if doc_storage_location is None:
        doc_storage_location = Path.cwd()
    else:
        doc_storage_location = Path(Path.cwd() / doc_storage_location)

    if img_path is None:
        img_path = Path.cwd() / "img"
    else:
        img_path = Path(img_path)

    html = get_article_html(cookie, article_id)
    INITIAL_STATE = re.findall(r'window.__INITIAL_STATE__=(.*);\(function', html)[0]
    INITIAL_STATE = json.loads(INITIAL_STATE)

    # with open('test.json', 'w', encoding='utf-8') as f:
    #     json.dump(INITIAL_STATE, f, ensure_ascii=False)

    title = get_module(INITIAL_STATE, 'title')
    author = get_module(INITIAL_STATE, 'author')
    module_top = get_module(INITIAL_STATE, 'module_top')
    module_content = get_module(INITIAL_STATE, 'module_content')

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    if title is not None:
        title = title['module_title']['text']
        doc_title = doc.add_heading(title, level=1)
        for run in doc_title.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(22 / 1.5)

    if author is not None:
        author_name = author['module_author']['name']
        author_time = author['module_author']['pub_time']

        if '编辑于' in author_time:
            author_time = datetime.strptime(author_time, '编辑于 %Y年%m月%d日 %H:%M')
        else:
            author_time = datetime.strptime(author_time, '%Y年%m月%d日 %H:%M')
    else:
        author_name = None
        author_time = None

    if author_name is not None and author_time is not None:
        doc_name = doc.add_paragraph(author_name)
        doc_time = doc.add_paragraph(author_time.strftime("%Y/%m/%d %H:%M"))

        for run in doc_name.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(17 / 1.5)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for run in doc_time.runs:
            run.font.color.rgb = RGBColor(148, 153, 160)
            run.font.size = Pt(13 / 1.5)

    top_pic_urls = list()
    if module_top is not None:
        module_tops = module_top['module_top']['display']['album']['pics']
        for pic in module_tops:
            top_pic_urls.append(pic['url'])

    add_image(doc, get_img(cookie, article_id, top_pic_urls, img_path))

    if module_content is not None:
        module_content = module_content['module_content']['paragraphs']

    for item in module_content:
        if item['para_type'] == 1:
            add_text(doc, item['text']['nodes'], item['align'])
        elif item['para_type'] == 2:
            pic_urls = list()
            pics = item['pic']['pics']
            for pic in pics:
                pic_urls.append(pic['url'])
            add_image(doc, get_img(cookie, article_id, pic_urls, img_path))
            logging.log(logging.INFO, 'add picture urls: {}'.format(pic_urls))
        elif item['para_type'] == 3:
            logging.log(logging.WARNING,'I don\'t know this para_type. If possible, please send this article to my GitHub ISSUES.(https://github.com/Guoziqi329/bilibili_crawling/issues)')
        elif item['para_type'] == 4:
            logging.log(logging.WARNING,'I don\'t know this para_type. If possible, please send this article to my GitHub ISSUES.(https://github.com/Guoziqi329/bilibili_crawling/issues)')
        elif item['para_type'] == 5:
            add_auto_numbered_data(doc, item['list']['items'])
            logging.log(logging.INFO, 'add auto_numbered_data')
        elif item['para_type'] == 6:
            add_link_card(doc, item)
            logging.log(logging.INFO,'add link_card')
        elif item['para_type'] == 7:
            add_code_text(doc, item)
            logging.log(logging.INFO,'add code_text')

    doc_storage_location.mkdir(parents=True, exist_ok=True)
    doc.save(str(doc_storage_location / document_name))

    logging.log(logging.INFO,f'{title}--{author_name}--{author_time} finished.')

    return "\n".join([s.text for s in doc.paragraphs])