"""
表格处理模块 - 提供用于在文档中添加和处理表格的功能

本模块提供了以下功能：
- 向文档添加表格（支持自动计数）
- 添加表格标题
- 设置表格样式和内容
- 处理表格行列数据
- 管理表格自动编号
"""

import os
import sys
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Any, Optional, Tuple, Union
import time
import logging

# 创建模块日志记录器
logger = logging.getLogger(__name__)

def log_exception(logger, exception, message="发生异常"):
    """记录异常信息，包括完整的堆栈跟踪
    
    Args:
        logger: 日志记录器
        exception: 异常对象
        message: 异常前缀消息
    """
    import traceback
    logger.error(f"{message}: {str(exception)}")
    logger.error(f"异常详情: {traceback.format_exc()}")

class TableConfig:
    """表格配置类，简化表格函数的参数传递"""
    
    def __init__(self, 
                 rows: int = 0,
                 cols: int = 0,
                 data: Optional[List[List[Any]]] = None,
                 headers: Optional[List[str]] = None,
                 caption: Optional[str] = None,
                 style: str = "Table Grid",
                 chapter_num: Optional[int] = None,
                 table_num: Optional[int] = None,
                 style_name: str = "表格标题",
                 auto_number: bool = True,
                 caption_position: str = "before",
                 table_prefix: Optional[str] = None,
                 merge_cells: Optional[List[Dict[str, Any]]] = None,
                 column_widths: Optional[List[Union[str, float]]] = None):
        """
        初始化表格配置
        
        Args:
            rows: 表格行数
            cols: 表格列数
            data: 表格数据，二维列表
            headers: 表头数据，列表
            caption: 表格标题
            style: 表格样式名称
            chapter_num: 章节编号（用于自动生成"表x-y"格式的编号）
            table_num: 表格编号，如果为None且auto_number=True，则自动使用当前计数
            style_name: 标题样式名称
            auto_number: 是否自动编号，如果为True，则使用表格管理器进行编号
            caption_position: 标题位置，"before"或"after"，默认为"before"（表格前面）
            table_prefix: 表格前缀，默认使用table_manager的前缀
            merge_cells: 需要合并的单元格列表，每个元素是包含row、col、row_span、col_span的字典
            column_widths: 列宽比例列表，可以是数字或百分比字符串，如["33%", "67%"]
        """
        self.rows = rows
        self.cols = cols
        self.data = data or []
        self.headers = headers
        self.caption = caption
        self.style = style
        self.chapter_num = chapter_num
        self.table_num = table_num
        self.style_name = style_name
        self.auto_number = auto_number
        self.caption_position = caption_position.lower()
        self.table_prefix = table_prefix
        self.merge_cells = merge_cells or []
        self.column_widths = column_widths
        
        # 记录日志
        logger.debug(f"创建表格配置: caption={caption}, chapter_num={chapter_num}, "
                    f"rows={rows}, cols={cols}, auto_number={auto_number}, "
                    f"column_widths={column_widths}")

    def calculate_dimensions(self) -> Tuple[int, int]:
        """计算表格的实际行数和列数"""
        # 计算列数
        cols = self.cols
        if self.headers:
            cols = max(cols, len(self.headers))
        
        if self.data:
            for row in self.data:
                cols = max(cols, len(row))
        
        # 计算行数
        rows = self.rows
        if rows <= 0:
            rows = len(self.data) + (1 if self.headers else 0)
            if rows <= 0:
                rows = 1
        
        logger.debug(f"计算的表格尺寸: {rows}行 x {cols}列")
        return rows, cols


class TableManager:
    """表格管理器类，用于管理表格编号"""
    
    def __init__(self, table_prefix="表"):
        """
        初始化表格管理器
        
        Args:
            table_prefix: 表格编号前缀，默认为"表"
        """
        # 按章节存储表格计数
        self.table_counters = {}
        # 存储最近使用的章节编号
        self.current_chapter = None
        # 表格前缀
        self.table_prefix = table_prefix
        logger.info(f"初始化表格管理器: 前缀='{table_prefix}'")
    
    def get_table_number(self, chapter_num):
        """
        获取指定章节的当前表格编号
        
        Args:
            chapter_num: 章节编号
            
        Returns:
            当前表格编号
        """
        # 如果章节不存在，则初始化为1
        if chapter_num not in self.table_counters:
            self.table_counters[chapter_num] = 1
        
        self.current_chapter = chapter_num
        
        # 记录日志
        count = self.table_counters[chapter_num]
        logger.debug(f"获取章节 {chapter_num} 的表格编号: {count}")
        
        return count
    
    def increment_table_number(self, chapter_num=None):
        """
        增加指定章节的表格计数
        
        Args:
            chapter_num: 章节编号，如果为None则使用最近的章节
            
        Returns:
            更新后的表格编号
        """
        # 如果未指定章节，使用当前章节
        if chapter_num is None:
            chapter_num = self.current_chapter
            logger.debug(f"使用当前章节 {chapter_num} 增加表格编号")
        
        # 确保章节存在
        if chapter_num not in self.table_counters:
            self.table_counters[chapter_num] = 1
            logger.debug(f"初始化章节 {chapter_num} 的表格计数为 1")
        else:
            # 递增计数器
            self.table_counters[chapter_num] += 1
            logger.debug(f"递增章节 {chapter_num} 的表格计数到 {self.table_counters[chapter_num]}")
        
        # 更新当前章节
        self.current_chapter = chapter_num
        
        return self.table_counters[chapter_num]
    
    def get_formatted_table_number(self, chapter_num=None):
        """
        获取格式化的表格编号字符串（如"表1-1"）
        
        Args:
            chapter_num: 章节编号，如果为None则使用最近的章节
            
        Returns:
            格式化的表格编号字符串
        """
        # 如果未指定章节，使用当前章节
        if chapter_num is None:
            chapter_num = self.current_chapter
        
        # 确保章节存在
        if chapter_num not in self.table_counters:
            self.table_counters[chapter_num] = 1
        
        # 修改格式，去除中间的空格，使用短横线连接
        formatted = f"{self.table_prefix}{chapter_num} - {self.table_counters[chapter_num]}"
        logger.debug(f"生成格式化表格编号: {formatted}")
        return formatted
    
    def set_table_prefix(self, prefix):
        """
        设置表格前缀
        
        Args:
            prefix: 表格前缀，如"表"、"Table"等
        """
        logger.debug(f"设置表格前缀: 从 '{self.table_prefix}' 到 '{prefix}'")
        self.table_prefix = prefix
        
    def reset(self):
        """重置表格编号（用于生成新文档时重置计数）"""
        logger.info("重置表格编号计数器")
        self.table_counters = {}
        self.current_chapter = None
        
    def with_temporary_state(self, chapter_num, table_num=None, prefix=None):
        """
        创建临时状态的上下文管理器
        
        Args:
            chapter_num: 章节编号
            table_num: 临时表格编号，如果为None则使用当前值
            prefix: 临时前缀，如果为None则使用当前值
            
        Returns:
            上下文管理器对象
        """
        logger.debug(f"创建临时状态上下文: chapter_num={chapter_num}, table_num={table_num}, prefix={prefix}")
        return _TableManagerContext(self, chapter_num, table_num, prefix)


class _TableManagerContext:
    """TableManager的上下文管理器，用于临时修改状态"""
    
    def __init__(self, manager, chapter_num, table_num=None, prefix=None):
        self.manager = manager
        self.chapter_num = chapter_num
        self.table_num = table_num
        self.prefix = prefix
        self.original_prefix = None
        self.original_counter = None
        
    def __enter__(self):
        # 保存原始状态
        self.original_prefix = self.manager.table_prefix
        self.original_counter = self.manager.table_counters.get(self.chapter_num, 1)
        
        # 设置临时状态
        if self.prefix is not None:
            self.manager.table_prefix = self.prefix
            
        if self.table_num is not None:
            self.manager.table_counters[self.chapter_num] = self.table_num
            
        logger.debug(f"进入临时状态: 章节={self.chapter_num}, "
                   f"编号={self.manager.table_counters[self.chapter_num]}, "
                   f"前缀='{self.manager.table_prefix}'")
        return self.manager
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        # 恢复原始状态
        self.manager.table_prefix = self.original_prefix
        self.manager.table_counters[self.chapter_num] = self.original_counter
        logger.debug("退出临时状态，已恢复原始值")


# 创建全局表格管理器实例
table_manager = TableManager()


def add_table_to_doc(doc, config: Optional[TableConfig] = None, **kwargs):
    """
    向文档添加表格并自动处理表格编号
    
    可以使用TableConfig对象或直接传入关键字参数
    
    Args:
        doc: 文档对象
        config: 表格配置对象
        **kwargs: 直接传入的配置参数，与TableConfig的参数相同
        
    Returns:
        tuple: (表格对象, 下一个表格编号)
    """
    start_time = time.time()
    try:
        # 记录函数调用
        logger.info(f"开始添加表格: caption={config.caption if config else kwargs.get('caption')}")
        
        # 如果未提供config，则从kwargs创建
        if config is None:
            logger.debug(f"从关键字参数创建TableConfig: {kwargs}")
            config = TableConfig(**kwargs)
        
        # 计算表格尺寸
        rows, cols = config.calculate_dimensions()
        logger.debug(f"表格尺寸: {rows}行 x {cols}列")
        
        # 自动编号处理
        next_table_num = None
        current_table_num = None  # 用于存储当前表格的编号
        
        if config.auto_number and not (config.chapter_num is not None and config.table_num is not None):
            if config.chapter_num is not None:
                # 按章节编号，获取当前表格编号
                current_table_num = table_manager.get_table_number(config.chapter_num)
                logger.debug(f"自动编号: 章节 {config.chapter_num} 的当前表格编号为 {current_table_num}")
            else:
                # 使用全局编号
                current_table_num = table_manager.get_table_number(1)  # 默认使用第1章
                logger.debug(f"自动编号: 使用全局编号 {current_table_num}")
        else:
            # 使用指定的表格编号
            current_table_num = config.table_num
            logger.debug(f"使用指定表格编号: {current_table_num}")
        
        # 构建表格标题
        full_caption = ""
        if config.caption:
            if config.chapter_num is not None and current_table_num is not None:
                # 使用表格管理器临时状态生成格式化编号
                with table_manager.with_temporary_state(
                        config.chapter_num, 
                        current_table_num, 
                        config.table_prefix) as temp_manager:
                    formatted_number = temp_manager.get_formatted_table_number(config.chapter_num)
                    full_caption = f"{formatted_number} {config.caption}"
                    logger.debug(f"生成表格标题: {full_caption}")
            else:
                # 仅使用标题
                full_caption = config.caption
                logger.debug(f"使用无编号标题: {full_caption}")
        
        # 添加表格标题（如果标题位置在表格前）
        if full_caption and config.caption_position == "before":
            logger.debug("添加表格前标题")
            caption_para = doc.add_paragraph(full_caption, style=config.style_name)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加表格
        logger.debug(f"添加表格: {rows}行 x {cols}列, 样式='{config.style}'")
        table = doc.add_table(rows=rows, cols=cols)
        table.style = config.style
        
        # 设置列宽
        if config.column_widths:
            logger.debug(f"设置列宽: {config.column_widths}")
        _set_column_widths(table, config.column_widths, doc)
        
        # 先合并单元格
        if config.merge_cells:
            logger.debug(f"合并单元格: {len(config.merge_cells)}个合并操作")
        _merge_table_cells(table, config.merge_cells)
        
        # 再填充表格内容
        if config.headers:
            logger.debug(f"填充表头: {len(config.headers)}列")
        if config.data:
            logger.debug(f"填充数据: {len(config.data)}行")
        _populate_table_content(table, config.headers, config.data)
        
        # 应用表格样式
        logger.debug("应用表格样式")
        _apply_table_style(table)
        
        # 添加表格标题（如果标题位置在表格后）
        if full_caption and config.caption_position == "after":
            logger.debug("添加表格后标题")
            caption_para = doc.add_paragraph(full_caption, style=config.style_name)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 递增表格编号（如果使用自动编号）
        if config.auto_number and config.chapter_num is not None:
            # 递增计数器
            next_table_num = table_manager.increment_table_number(config.chapter_num)
            logger.debug(f"递增表格编号: 章节 {config.chapter_num} 的下一个表格编号为 {next_table_num}")
        else:
            # 如果未使用自动编号，返回传入的表格编号加1
            if current_table_num is not None:
                next_table_num = current_table_num + 1
                logger.debug(f"手动递增表格编号: 下一个表格编号为 {next_table_num}")
        
        # 记录完成信息
        elapsed = time.time() - start_time
        logger.info(f"表格添加完成: '{full_caption}', 耗时: {elapsed:.2f}秒")
        return table, next_table_num
        
    except Exception as e:
        # 记录异常信息
        log_exception(logger, e, "添加表格失败")
        # 重新抛出异常
        raise


def _set_column_widths(table, column_widths, doc):
    """设置表格列宽"""
    if not column_widths:
        return
    
    try:
        # 导入需要的模块
        from docx.shared import Cm, Inches, Twips
        
        # 获取表格宽度（考虑页面宽度和边距）
        section = doc.sections[0]
        # 计算页面可用宽度(厘米)
        page_width_cm = section.page_width.cm
        left_margin_cm = section.left_margin.cm
        right_margin_cm = section.right_margin.cm
        available_width_cm = page_width_cm - left_margin_cm - right_margin_cm
        
        # 确保表格总宽度略小于可用宽度，避免溢出
        table_width_cm = available_width_cm * 0.98
        
        # 记录调试信息
        logger.debug(f"设置表格列宽 - 页面宽度: {page_width_cm:.2f}厘米, 左边距: {left_margin_cm:.2f}厘米, 右边距: {right_margin_cm:.2f}厘米")
        logger.debug(f"可用宽度: {available_width_cm:.2f}厘米, 表格宽度: {table_width_cm:.2f}厘米")
        logger.debug(f"列宽比例参数: {column_widths}")
        
        # 计算总比例
        total_ratio = 0
        ratios = []
        
        for width in column_widths:
            if isinstance(width, str) and "%" in width:
                # 处理百分比格式的宽度
                ratio = float(width.strip("%")) / 100.0
                ratios.append(ratio)
                total_ratio += ratio
            else:
                # 处理数值格式的宽度
                try:
                    ratio = float(width)
                    ratios.append(ratio)
                    total_ratio += ratio
                except (ValueError, TypeError):
                    # 无效的宽度值，使用均等宽度
                    logger.warning(f"无效的列宽值: {width}，使用均等宽度")
                    total_ratio = 1.0
                    ratios = [1.0/len(table.columns)] * len(table.columns)
                    break
        
        # 如果总比例不为1，归一化处理
        if total_ratio != 1.0 and total_ratio > 0:
            logger.debug(f"总比例为 {total_ratio}，进行归一化处理")
            ratios = [r / total_ratio for r in ratios]
        
        # 确保有足够的比例值，不足则补充均等值
        while len(ratios) < len(table.columns):
            remaining_ratio = 1.0 - sum(ratios)
            remaining_cols = len(table.columns) - len(ratios)
            ratios.append(remaining_ratio / remaining_cols)
        
        # 尝试更强制的方式设置列宽
        # 方式1: 使用twips单位(更精确)
        table_width_twips = int(table_width_cm * 567)  # 1厘米约等于567 twips
        
        # 禁用自动适应
        table.autofit = False
        table.allow_autofit = False
        
        # 先设置表格的总宽度
        try:
            table._tbl.xpath("./w:tblPr")[0].xpath("./w:tblW")[0].set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", str(table_width_twips))
            table._tbl.xpath("./w:tblPr")[0].xpath("./w:tblW")[0].set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "dxa")
        except Exception as e:
            logger.error(f"设置表格总宽度时出错: {str(e)}")
        
        # 应用列宽
        logger.debug("设置列宽:")
        for i in range(len(table.columns)):
            if i < len(ratios):
                # 计算该列应占表格宽度的比例
                col_ratio = ratios[i]
                # 计算列宽(twips)
                col_width_twips = int(table_width_twips * col_ratio)
                
                try:
                    # 直接设置列宽(XML方式)
                    grid_col = table._tbl.xpath('./w:tblGrid/w:gridCol')[i]
                    grid_col.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", str(col_width_twips))
                    
                    # 设置列属性
                    for cell in table.column_cells(i):
                        tc = cell._tc
                        tcW = tc.xpath('./w:tcPr/w:tcW')[0]
                        tcW.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", str(col_width_twips))
                        tcW.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "dxa")
                    
                    logger.debug(f"  列 {i+1}: 比例 {col_ratio:.4f}, 宽度 {col_width_twips/567:.2f}厘米 ({col_width_twips} twips)")
                except Exception as e:
                    logger.error(f"  设置列 {i+1} 宽度时出错: {str(e)}")
            else:
                # 如果列数超过了比例数，平均分配剩余宽度
                remaining_width = table_width_twips - sum(int(table_width_twips * r) for r in ratios[:i])
                remaining_cols = len(table.columns) - i
                col_width_twips = int(remaining_width / remaining_cols)
                
                try:
                    # 直接设置列宽(XML方式)
                    grid_col = table._tbl.xpath('./w:tblGrid/w:gridCol')[i]
                    grid_col.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", str(col_width_twips))
                    
                    # 设置列属性
                    for cell in table.column_cells(i):
                        tc = cell._tc
                        tcW = tc.xpath('./w:tcPr/w:tcW')[0]
                        tcW.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", str(col_width_twips))
                        tcW.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "dxa")
                    
                    logger.debug(f"  列 {i+1}: (剩余平均分配) 宽度 {col_width_twips/567:.2f}厘米 ({col_width_twips} twips)")
                except Exception as e:
                    logger.error(f"  设置列 {i+1} 宽度时出错: {str(e)}")
        
        # 禁用表格自动调整
        try:
            tbl_layout = table._element.xpath('./w:tblPr/w:tblLayout')
            if not tbl_layout:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                layout_element = OxmlElement('w:tblLayout')
                layout_element.set(qn('w:type'), 'fixed')
                table._element.xpath('./w:tblPr')[0].append(layout_element)
            else:
                tbl_layout[0].set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "fixed")
        except Exception as e:
            logger.error(f"设置表格布局固定时出错: {str(e)}")
        
        logger.debug(f"表格总宽度设置为: {table_width_twips/567:.2f}厘米 ({table_width_twips} twips)")
        
        # 设置表格整体宽度的备选方法
        try:
            table.width = Twips(table_width_twips)
        except Exception as e:
            logger.error(f"使用Twips设置表格宽度时出错: {str(e)}")
    
    except Exception as e:
        log_exception(logger, e, "设置表格列宽失败")
        
    # 继续模块的其他函数...
    # ... 使用类似的方式为每个函数添加日志记录


def _populate_table_content(table, headers, data):
    """填充表格内容"""
    row_index = 0
    
    # 确保表格是否已经有行
    if len(table.rows) == 0:
        return
    
    # 填充表头
    if headers:
        header_row = table.rows[row_index]
        for col_index, header_text in enumerate(headers):
            if col_index < len(header_row.cells):
                cell = header_row.cells[col_index]
                # 先清空单元格，再设置内容
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(header_text))
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_index += 1
    
    # 填充数据
    if data:
        for data_row in data:
            if row_index < len(table.rows):  # 确保行索引不超过表格实际行数
                table_row = table.rows[row_index]
                for col_index, cell_data in enumerate(data_row):
                    if col_index < len(table_row.cells):
                        cell = table_row.cells[col_index]
                        # 检查单元格是否有内容，已有内容的可能是合并的单元格
                        if not cell.text.strip():
                            _set_cell_content(cell, cell_data)
                row_index += 1


def _set_cell_content(cell, content):
    """设置单元格内容，处理文本和图片"""
    # 检查单元格数据是否为图片路径
    if isinstance(content, str) and (content.startswith("fig") or 
                                   content.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))):
        try:
            # 导入必要模块
            import os
            from docx.shared import Cm
            
            
            # 清空单元格内容
            cell.text = ""
            
            # 构建多种可能的图片路径
            image_path = content
            
            
            # 如果找到图片，添加到单元格
            if os.path.exists(image_path):
                p = cell.paragraphs[0]
                run = p.add_run()
                # 设置图片宽度为单元格宽度的80%
                cell_width = 3.0  # 默认值，单位厘米
                run.add_picture(image_path, width=Cm(cell_width * 0.8))
                # 居中对齐图片
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 图片未找到，使用文本代替
                p = cell.paragraphs[0]
                run = p.add_run(f"[图片: {content}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            # 处理图片添加失败的情况
            p = cell.paragraphs[0]
            run = p.add_run(f"[图片添加失败: {str(e)}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # 普通文本数据 - 使用更安全的方式设置
        # 先清空单元格
        cell.text = ""
        # 按换行符分割文本并添加
        if isinstance(content, str) and "\n" in content:
            lines = content.split("\n")
            for i, line in enumerate(lines):
                if i == 0:
                    # 第一行使用现有段落
                    p = cell.paragraphs[0]
                    p.add_run(line)
                else:
                    # 后续行添加新段落
                    p = cell.add_paragraph()
                    p.add_run(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 单行文本
            p = cell.paragraphs[0]
            p.add_run(str(content))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _merge_table_cells(table, merge_cells):
    """处理单元格合并"""
    if not merge_cells:
        return
        
    rows = len(table.rows)
    cols = len(table.columns) if rows > 0 else 0
    
    for merge_info in merge_cells:
        row = merge_info.get("row", 0)
        col = merge_info.get("col", 0)
        row_span = merge_info.get("row_span", 1)
        col_span = merge_info.get("col_span", 1)
        bold = merge_info.get("bold", False)
        text = merge_info.get("text", None)  # 可选的直接指定文本
        
        # 检查行列索引有效性
        if row < 0 or row >= rows or col < 0 or col >= cols:
            logger.warning(f"警告: 无效的合并单元格索引 row={row}, col={col}, 已忽略")
            continue
        
        # 检查跨度有效性
        if row + row_span > rows or col + col_span > cols:
            logger.warning(f"警告: 合并单元格超出表格范围 row={row}, col={col}, row_span={row_span}, col_span={col_span}, 已调整")
            row_span = min(row_span, rows - row)
            col_span = min(col_span, cols - col)
        
        # 获取目标单元格
        target_cell = table.cell(row, col)
        
        # 保存原始内容 (如果未指定文本)
        if text is None:
            target_text = target_cell.text
        else:
            # 使用指定的文本替换
            target_text = text
        
        # 清空将要合并的其他单元格内容，防止合并后出现多余的制表符
        for r in range(row, row + row_span):
            for c in range(col, col + col_span):
                if r == row and c == col:
                    continue  # 跳过目标单元格
                cell = table.cell(r, c)
                cell.text = ""
        
        # 合并行
        if row_span > 1:
            for i in range(1, row_span):
                target_cell = target_cell.merge(table.cell(row + i, col))
        
        # 合并列
        if col_span > 1:
            for i in range(1, col_span):
                target_cell = target_cell.merge(table.cell(row, col + i))
        
        # 合并后设置内容
        target_cell.text = ""  # 先清空
        
        # 按换行符分割文本并添加
        if "\n" in target_text:
            lines = target_text.split("\n")
            for i, line in enumerate(lines):
                if i == 0:
                    # 第一行使用现有段落
                    p = target_cell.paragraphs[0]
                    p.add_run(line)
                else:
                    # 后续行添加新段落
                    p = target_cell.add_paragraph()
                    p.add_run(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 单行文本
            p = target_cell.paragraphs[0]
            p.add_run(target_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 应用粗体样式
        if bold:
            for paragraph in target_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_table_style(table):
    """应用表格样式"""
    for row in table.rows:
        for cell in row.cells:
            # 设置单元格文本垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 设置所有段落对齐方式
            for paragraph in cell.paragraphs:
                if not any(run.bold for run in paragraph.runs):  # 保持表头的居中样式
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_data_table(doc, title, headers, data, chapter_num, table_count=None, table_prefix=None, 
                  merge_cells=None, column_widths=None):
    """
    添加数据表格到文档中（简化API，兼容旧版本）
    
    Args:
        doc: 文档对象
        title: 表格标题
        headers: 表头列表，可以为None
        data: 表格数据（二维列表）
        chapter_num: 章节编号
        table_count: 表格计数，如果为None则自动获取当前计数
        table_prefix: 表格前缀，如"表"、"Table"等
        merge_cells: 需要合并的单元格列表
        column_widths: 列宽比例列表
        
    Returns:
        下一个表格的编号
    """
    # 如果没有提供表格计数，从表格管理器获取当前计数
    if table_count is None:
        # 这里不做递增，因为add_table_to_doc中会递增
        current_count = table_manager.get_table_number(chapter_num)
    else:
        # 如果提供了表格计数，直接使用
        current_count = table_count
    
    # 处理column_widths参数，确保它是正确的格式
    if column_widths:
        # 规范化列宽值
        processed_widths = []
        for width in column_widths:
            if isinstance(width, str) and "%" in width:
                # 百分比格式保持原样
                processed_widths.append(width)
            elif isinstance(width, (int, float)):
                # 数值转为百分比字符串
                processed_widths.append(f"{width * 100:.0f}%")
            else:
                # 其他格式尝试转换为浮点数
                try:
                    value = float(width)
                    processed_widths.append(f"{value * 100:.0f}%")
                except (ValueError, TypeError):
                    # 无效值使用默认比例
                    pass
        
        if processed_widths:
            column_widths = processed_widths
            logger.debug(f"处理后的列宽参数: {column_widths}")
    
    config = TableConfig(
        data=data,
        headers=headers,
        caption=title,
        chapter_num=chapter_num,
        table_num=current_count,
        auto_number=(table_count is None),  # 只有在不提供table_count时才自动编号
        table_prefix=table_prefix,
        merge_cells=merge_cells,
        column_widths=column_widths
    )
    
    table, next_table_num = add_table_to_doc(doc, config)
    
    # 设置表格字体和单元格样式
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # 将所有段落设置为居中对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 设置首行缩进为0
                paragraph.paragraph_format.first_line_indent = 0
                # 设置行间距为单倍行距
                paragraph.paragraph_format.line_spacing = 1.0
                # 设置段前段后间距为0
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                
                for run in paragraph.runs:
                    # 设置中文字体为宋体
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
                    # 确保在非中文环境下也能正确显示
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    
    # 如果表格是两列结构且没有指定列宽，使用默认的33%和67%
    if len(table.columns) == 2 and not column_widths:
        logger.debug("检测到两列表格且未指定列宽，应用默认列宽比例: 33%/67%")
        _set_column_widths(table, ["33%", "67%"], doc)
    
    return next_table_num


def reset_table_numbering():
    """重置表格编号（用于生成新文档时重置计数）"""
    table_manager.reset()


def get_current_table_number(chapter_num=None):
    """
    获取当前表格编号
    
    Args:
        chapter_num: 章节编号，如果为None则使用最近的章节
        
    Returns:
        当前表格编号
    """
    if chapter_num is None:
        chapter_num = table_manager.current_chapter
        
    if chapter_num is None:
        return None
        
    return table_manager.get_table_number(chapter_num)


def set_table_column_widths(table, widths):
    """
    设置表格列宽
    
    Args:
        table: 表格对象
        widths: 列宽列表，单位为厘米
    """
    for i, width in enumerate(widths):
        if i < len(table.columns):
            table.columns[i].width = Cm(width)


def merge_table_cells(table, start_row, start_col, end_row, end_col):
    """
    合并表格单元格
    
    Args:
        table: 表格对象
        start_row: 起始行索引
        start_col: 起始列索引
        end_row: 结束行索引
        end_col: 结束列索引
    """
    # 检查索引是否有效
    if start_row < 0 or start_col < 0 or end_row >= len(table.rows) or end_col >= len(table.columns):
        logger.warning("警告: 表格合并索引超出范围")
        return
    
    # 获取合并区域的第一个单元格
    cell = table.cell(start_row, start_col)
    
    # 执行单元格合并
    cell.merge(table.cell(end_row, end_col))


def apply_cell_formatting(cell, text=None, bold=False, italic=False, alignment=None, 
                         vertical_alignment=None, shading=None, font_size=None):
    """
    应用单元格格式设置
    
    Args:
        cell: 单元格对象
        text: 要设置的文本（如果为None则保持原文本）
        bold: 是否加粗
        italic: 是否斜体
        alignment: 水平对齐方式
        vertical_alignment: 垂直对齐方式
        shading: 单元格底纹颜色
        font_size: 字体大小（磅）
    """
    # 设置文本（如果提供）
    if text is not None:
        cell.text = str(text)
    
    # 应用文本格式
    for paragraph in cell.paragraphs:
        # 水平对齐
        if alignment is not None:
            paragraph.alignment = alignment
        
        # 应用字体格式
        for run in paragraph.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if font_size is not None:
                run.font.size = Pt(font_size)
    
    # 设置垂直对齐
    if vertical_alignment is not None:
        cell.vertical_alignment = vertical_alignment
    
    # 设置底纹
    if shading is not None:
        cell._tc.get_or_add_tcPr().append(shading)


def add_auto_numbered_table(doc, title, headers=None, data=None, chapter_num=None, 
                            rows=None, cols=None, style="Table Grid", style_name="表格标题",
                            caption_position="before", table_prefix=None, column_widths=None):
    """
    添加带自动编号和标题的表格，更简便地实现表格的标准化添加
    
    Args:
        doc: 文档对象
        title: 表格标题
        headers: 表头列表，如果为None则不添加表头
        data: 表格数据（二维列表），如果为None则创建空表格
        chapter_num: 章节编号，如果为None则使用最近的章节
        rows: 表格行数，如果为None则根据data和headers自动计算
        cols: 表格列数，如果为None则根据headers或data第一行自动计算
        style: 表格样式名称
        style_name: 标题样式名称
        caption_position: 标题位置，"before"或"after"
        table_prefix: 表格前缀
        column_widths: 列宽列表（厘米），如果提供则自动设置列宽
        
    Returns:
        tuple: (表格对象, 下一个表格编号)
    
    Examples:
        # 基本用法
        table, _ = add_auto_numbered_table(
            doc=doc,
            title="水质监测数据",
            headers=["监测点", "pH值", "溶解氧(mg/L)"],
            data=[
                ["点位1", "7.2", "6.5"],
                ["点位2", "7.5", "6.8"]
            ],
            chapter_num=2
        )
        
        # 指定列宽的用法
        table, _ = add_auto_numbered_table(
            doc=doc,
            title="财务数据",
            headers=["项目", "金额(万元)"],
            data=[["项目1", "500"], ["项目2", "800"]],
            chapter_num=4,
            column_widths=[5.0, 3.0]  # 第一列5厘米宽，第二列3厘米宽
        )
    """
    # 自动计算行数和列数
    if rows is None:
        if data:
            rows = len(data) + (1 if headers else 0)
        else:
            rows = 1 if headers else 0
    
    if cols is None:
        if headers:
            cols = len(headers)
        elif data and data[0]:
            cols = len(data[0])
        else:
            cols = 0
    
    # 确保行列数合理
    if rows <= 0:
        rows = 1
    if cols <= 0:
        cols = 1
    
    # 添加表格
    table, next_table_num = add_table_to_doc(
        doc=doc,
        rows=rows,
        cols=cols,
        data=data,
        headers=headers,
        caption=title,
        style=style,
        chapter_num=chapter_num,
        table_num=None,  # 自动获取当前编号
        style_name=style_name,
        auto_number=True,
        caption_position=caption_position,
        table_prefix=table_prefix
    )
    
    # 设置列宽（如果提供）
    if column_widths and len(column_widths) > 0:
        set_table_column_widths(table, column_widths)
    
    return table, next_table_num

