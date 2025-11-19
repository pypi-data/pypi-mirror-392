"""
图像处理模块
提供图像加载和处理功能
"""

import io
import logging
import os
import platform
from typing import Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont

logger = logging.getLogger(__name__)


def get_system_chinese_fonts():
    """
    获取系统可用的中文字体路径列表

    Returns:
        list: 按优先级排序的字体路径列表
    """
    system = platform.system()
    font_paths = []

    if system == "Windows":
        font_paths.extend(
            [
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
                "C:/Windows/Fonts/simsun.ttc",  # 宋体
                "C:/Windows/Fonts/msyh.ttc",  # 微软雅黑
                "C:/Windows/Fonts/simkai.ttf",  # 楷体
                "C:/Windows/Fonts/simfang.ttf",  # 仿宋
            ]
        )
    elif system == "Linux":
        font_paths.extend(
            [
                "/root/.fonts/SimHei.ttf",  # SimHei
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",  # 文泉驿微米黑
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # 文泉驿正黑
                "/usr/share/fonts/truetype/arphic/ukai.ttc",  # AR PL UKai
                "/usr/share/fonts/truetype/arphic/uming.ttc",  # AR PL UMing
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",  # Noto Sans CJK
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",  # DejaVu Sans
                "/System/Fonts/PingFang.ttc",  # Ubuntu 可能的位置
            ]
        )
    elif system == "Darwin":  # macOS
        font_paths.extend(
            [
                "/System/Library/Fonts/PingFang.ttc",  # 苹方
                "/System/Library/Fonts/Hiragino Sans GB.ttc",  # 冬青黑体
                "/System/Library/Fonts/STHeiti Light.ttc",  # 华文黑体
            ]
        )

    return font_paths


def test_font_chinese_support(font_path, test_text="测试"):
    """
    测试字体是否支持中文

    Args:
        font_path: 字体文件路径
        test_text: 测试文本

    Returns:
        bool: 支持中文返回True
    """
    try:
        font = ImageFont.truetype(font_path, 20)
        test_img = Image.new("RGB", (100, 50))
        draw = ImageDraw.Draw(test_img)

        # 尝试绘制中文文本
        try:
            draw.textsize(test_text, font=font)
            return True
        except:
            try:
                # PIL 10.0.0+ 使用 textbbox
                draw.textbbox((0, 0), test_text, font=font)
                return True
            except:
                return False
    except:
        return False


class ImageProcessor:
    """图像处理器类"""

    def __init__(self, path_manager=None):
        """初始化图像处理器

        Args:
            path_manager: 路径管理器对象
        """
        self.path_manager = path_manager

    def load_image(self, image_path: str) -> Optional[Image.Image]:
        """加载图像

        Args:
            image_path: 图像文件路径

        Returns:
            PIL图像对象，失败返回None
        """
        try:
            return Image.open(image_path)
        except Exception as e:
            logger.error(f"加载图像失败: {str(e)}")
            return None

    def resize_image(
        self, image: Image.Image, width: int = None, height: int = None
    ) -> Image.Image:
        """调整图像大小

        Args:
            image: PIL图像对象
            width: 目标宽度（像素）
            height: 目标高度（像素）

        Returns:
            调整大小后的图像对象
        """
        if not width and not height:
            return image

        if width and height:
            return image.resize((width, height), Image.LANCZOS)

        # 等比例缩放
        w, h = image.size
        if width:
            new_height = int(h * width / w)
            return image.resize((width, new_height), Image.LANCZOS)
        else:
            new_width = int(w * height / h)
            return image.resize((new_width, height), Image.LANCZOS)

    def add_image_to_document(
        self,
        doc,
        image_path: str,
        width_cm: float = 12.0,
        caption: str = None,
        style_name: str = "Normal",
    ) -> bool:
        """将图像添加到文档

        Args:
            doc: 文档对象
            image_path: 图像文件路径
            width_cm: 图像宽度（厘米）
            caption: 图像说明文字
            style_name: 说明文字的样式名称

        Returns:
            是否成功添加
        """
        try:
            if not os.path.exists(image_path):
                logger.error(f"图像文件不存在: {image_path}")
                return False

            # 添加图像
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1  # 居中
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width_cm))

            # 添加说明文字
            if caption:
                caption_para = doc.add_paragraph(caption)
                caption_para.alignment = 1  # 居中
                caption_para.style = doc.styles[style_name]

            return True
        except Exception as e:
            logger.error(f"添加图像到文档失败: {str(e)}")
            return False


def generate_placeholder_image(
    width=10,
    height=7.5,
    bg_color=(240, 240, 240),
    text="用户手动添加",
    text_color=(150, 150, 150),
    font_size=30,
    return_bytes=True,
):
    """
    生成占位图片

    Args:
        width: 图片宽度（厘米）
        height: 图片高度（厘米）
        bg_color: 背景颜色（RGB元组）
        text: 显示的文本
        text_color: 文本颜色（RGB元组）
        font_size: 字体大小
        return_bytes: 是否返回字节流（否则返回PIL图片对象）

    Returns:
        BytesIO对象或PIL图片对象
    """
    # 转换厘米为像素（假设72DPI）
    width_px = int(width * 28.35)  # 1厘米约等于28.35像素
    height_px = int(height * 28.35)

    # 创建图片
    img = Image.new("RGB", (width_px, height_px), bg_color)
    draw = ImageDraw.Draw(img)

    # 尝试加载字体
    font = None
    chinese_supported = False

    try:
        # 获取系统中文字体路径
        font_paths = get_system_chinese_fonts()

        # 尝试加载支持中文的字体
        for path in font_paths:
            if os.path.exists(path):
                try:
                    test_font = ImageFont.truetype(path, font_size)
                    if test_font_chinese_support(path):
                        font = test_font
                        chinese_supported = True
                        logger.info(f"成功加载中文字体: {os.path.basename(path)}")
                        break
                except Exception as e:
                    logger.debug(f"字体加载失败 {os.path.basename(path)}: {str(e)}")
                    continue

        # 如果没有找到中文字体，使用默认字体并转换文本
        if font is None:
            font = ImageFont.load_default()
            chinese_supported = False
            logger.warning(f"未找到中文字体，使用默认字体。系统: {platform.system()}")

    except Exception as e:
        logger.error(f"字体加载出错: {str(e)}")
        font = ImageFont.load_default()
        chinese_supported = False

    # 如果不支持中文，将文本转换为英文
    if not chinese_supported:
        if "用户手动添加" in text:
            text = "Please add manually"
        elif "图片加载失败" in text:
            text = text.replace("图片加载失败", "Image load failed")
        logger.info(f"中文字体不可用，文本已转换为英文: {text}")

    # 计算文本大小以居中显示
    try:
        text_width, text_height = draw.textsize(text, font=font)
    except:
        # PIL 9.0.0及更高版本中的变化
        text_width, text_height = draw.textbbox((0, 0), text, font=font)[2:4]

    position = ((width_px - text_width) // 2, (height_px - text_height) // 2)

    # 绘制文本
    draw.text(position, text, font=font, fill=text_color)

    if return_bytes:
        # 转换为字节流
        img_bytes = io.BytesIO()
        img.save(img_bytes, format="PNG")
        img_bytes.seek(0)
        return img_bytes
    else:
        return img


def is_image_file(file_path):
    """
    判断文件是否为图片文件

    Args:
        file_path: 文件路径

    Returns:
        是图片文件返回True，否则返回False
    """
    # 首先检查文件是否存在且不是目录
    if not file_path or not os.path.exists(file_path) or os.path.isdir(file_path):
        return False

    # 检查文件扩展名
    valid_extensions = [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"]
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext not in valid_extensions:
        return False

    # 尝试用PIL打开文件
    try:
        img = Image.open(file_path)
        img.verify()  # 验证图片文件
        return True
    except Exception:
        return False


def add_image_to_doc(
    doc,
    image_path,
    width=None,
    height=None,
    caption=None,
    chapter_num=None,
    figure_count=None,
    style_name="图片标题",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    dpi=72,
):
    """
    向文档添加图片，如果图片不存在则添加占位图片

    Args:
        doc: 文档对象
        image_path: 图片路径
        width: 图片宽度（厘米，如果为None则使用图片原始宽度）
        height: 图片高度（厘米，如果为None则使用图片原始高度）
        caption: 图片标题
        chapter_num: 章节编号（用于自动生成"图x.y"格式的编号）
        figure_count: 图片编号
        style_name: 标题样式名称
        alignment: 图片对齐方式
        dpi: 图片分辨率（用于计算尺寸）

    Returns:
        如果提供了chapter_num和figure_count，则返回更新后的figure_count，否则返回None
    """
    # 检查图片是否存在且是有效的图片文件
    image_exists = False
    if image_path == "skip":
        logger.info("水质等级图片不存在，跳过")
        return figure_count
    if image_path and os.path.exists(image_path) and not os.path.isdir(image_path):
        # 进一步验证是否为有效图片文件
        try:
            with Image.open(image_path) as img:
                # 如果能打开，说明是有效图片
                image_exists = True
        except Exception as e:
            logger.warning(f"警告: 文件 '{image_path}' 不是有效的图片文件: {str(e)}")

    # 创建一个专用于图片的段落
    paragraph = doc.add_paragraph()

    # 完全重置所有段落格式设置，确保图片能完整显示
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.right_indent = 0
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = 0

    # 设置段落对齐方式
    paragraph.alignment = alignment

    # 处理图片插入
    run = paragraph.add_run()

    # 如果图片不存在，创建占位图片
    if not image_exists:
        logger.warning(f"警告: 图片 '{image_path}' 不存在，将使用占位图片替代")
        image_data = generate_placeholder_image(
            width=width or 10,
            height=width * 0.75 if width else 7.5,  # 使用4:3比例
            text="用户手动添加",
        )

        # 只指定宽度
        if width:
            run.add_picture(image_data, width=Cm(width))
        else:
            run.add_picture(image_data, width=Cm(10))
    else:
        # 添加实际图片到文档
        try:
            # 只指定宽度，保持原始比例
            if width:
                run.add_picture(image_path, width=Cm(width))
            else:
                # 使用原始尺寸
                run.add_picture(image_path)
            logger.info(f"成功添加图片: {image_path}")
        except Exception as e:
            logger.error(f"错误: 添加图片 '{image_path}' 时出现问题: {str(e)}")
            # 使用占位图片代替
            image_bytes = generate_placeholder_image(
                width=width or 10,
                height=None,
                text=f"图片加载失败: {os.path.basename(image_path)}",
            )

            run.add_picture(image_bytes, width=Cm(width or 10))

    # 添加图片标题
    if caption and chapter_num is not None and figure_count is not None:
        caption_para = doc.add_paragraph(f"图 {chapter_num} - {figure_count} {caption}")
        caption_para.style = doc.styles[style_name]
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 增加图片计数
        updated_figure_count = figure_count + 1
        return updated_figure_count
    elif caption:
        # 如果没有提供章节和图片编号，仅添加标题
        caption_para = doc.add_paragraph(caption)
        caption_para.style = doc.styles[style_name]
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 如果提供了chapter_num和figure_count但没有caption，仍然增加计数
    if chapter_num is not None and figure_count is not None:
        return figure_count + 1

    return None


def add_watermark(doc, image_path, opacity=0.1):
    """
    向文档添加水印

    Args:
        doc: 文档对象
        image_path: 水印图片路径
        opacity: 水印透明度（0-1之间的浮点数）

    Returns:
        成功返回True，失败返回False
    """
    # 检查图片是否存在
    if not os.path.exists(image_path):
        logger.warning(f"警告: 水印图片 '{image_path}' 不存在")
        return False

    try:
        # 访问文档的节
        section = doc.sections[0]  # 默认使用第一个节

        # 获取页眉
        header = section.header

        # 在页眉中添加一个段落
        paragraph = (
            header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        )

        # 设置段落属性
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加水印图片
        run = paragraph.add_run()

        # 设置图片的尺寸（这里根据页面宽度调整）
        width_cm = 15  # 水印宽度（厘米）

        # 使用PIL处理水印透明度
        try:
            with Image.open(image_path) as img:
                # 转换为RGBA以支持透明度
                if img.mode != "RGBA":
                    img = img.convert("RGBA")

                # 调整透明度
                data = img.getdata()
                new_data = []
                for item in data:
                    # 调整Alpha通道
                    new_data.append(
                        (
                            item[0],
                            item[1],
                            item[2],
                            int(item[3] * opacity)
                            if len(item) > 3
                            else int(255 * opacity),
                        )
                    )

                img.putdata(new_data)

                # 保存为临时文件
                temp_watermark = io.BytesIO()
                img.save(temp_watermark, format="PNG")
                temp_watermark.seek(0)

                # 添加处理后的水印图片
                run.add_picture(temp_watermark, width=Cm(width_cm))
        except Exception as e:
            logger.warning(f"警告: 处理水印透明度时出错: {str(e)}，将直接使用原始图片")
            # 直接使用原始图片
            run.add_picture(image_path, width=Cm(width_cm))

        return True
    except Exception as e:
        logger.error(f"错误: 添加水印时出现问题: {str(e)}")
        return False


def load_image_with_fallback(
    image_path, placeholder_width=10, placeholder_height=7.5, placeholder_text=None
):
    """
    加载图片，如果不存在则返回占位图片

    Args:
        image_path: 图片路径
        placeholder_width: 占位图片宽度（厘米）
        placeholder_height: 占位图片高度（厘米）
        placeholder_text: 占位图片文本（如果为None则使用默认文本）

    Returns:
        BytesIO对象
    """
    if os.path.exists(image_path):
        # 如果图片存在，返回文件路径
        return image_path
    else:
        # 如果图片不存在，创建占位图片
        text = placeholder_text or f"图片未找到: {os.path.basename(image_path)}"
        return generate_placeholder_image(
            width=placeholder_width, height=placeholder_height, text=text
        )
