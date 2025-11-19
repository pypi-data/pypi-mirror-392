#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档页面工具模块 - 提供页面设置、页眉页脚处理等功能
"""
import logging
import os
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ..utils.text import is_chinese
import sys
from docx import Document
import time
# 从配置模块导入相关设置
from ..config.fonts import SONG_FONT_NAMES
from ..config.styles import PAGE_SETTINGS

# 获取模块日志记录器
logger = logging.getLogger(__name__)

def set_page_size_to_a4(doc):
    """设置文档页面大小为A4"""
    section = doc.sections[0]
    # 设置页面大小为A4
    section.page_height = PAGE_SETTINGS["page_height"]
    section.page_width = PAGE_SETTINGS["page_width"]
    # 设置页边距
    section.left_margin = PAGE_SETTINGS["left_margin"]
    section.right_margin = PAGE_SETTINGS["right_margin"]
    section.top_margin = PAGE_SETTINGS["top_margin"]
    section.bottom_margin = PAGE_SETTINGS["bottom_margin"]
    # 设置页眉页脚距离
    section.header_distance = PAGE_SETTINGS["header_distance"]
    section.footer_distance = PAGE_SETTINGS["footer_distance"]

    logger.info("已将文档纸张大小设置为A4格式")
    return doc

def add_page_number(run):
    """添加页码域码"""
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._element.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    
    run.font.size = Pt(10)
    run.font.name = SONG_FONT_NAMES[0]


def add_header_footer(doc, header_text=None, company_name=None):
    """为文档添加页眉和页脚"""
    logger.info("添加页眉和页脚...")
    
    # 对每个section添加页眉和页脚
    for i, section in enumerate(doc.sections):
        # 第一节（封面和公司简介）不显示页眉页脚和页码
        if i == 0:
            # 设置不显示页码 - 使用XML方式设置
            section.different_first_page_header_footer = False  # 修改这里，使整个节都不显示页码
            
            # 使用XML直接设置页码格式为none
            sectPr = section._sectPr
            pgNumType_elements = sectPr.xpath('./w:pgNumType')
            
            if pgNumType_elements:
                # 如果已存在，设置格式为none
                pgNumType = pgNumType_elements[0]
                # 删除fmt属性实现隐藏页码
                if pgNumType.get(qn('w:fmt')):
                    del pgNumType.attrib[qn('w:fmt')]
            else:
                # 如果不存在，创建一个新元素但不设置格式
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            
            # 彻底清空页眉
            header = section.header
            header.is_linked_to_previous = False  # 确保与前一节断开链接
            # 清空所有段落内容
            for p in header.paragraphs:
                p.clear()
                # 删除所有运行对象
                for run in p.runs:
                    p._p.remove(run._r)
            
            # 彻底清空页脚
            footer = section.footer
            footer.is_linked_to_previous = False  # 确保与前一节断开链接
            # 清空所有段落内容
            for p in footer.paragraphs:
                p.clear()
                # 删除所有运行对象
                for run in p.runs:
                    p._p.remove(run._r)
            
            continue
        
        # 第二节（目录）只显示大写罗马数字页码，不显示页眉
        elif i == 1:
            # 设置页码从1开始，使用大写罗马数字
            section.different_first_page_header_footer = False
            
            # 使用XML设置页码格式为大写罗马数字
            sectPr = section._sectPr
            pgNumType_elements = sectPr.xpath('./w:pgNumType')
            
            if pgNumType_elements:
                pgNumType = pgNumType_elements[0]
            else:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            
            pgNumType.set(qn('w:start'), '1')
            pgNumType.set(qn('w:fmt'), 'upperRoman')  # 设置为大写罗马数字
            
            # 清空页眉
            header = section.header
            header.is_linked_to_previous = False  # 确保与前一节断开链接
            for p in header.paragraphs:
                p.clear()
            
            # 设置页脚（大写罗马数字页码）
            footer = section.footer
            footer.is_linked_to_previous = False  # 确保与前一节断开链接
            for p in footer.paragraphs:
                p.clear()
            
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 页码格式：- I -
            run = footer_paragraph.add_run("- ")
            run.font.size = Pt(10)
            run.font.name = SONG_FONT_NAMES[0]
            
            # 添加页码域
            run = footer_paragraph.add_run()
            add_page_number(run)
            
            run = footer_paragraph.add_run(" -")
            run.font.size = Pt(10)
            run.font.name = SONG_FONT_NAMES[0]
        
        # 第三节及以后（正文章节）显示页眉和阿拉伯数字页码
        else:
            # 设置页码从1开始，使用阿拉伯数字
            section.different_first_page_header_footer = False
            
            # 使用XML设置页码格式为阿拉伯数字
            sectPr = section._sectPr
            pgNumType_elements = sectPr.xpath('./w:pgNumType')
            
            if pgNumType_elements:
                pgNumType = pgNumType_elements[0]
            else:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            
            pgNumType.set(qn('w:start'), '1')
            pgNumType.set(qn('w:fmt'), 'decimal')  # 设置为阿拉伯数字
            
            # 设置页眉
            header = section.header
            header.is_linked_to_previous = False  # 确保与前一节断开链接
            for p in header.paragraphs:
                p.clear()
            
            if header_text or company_name:
                header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 创建页眉文本
                if company_name:
                    if header_text:
                        text = f"{company_name} - {header_text}"
                    else:
                        text = company_name
                else:
                    text = header_text if header_text else ""
                
                # 使用混合字体添加页眉
                if text:
                    current_text = ""
                    current_is_chinese = None
                    
                    def add_current_text():
                        nonlocal current_text, current_is_chinese
                        if current_text:
                            run = header_paragraph.add_run(current_text)
                            run.font.size = Pt(9)  # 小五号字体
                            
                            if current_is_chinese:
                                run.font.name = SONG_FONT_NAMES[0]
                                rpr = run._element.get_or_add_rPr()
                                rpr.rFonts.set(qn('w:eastAsia'), SONG_FONT_NAMES[0])
                            else:
                                run.font.name = "Times New Roman"
                                rpr = run._element.get_or_add_rPr()
                                rpr.rFonts.set(qn('w:ascii'), "Times New Roman")
                                rpr.rFonts.set(qn('w:hAnsi'), "Times New Roman")
                        
                        current_text = ""
                        current_is_chinese = None
                    
                    for char in text:
                        char_is_chinese = is_chinese(char)
                        
                        if current_is_chinese is None:
                            current_is_chinese = char_is_chinese
                            current_text += char
                        elif current_is_chinese == char_is_chinese:
                            current_text += char
                        else:
                            add_current_text()
                            current_is_chinese = char_is_chinese
                            current_text = char
                    
                    add_current_text()
                    
                    # 添加页眉下划线
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')  # 设置线宽为6磅
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), 'auto')
                    pBdr.append(bottom)
                    
                    # 将边框添加到段落属性中
                    header_paragraph._p.get_or_add_pPr().append(pBdr)
            
            # 设置页脚（阿拉伯数字页码）
            footer = section.footer
            footer.is_linked_to_previous = False  # 确保与前一节断开链接
            for p in footer.paragraphs:
                p.clear()
            
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 页码格式：- 1 -
            run = footer_paragraph.add_run("- ")
            run.font.size = Pt(10)
            run.font.name = SONG_FONT_NAMES[0]
            
            # 添加页码域
            run = footer_paragraph.add_run()
            add_page_number(run)
            
            run = footer_paragraph.add_run(" -")
            run.font.size = Pt(10)
            run.font.name = SONG_FONT_NAMES[0]
    
    # 再次确保第一节不显示页眉页脚
    if doc.sections:
        first_section = doc.sections[0]
        # 使用XML直接设置页码格式为none
        sectPr = first_section._sectPr
        pgNumType_elements = sectPr.xpath('./w:pgNumType')
        
        if pgNumType_elements:
            pgNumType = pgNumType_elements[0]
            # 删除fmt属性实现隐藏页码
            if pgNumType.get(qn('w:fmt')):
                del pgNumType.attrib[qn('w:fmt')]
        else:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        
        # 彻底清空页眉
        header = first_section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        
        # 彻底清空页脚
        footer = first_section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
    
    logger.info("页眉和页脚添加完成")
    return doc

def add_toc(doc):
    """添加目录"""
    # 确保第一节没有页眉页脚（保存当前状态）
    first_section_has_no_header_footer = False
    if doc.sections and len(doc.sections) > 0:
        first_section = doc.sections[0]
        # 检查第一节的页眉页脚是否为空
        if first_section.header and not any(p.text for p in first_section.header.paragraphs) and \
           first_section.footer and not any(p.text for p in first_section.footer.paragraphs):
            first_section_has_no_header_footer = True
    
    # 添加居中显示的目录标题
    toc_title = doc.add_paragraph("目录", style="一级标题")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加段落
    paragraph = doc.add_paragraph()
    
    # 创建TOC域代码
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    
    # 设置TOC指令，使用大纲级别
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # 使用大纲级别方式创建目录，添加\o "1-3"参数指定只包含1-3级标题
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)
    
    # 结束域代码
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    
    # 添加提示文本（如果无法自动更新域时显示）
    from docx.shared import RGBColor
    hint_para = doc.add_paragraph("如未自动更新目录，请右键点击并选择\"更新域\"或按F9键更新目录。", style="正文")
    hint_para.runs[0].italic = True
    hint_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # 设置中文字体
    for run in hint_para.runs:
        run.font.name = SONG_FONT_NAMES[0]
        # 设置东亚文字字体
        rpr = run._element.get_or_add_rPr()
        rpr.rFonts.set(qn('w:eastAsia'), SONG_FONT_NAMES[0])
    
    # 添加分节符，而非分页符
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    
    # 恢复第一节的页眉页脚设置（如果之前是没有的）
    if first_section_has_no_header_footer and doc.sections and len(doc.sections) > 0:
        first_section = doc.sections[0]
        
        # 清空页眉
        header = first_section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        
        # 清空页脚
        footer = first_section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        
        # 使用XML直接设置页码格式为none
        sectPr = first_section._sectPr
        pgNumType_elements = sectPr.xpath('./w:pgNumType')
        
        if pgNumType_elements:
            # 如果已存在，设置格式为none（通过删除格式属性实现）
            pgNumType = pgNumType_elements[0]
            if pgNumType.get(qn('w:fmt')):
                del pgNumType.attrib[qn('w:fmt')]
        else:
            # 如果不存在，创建一个新元素但不设置格式
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
    
    return doc

def add_section_break(doc, continuous=False):
    """添加分节符"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    if continuous:
        run._r.add_br(WD_BREAK.SECTION_CONTINUOUS)
    else:
        run._r.add_br(WD_BREAK.NEW_PAGE)
    return doc 


def add_watermark_to_doc(doc_path, watermark_text="无锡谱视界科技有限公司", font_size=65, color=(200, 0, 0), diagonal=True):
    """
    使用Spire.Doc为Word文档添加文本水印
    
    参数:
        doc_path (str): Word文档的路径
        watermark_text (str): 水印文本内容，默认为"保密文件"
        font_size (int): 水印字体大小，默认为65
        color (tuple): 水印颜色，RGB格式，默认为(200, 0, 0)，即暗红色
        diagonal (bool): 是否使用对角线布局，默认为True
                        True表示对角线布局，False表示水平布局
    
    返回:
        bool: 操作是否成功
    """
    try:
        # 导入Spire.Doc库
        from spire.doc import Document
        from spire.doc import TextWatermark, Color, WatermarkLayout
        
        # 加载文档
        doc = Document()
        doc.LoadFromFile(doc_path)
        
        # 创建水印
        watermark = TextWatermark()
        watermark.Text = watermark_text
        watermark.FontSize = font_size
        
        # 设置水印颜色
        r, g, b = color
        watermark.Color = Color.get_Red()  # 暂时使用红色，避免FromArgb问题
        
        # 设置水印布局（使用枚举值）
        watermark.Layout = WatermarkLayout.Diagonal if diagonal else WatermarkLayout.Horizontal
        
        # 添加水印到文档
        doc.Watermark = watermark
        
        # 保存文档
        doc.SaveToFile(doc_path)
        doc.Close()
        
        return True
        
    except ImportError:
        logger.error("错误: 未找到Spire.Doc库，请使用 'pip install spire.doc' 安装")
        return False
    except Exception as e:
        logger.error(f"添加水印时发生错误: {e}")
        return False 

def update_fields_in_word(doc_path, retry_count=3):
    """
    使用COM自动化更新Word文档中的所有域，包括目录
    """
    logger.info(f"正在自动更新文档中的所有域...")
    
    try:
        if sys.platform == 'win32':
            # Windows系统使用COM自动化
            try:
                import win32com.client
            except ImportError:
                logger.error("Windows系统需要安装pywin32依赖才能使用目录自动更新功能")
                logger.info("解决方案:")
                logger.info("1. pip install pywin32")
                logger.info("2. pip install aerospot-autoreport[windows]")
                logger.info("3. pip install aerospot-autoreport[all]")
                return False
            
            # 处理WSL路径转换
            windows_path = doc_path
            if doc_path.startswith('/mnt/'):
                parts = doc_path.split('/')
                if len(parts) >= 3 and parts[2].isalpha():
                    drive_letter = parts[2].upper()
                    windows_path = f"{drive_letter}:\\" + "\\".join(parts[3:])
                    logger.info(f"转换WSL路径: {doc_path} -> {windows_path}")
            
            # 确保使用绝对路径
            windows_path = os.path.abspath(windows_path)
            logger.info(f"使用文件路径: {windows_path}")
            
            for _ in range(retry_count):
                try:
                    word = win32com.client.DispatchEx("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = False
                    doc = word.Documents.Open(windows_path)
                    
                    # 更新所有域
                    for i in range(doc.Fields.Count):
                        doc.Fields(i+1).Update()
                    
                    # 更新目录
                    try:
                        doc.TablesOfContents(1).Update()
                    except:
                        logger.warning("目录更新失败，可能文档中没有目录")
                    
                    # 设置目录字体为宋体
                    logger.info("正在设置目录字体为宋体...")
                    if doc.TablesOfContents.Count > 0:
                        toc = doc.TablesOfContents(1)
                        toc.Range.Font.Name = SONG_FONT_NAMES[0]
                        toc.Range.Font.NameAscii = SONG_FONT_NAMES[0]
                        toc.Range.Font.NameFarEast = SONG_FONT_NAMES[0]
                    
                    # 保存文档
                    doc.Save()
                    doc.Close()
                    word.Quit()
                    
                    logger.info(f"域更新完成！目录字体已设置为宋体。")
                    return True
                
                except Exception as e:
                    logger.error(f"更新域时出错 (尝试 {_+1}/{retry_count}): {str(e)}")
                    time.sleep(1)
                    
                finally:
                    try:
                        doc.Close(False)
                    except:
                        pass
                    try:
                        word.Quit()
                    except:
                        pass
            
            logger.warning("无法使用COM自动化更新域")
            return False
        
        else:
            # 非Windows系统使用替代方法 - 这里只设置目录字体，不更新域
            logger.info("非Windows系统，使用替代方法设置目录字体...")
            try:
                doc = Document(doc_path)
                # 查找并修改TOC样式的字体
                for i in range(1, 4):
                    toc_style_name = f'TOC{i}'
                    if toc_style_name in doc.styles:
                        style = doc.styles[toc_style_name]
                        font = style.font
                        font.name = SONG_FONT_NAMES[0]
                        font.size = Pt(10.5)  # 五号字体
                        
                        # 设置东亚文字字体
                        rpr = style._element.get_or_add_rPr()
                        if hasattr(rpr, 'rFonts'):
                            rpr.rFonts.set(qn('w:eastAsia'), SONG_FONT_NAMES[0])
                
                doc.save(doc_path)
                logger.info("目录字体已设置为宋体。")
                return True
            
            except Exception as e:
                logger.error(f"设置目录字体时出错: {str(e)}")
                return False
    
    except Exception as e:
        logger.error(f"更新域失败: {str(e)}")
        return False