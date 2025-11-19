"""
报告生成器模块
负责根据配置生成Word格式的报告
"""
import os
import logging
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK
from .document.tables import add_data_table
from .document.paragraphs import add_paragraph_with_style
from .document.styles import create_styles, create_toc_styles
from .document.paragraphs import add_mixed_font_paragraph
from .utils.font import set_default_font
from .config.company_info import DEFAULT_COMPANY_INFO

from .document.pages import (
    # 样式相关
    # 页面设置相关
    set_page_size_to_a4,
    add_toc,
    add_header_footer,
    update_fields_in_word,
    add_watermark_to_doc
)


from .document.images import add_image_to_doc

logger = logging.getLogger(__name__)

class ReportGenerator:
    """报告生成器类"""
    
    def __init__(self, output_path: str, report_structure: Dict[str, Any], update_data: dict):
        """
        初始化报告生成器
        
        Args:
            output_dir: 输出目录
            company_info: 公司信息配置
            report_structure: 报告结构配置
        """
        self.output_path = output_path
        self.report_structure = report_structure
        self.doc = None
        self.update_data = update_data
    
    def generate(self) -> Optional[str]:
        """
        生成完整的报告
        
        Returns:
            str: 生成的报告文件路径，失败返回None
        """
        try:
            # 创建新文档
            self.doc = Document()
            
            # 设置文档样式
            self._setup_document_styles()
            
            # 生成报告内容
            self._create_cover_page()
            self._create_company_profile()
            self._create_toc()
            self._create_chapters()
            
            add_header_footer(self.doc, self.report_structure["title"], self.update_data.get('company', {}).get('name'))

            # 特别确保第一节不显示页码（封面和公司简介）
            if self.doc.sections:
                first_section = self.doc.sections[0]
                
                # 使用XML直接设置页码格式为none
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                
                sectPr = first_section._sectPr
                pgNumType_elements = sectPr.xpath('./w:pgNumType')
                
                if pgNumType_elements:
                    pgNumType = pgNumType_elements[0]
                    # 删除fmt属性实现隐藏页码
                    if pgNumType.get(qn('w:fmt')):
                        del pgNumType.attrib[qn('w:fmt')]
                else:
                    pgNumType = OxmlElement('w:pgNumType')
                    sectPr.append(pgNumType)
                
                # 彻底清空第一节的页眉页脚
                header = first_section.header
                header.is_linked_to_previous = False
                for p in header.paragraphs:
                    p.clear()
                footer = first_section.footer
                footer.is_linked_to_previous = False
                for p in footer.paragraphs:
                    p.clear()

            # 保存文档
            logger.info(f"保存报告到: {self.output_path}")
            # 确保输出目录存在
            self.doc.save(self.output_path)
            
            # 临时文件用于处理水印
            temp_output_file = None
            
            # 检查是否需要添加水印
            if self.update_data.get("company_info").get("watermark_enabled", False) and self.update_data.get("company_info").get("watermark_use_spire", False):
                try:                    
                    # 创建临时文件用于处理
                    import tempfile
                    temp_fd, temp_output_file = tempfile.mkstemp(suffix='.docx')
                    os.close(temp_fd)
                    
                    # 复制原始文件到临时文件
                    import shutil
                    shutil.copy2(self.output_path, temp_output_file)
                    
                    # 获取水印配置
                    watermark_text = self.update_data.get("company_info").get("watermark_text", "无锡谱视界科技有限公司")
                    watermark_size = self.update_data.get("company_info").get("watermark_size", 65)
                    watermark_color = self.update_data.get("company_info").get("watermark_color", (200, 0, 0))
                    watermark_diagonal = self.update_data.get("company_info").get("watermark_diagonal", True)
                    
                    # 添加水印到临时文件
                    success = add_watermark_to_doc(
                        doc_path=temp_output_file,
                        watermark_text=watermark_text,
                        font_size=watermark_size,
                        color=watermark_color,
                        diagonal=watermark_diagonal
                    )
                    
                    if success:
                        logger.info(f"已使用Spire.Doc添加水印: {watermark_text}")
                        
                        # 再次打开临时文件，删除Spire.Doc添加的评估警告
                        doc = Document(temp_output_file)
                        
                        # 检查首页第一段是否为评估警告
                        if doc.paragraphs and "Evaluation Warning" in doc.paragraphs[0].text:
                            # 删除评估警告段落
                            p = doc.paragraphs[0]._element
                            p.getparent().remove(p)
                            logger.info("已删除Spire.Doc评估警告文本")
                        
                        # 保存处理后的文档到最终位置
                        doc.save(self.output_path)
                        
                    else:
                        logger.warning("添加水印失败，请检查Spire.Doc是否正确安装")
                except ImportError:
                    logger.warning("未找到watermark模块，无法添加水印")
                except Exception as e:
                    logger.error(f"添加水印时发生错误: {e}")
                finally:
                    # 清理临时文件
                    if temp_output_file and os.path.exists(temp_output_file):
                        try:
                            os.remove(temp_output_file)
                        except:
                            pass
            
            # 更新域
            update_fields_in_word(self.output_path)
            
            logger.info(f"报告生成完成: {self.output_path}")
            return self.output_path
            
        except Exception as e:
            logger.error(f"生成报告失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    def _setup_document_styles(self):
        """设置文档样式"""
        # 设置页面大小为A4
        self.doc = set_page_size_to_a4(self.doc)
        # 设置默认字体为宋体
        self.doc = set_default_font(self.doc)
        # 创建样式
        self.doc = create_styles(self.doc)
        # 设置目录样式
        self.doc = create_toc_styles(self.doc)
        
        
    
    def _create_cover_page(self):
        """创建封面页"""
        logger.info("添加封面...")
        
        # 添加Logo
        logo_path = self.update_data.get('image_resources', {}).get('logo_path')
        if logo_path and os.path.exists(logo_path):
            logger.info(f"找到logo图片文件: {logo_path}")
        else:
            logo_path = DEFAULT_COMPANY_INFO.get('logo_path', '')
            logger.error(f"未找到logo图片文件: {logo_path}，使用默认logo：{logo_path}")
        
        logo_paragraph = self.doc.add_paragraph()
        logo_paragraph.paragraph_format.left_indent = 0
        logo_paragraph.paragraph_format.right_indent = 0
        logo_paragraph.paragraph_format.space_before = 0
        logo_paragraph.paragraph_format.space_after = 0
        logo_paragraph.paragraph_format.line_spacing = 1
        logo_paragraph.paragraph_format.first_line_indent = 0
        run = logo_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(2.5))
        logger.info(f"已添加logo图片: {logo_path}")
        
        # 添加空行
        add_paragraph_with_style(self.doc, "\n\n\n\n", "正文")
        
        # 添加报告标题
        add_mixed_font_paragraph(self.doc, self.report_structure["title"], "文档标题")
        
        # 添加空行
        add_paragraph_with_style(self.doc, "\n\n\n\n\n", "正文")
        
        # 添加公司名称和日期
        add_mixed_font_paragraph(self.doc, self.update_data.get('company_info', {}).get('name', '未找到 公司名称'), "公司名称")
        add_mixed_font_paragraph(self.doc, self.update_data.get('company_info', {}).get('date', '未找到 日期'), "日期")
        
        add_paragraph_with_style(self.doc, "\n", "正文")

        # 添加底部图片
        bottom_img_path = DEFAULT_COMPANY_INFO.get('watermark_path', '')
    
        # 为水印创建无格式段落，但保持居中
        watermark_paragraph = self.doc.add_paragraph()
        watermark_paragraph.paragraph_format.left_indent = 0
        watermark_paragraph.paragraph_format.right_indent = 0
        watermark_paragraph.paragraph_format.space_before = 0
        watermark_paragraph.paragraph_format.space_after = 0
        watermark_paragraph.paragraph_format.line_spacing = 1
        watermark_paragraph.paragraph_format.first_line_indent = 0
        
        if bottom_img_path and os.path.exists(bottom_img_path):
            # 获取页面宽度
            section = self.doc.sections[0]
            page_width = section.page_width - section.left_margin - section.right_margin
            
            # 添加图片并设置宽度为页面宽度
            run = watermark_paragraph.add_run()
            run.add_picture(bottom_img_path, width=Inches(page_width/914400))  # 914400是1英寸的EMU值
            logger.info(f"已添加封面底部图片: {bottom_img_path}")
        else:
            logger.warning(f"未找到封面底部图片文件")
        
        # 添加分页符
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        

    
    def _create_company_profile(self):
        """创建公司简介页"""
        logger.info("添加公司简介...")
        
        # 添加公司简介文本
        if "profile" in self.update_data.get('company_info', '未找到公司简介'):
            paragraphs = self.update_data.get('company_info').get("profile").strip().split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    self.doc.add_paragraph(para_text.strip(), style="正文")
        
        add_paragraph_with_style(self.doc, "\n\n", "正文")

        # 添加联系信息
        address_para = self.doc.add_paragraph(style="地址")
        label_run = address_para.add_run("地 址：")
        label_run.font.name = "宋体"
        content_run = address_para.add_run(self.update_data.get('company_info')["address"])
        content_run.font.name = "宋体"
        
        # 邮箱
        email_para = self.doc.add_paragraph(style="邮件")
        label_run = email_para.add_run("邮 箱：")
        label_run.font.name = "宋体"
        content_run = email_para.add_run(self.update_data.get('company_info')["email"])
        content_run.font.name = "Times New Roman"  # 邮箱使用Times New Roman字体
        
        # 电话
        phone_para = self.doc.add_paragraph(style="电话")
        label_run = phone_para.add_run("电 话：")
        label_run.font.name = "宋体"
        content_run = phone_para.add_run(self.update_data.get('company_info')["phone"])
        content_run.font.name = "Times New Roman"  # 电话号码使用Times New Roman字体
        
        # 确保第一节没有页眉页脚和页码
        if self.doc.sections:
            first_section = self.doc.sections[0]
            
            # 使用XML直接设置页码格式为none
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            sectPr = first_section._sectPr
            pgNumType_elements = sectPr.xpath('./w:pgNumType')
            
            if pgNumType_elements:
                pgNumType = pgNumType_elements[0]
                # 删除fmt属性实现隐藏页码
                if pgNumType.get(qn('w:fmt')):
                    del pgNumType.attrib[qn('w:fmt')]
            else:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            
            # 清空页眉
            header = first_section.header
            for p in header.paragraphs:
                p.clear()
            
            # 清空页脚
            footer = first_section.footer
            for p in footer.paragraphs:
                p.clear()
        
        # 封面和目录之间添加分节符（新的一页）
        self.doc.add_section(WD_SECTION_START.NEW_PAGE)
    
    def _create_toc(self):
        """创建目录"""
        logger.info("添加目录...")
        add_toc(self.doc)

    
    def _create_chapters(self):
        """创建章节内容"""
        logger.info("添加章节内容...")
        
        
        # 遍历章节
        chapters_to_add = []
        for chapter in self.report_structure.get("chapters", []):
            chapters_to_add.append(chapter)
        
        for chapter_idx, chapter in enumerate(chapters_to_add):
            chapter_num = chapter_idx + 1

            # 添加章节标题
            chapter_title = f"{chapter_num}. {chapter['title']}"
            self.doc.add_paragraph(chapter_title, style="一级标题")
            
            # 处理章节直接内容
            if "content" in chapter:
                content = chapter["content"]
                if isinstance(content, list):
                    for para in content:
                        if para.strip():
                            self.doc.add_paragraph(para, style="正文")
                elif isinstance(content, str) and content.strip():
                    self.doc.add_paragraph(content, style="正文")
            
            # 处理小节
            figure_count = 1  # 图片计数器
            table_count = 1   # 表格计数器
            
            for section_idx, section in enumerate(chapter.get('sections', [])):
                # 添加小节标题
                add_paragraph_with_style(self.doc, f"{chapter_num}.{section_idx+1} {section['name']}", "二级标题")
                
                # 处理小节内容 - 可以是单一content或有序items
                if "items" in section and section["items"]:
                    # 初始化子段落索引计数器
                    subsection_idx = 0
                    
                    # 使用有序items列表处理内容
                    for item in section["items"]:
                        item_type = item.get("type", "text")
                        
                        # 处理文本项
                        if item_type == "text" and "content" in item and item["content"]:
                            content = item["content"]
                            if isinstance(content, list):
                                for para in content:
                                    if para.strip():
                                        self.doc.add_paragraph(para, style="正文")
                            elif isinstance(content, str) and content.strip():
                                self.doc.add_paragraph(content, style="正文")
                        
                        # 处理子段落（三级标题及内容）
                        elif item_type == "subsection" and "name" in item:
                            # 更新子段落索引
                            subsection_idx += 1
                            
                            # 添加三级标题
                            add_paragraph_with_style(self.doc, f"{chapter_num}.{section_idx+1}.{subsection_idx} {item['name']}", "三级标题")
                            
                            # 处理子段落内容
                            if "items" in item and item["items"]:
                                for sub_item in item["items"]:
                                    sub_item_type = sub_item.get("type", "text")
                                    
                                    # 处理文本
                                    if sub_item_type == "text" and "content" in sub_item and sub_item["content"]:
                                        add_paragraph_with_style(self.doc, sub_item["content"], "正文")
                                        content = sub_item["content"]
                                        if isinstance(content, list):
                                            for para in content:
                                                if para.strip():
                                                    self.doc.add_paragraph(para, style="正文")
                                        elif isinstance(content, str) and content.strip():
                                            self.doc.add_paragraph(content, style="正文")
                                    
                                    # 处理图片
                                    elif sub_item_type == "image" and "path" in sub_item and "caption" in sub_item:
                                        from docx.shared import Inches, Cm
                                        image_width = 14.64
                                        
                                        # 处理图片路径
                                        image_path = sub_item["path"]
                                        if not os.path.exists(image_path):
                                            logger.error(f"未找到图片：{image_path}")
                                            
                                        # 添加图片
                                        figure_count = add_image_to_doc(
                                            self.doc, 
                                            image_path, 
                                            width=image_width,
                                            caption=sub_item["caption"],
                                            chapter_num=chapter_num,
                                            figure_count=figure_count
                                        )
                                    
                                    # 处理表格
                                    elif sub_item_type == "table" and "name" in sub_item:
                                        title = sub_item.get("name", "")
                                        headers = sub_item.get("headers", None)
                                        data = sub_item.get("data", [])
                                        merge_cells = sub_item.get("merge_cells", None)
                                        column_widths = sub_item.get("column_widths", None)
                                        
                                        try:
                                            table_count = add_data_table(
                                                self.doc,
                                                title,
                                                headers,
                                                data,
                                                chapter_num,
                                                None,  # 不传递table_count，启用自动编号
                                                merge_cells=merge_cells,
                                                column_widths=column_widths
                                            )
                                        except Exception as e:
                                            logger.error(f"添加表格 '{title}' 时出现问题: {str(e)}")
                                    elif "content" in item and item["content"]:
                                        # 处理子段落的简单内容
                                        if isinstance(item["content"], list):
                                            for paragraph_text in item["content"]:
                                                if paragraph_text.strip():
                                                    add_paragraph_with_style(self.doc, paragraph_text, "正文")
                                        else:
                                            if item["content"].strip():
                                                add_paragraph_with_style(self.doc, item["content"], "正文")
                                
                        # 处理图片项
                        elif item_type == "image" and "path" in item and "caption" in item:
                            # 使用固定宽度值
                            image_width = 14.64  # 约14.64厘米宽
                            
                            # 处理图片路径
                            image_path = item["path"]
                            
                            # 特殊处理第2章的卫星图和航点规划图
                            if chapter_num == 2 and "现场图片" in section.get("name", ""):
                                caption_text = item.get("caption", "").lower()
                                if "卫星图" in caption_text and not image_path and "satellite_img" in self.update_data.get("company_info"):
                                    image_path = self.update_data.get("company_info").get("satellite_img")
                                    logger.info(f"添加satellite_img: {image_path}")
                                elif "航点规划图" in caption_text and not image_path and "wayline_img" in self.update_data.get("company_info"):
                                    image_path = self.update_data.get("company_info").get("wayline_img")
                                    logger.info(f"添加wayline_img: {image_path}")
                            
                            
                            # 添加图片并更新计数
                            figure_count = add_image_to_doc(
                                self.doc, 
                                image_path, 
                                width=image_width,
                                caption=item["caption"],
                                chapter_num=chapter_num,
                                figure_count=figure_count
                            )
                        
                        # 处理表格项
                        elif item_type == "table" and "name" in item:
                            # 提取数据和头部信息
                            title = item.get("name", "")
                            headers = item.get("headers", None)
                            data = item.get("data", [])
                            
                            # 提取合并单元格和列宽信息
                            merge_cells = item.get("merge_cells", None)
                            column_widths = item.get("column_widths", None)
                            
                            # 添加表格
                            try:
                                table_count = add_data_table(
                                    self.doc,
                                    title,
                                    headers,
                                    data,
                                    chapter_num,
                                    None,  # 不传递table_count，启用自动编号
                                    merge_cells=merge_cells,
                                    column_widths=column_widths
                                )
                                logger.info(f"成功添加表格: {title}")
                            except Exception as e:
                                logger.error(f"添加表格 '{title}' 时出现问题: {str(e)}")
                                continue
                
            
            # 添加章节结束的分页符，最后一章不添加
            if chapter_idx < len(chapters_to_add) - 1:
                self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        return self.doc