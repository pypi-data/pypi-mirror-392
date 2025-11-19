"""
段落样式模块 - 提供段落样式相关的功能
"""
import logging
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from ..config.fonts import HEI_FONT_NAMES, SONG_FONT_NAMES, TIMES_NEW_ROMAN
from ..utils.text import is_numeric, is_chinese
from ..utils.font import apply_font

# 获取模块日志记录器
logger = logging.getLogger(__name__)



def add_paragraph_with_style(doc: Document, text: str, style_name: str, alignment=None) -> None:
    """添加带样式的段落
    
    Args:
        doc: 文档对象
        text: 段落文本
        style_name: 样式名称
        alignment: 对齐方式
        
    Returns:
        添加的段落对象
    """
    paragraph = doc.add_paragraph(style=style_name)
    
    if alignment is not None:
        paragraph.alignment = alignment
    
    if text:
        paragraph.add_run(text)
    
    return paragraph

def add_date_with_bold_numbers(doc, date_text, style_name):
    """添加日期文本，数字使用黑体加粗显示
    
    Args:
        doc: 文档对象
        date_text: 日期文本
        style_name: 样式名称
        
    Returns:
        返回添加的段落对象
    """
    paragraph = doc.add_paragraph()
    paragraph.style = style_name
    
    # 获取字体大小
    style = doc.styles[style_name]
    font_size = style.font.size
    
    # 逐字符处理
    for char in date_text:
        run = paragraph.add_run(char)
        # 设置字体大小
        run.font.size = font_size
        
        if is_numeric(char):
            # 数字使用黑体，并加粗
            run.font.name = HEI_FONT_NAMES[0]
            run.font.bold = True
            
            # 设置东亚字体为黑体
            rpr = run._element.get_or_add_rPr()
            rpr.rFonts.set(qn('w:eastAsia'), HEI_FONT_NAMES[0])
        else:
            # 非数字使用宋体，不加粗
            run.font.name = SONG_FONT_NAMES[0]
            run.font.bold = False
            
            # 设置东亚字体为宋体
            rpr = run._element.get_or_add_rPr()
            rpr.rFonts.set(qn('w:eastAsia'), SONG_FONT_NAMES[0])
    
    return paragraph

def add_mixed_font_paragraph(doc, text, style='Normal', alignment=None):
    """添加包含中英文混合的段落，自动使用不同字体
    
    Args:
        doc: 文档对象
        text: 段落文本
        style: 段落样式名称
        alignment: 对齐方式
        
    Returns:
        添加的段落对象
    """
    paragraph = doc.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    
    current_text = ""
    current_is_chinese = None
    
    def add_current_text():
        nonlocal current_text, current_is_chinese
        if current_text:
            run = paragraph.add_run(current_text)
            
            if current_is_chinese:
                apply_font(run, SONG_FONT_NAMES[0])
            else:
                apply_font(run, TIMES_NEW_ROMAN[0])
        
        current_text = ""
        current_is_chinese = None
    
    for char in text:
        char_is_chinese = is_chinese(char)
        
        if current_is_chinese is None:
            current_is_chinese = char_is_chinese
            current_text += char
        elif current_is_chinese == char_is_chinese:
            current_text += char
        else:
            add_current_text()
            current_is_chinese = char_is_chinese
            current_text = char
    
    add_current_text()
    return paragraph


# 导出所有符号以保持向后兼容性
__all__ = [
    'add_paragraph_with_style',
    'add_date_with_bold_numbers',
    'add_mixed_font_paragraph'
]