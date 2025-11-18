# %% Required modules
import re
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import numpy as np


# %% parallelCoordinatesPlot function
def parallelCoordinatesPlot(
    filePath,
    numAssignments,
    startColIndex,
    normalize=True,
    replaceMissingWithZero=True,
    ydIndex=None,
):
    """This functions creates a parallel coordinates plot of student scores to
    visualize the point at which students start struggling.

    Parameters
    ----------
    filePath : Str
        Full path to the location of the Canvas gradebook file.
    numAssignments : Int
        The number of assignments for which you want to create box plots.
    startColIndex : Int
        The column index where the assignments start. This assumes that all
        of the assignment columns are adjacent.
    normalize : Bool
        Whether to convert the scores to a 0-1 scale so that assignments
        of differnt point values can be compared. The default is True.
    replaceMissingWithZero : Bool
        Whether to replace missing assignment scores with 0. The default is True.
    ydIndex : Int
        The column index for the yellowdig assignment, if there is one.

    Returns
    -------
    An interactive parallel coordinates plot of student scores.
    """
    df = pd.read_csv(filePath)

    # Extract only student name and homework columns
    colNums = [0]
    if ydIndex is not None:
        colNums.append(ydIndex)
    colNums.extend(list(range(startColIndex, startColIndex + numAssignments)))

    # Create new column names for homework assignments
    colNames = ["Student"]
    hwCols = ["h" + str(i) for i in range(1, numAssignments + 1)]
    if ydIndex is not None:
        hwCols.insert(0, "yd")
    colNames.extend(hwCols)
    df = df.iloc[:, colNums]
    df.columns = colNames

    # Calculate divisors, including a divisor for the total score
    divisors = df.iloc[
        0, 1:
    ]  # Get the first row, which contains the value of the assignments
    divisors = divisors.tolist()
    divisors.append(sum(divisors))

    # Create a total column that sums the scores for each student
    df["total"] = df.iloc[:, 1:].apply("sum", axis=1)

    # Remove the first row, which contains the assignment values
    df = df.iloc[1:, :]

    # Convert hw columns to numeric just in case they are not
    df[hwCols] = df[hwCols].apply(pd.to_numeric)

    # Loop through each row to calculate the percentage for each assignment
    if normalize:
        for c in range(1, df.shape[1]):
            df.iloc[:, c] = df.iloc[:, c] / divisors[c - 1]
    if replaceMissingWithZero:
        df = df.astype(float).fillna(0)  # Fill in missing values with 0

    # Remove test student
    df = df.query('Student.str.contains("Test") == False')

    fig = px.parallel_coordinates(
        df, color="total", title="Parallel Coordinates Plot of Scores"
    )
    fig.show()


# %% tidyUpCols Function
def tidyUpCols(myList, keepNums=False):
    """
    Parameters
    ----------
    myList : List
        List of column names, like df.columns
    keepNums : Bool
        Whether we want to keep numbers in column names. The default is False.

    Returns
    -------
    ml2 : List
        List of tidy column names.
    """
    ml2 = []
    for i in range(len(myList)):
        if myList[i] != None:
            ti = (
                myList[i]
                .strip()
                .lower()
                .replace(".", "")
                .replace("/", "_")
                .replace(" ", "_")
                .replace("$", "")
            )
        else:
            continue  # Goes to the next iteration in the for loop

        if "--" in ti:
            ti2 = ti.split("--")
            [ml2.append(x) for x in ti2]
        elif keepNums == True:
            ti = re.sub("[^a-zA-Z_0-9]", "", ti)
            ml2.append(ti)
        else:
            ti = re.sub("[^a-zA-Z_]", "", ti)
            ml2.append(ti)
    return ml2


# %% Abbreviate column names function
def abbreviate_col_names(col_name):
    # If Yellowdig is in the name, then return YD
    if col_name.find("Yellowdig") >= 0:
        return "YD"
    elif col_name == "Student":
        return "Student"
    else:
        # Remove the stuff in parentheses at the end
        tname = col_name.split(" (")[0]
        # Create a list of whole words or digits
        parts = re.findall(r"[A-Za-z]+|\d+", tname)
        # Create an empty string
        abbreviation = ""
        for part in parts:
            if part.isdigit():
                # If it's a digit, then add it
                abbreviation += part
            elif part[0].isupper():
                # If the first letter is uppercase, then add it
                abbreviation += part[0]
        return abbreviation


# %% assignmentPlots function
def assignmentPlots(
    filePath,
    numAssignments,
    startColIndex,
    normalize=True,
    replaceMissingWithZero=True,
    ydIndex=None,
):
    """
    This function uses an export from canvas to create boxplots and stripplots
    of assignments.

    Parameters
    ----------
    filePath : Str
        Full path to the location of the Canvas gradebook file.
    numAssignments : Int
        The number of assignments for which you want to create box plots.
    startColIndex : Int
        The column index where the assignments start. This assumes that all
        of the assignment columns are adjacent.
    normalize : Bool
        Whether to convert the scores to a 0-1 scale so that assignments
        of differnt point values can be compared. The default is True.
    replaceMissingWithZero : Bool
        Whether to replace missing assignment scores with 0. The default is True.
    ydIndex : Int
        The column index for the yellowdig assignment, if there is one.

    Returns
    -------
    A box and whisker plot on the left and a stripplot on the right.

    """
    df = pd.read_csv(filePath)

    # Extract only student name and homework columns
    colNums = [0]
    if ydIndex is not None:
        colNums.append(ydIndex)
    colNums.extend(list(range(startColIndex, startColIndex + numAssignments)))

    # Create new column names for homework assignments
    colNames = ["Student"]
    hwCols = ["h" + str(i) for i in range(1, numAssignments + 1)]
    if ydIndex is not None:
        hwCols.insert(0, "yd")
    colNames.extend(hwCols)
    df = df.iloc[:, colNums]
    df.columns = colNames

    # Find the row that contains Points Possible
    points_row = df.iloc[:, 0].str.contains("Points")
    points_row = points_row[points_row == True]
    points_row = points_row.index[0]

    # Calculate divisors, including a divisor for the total score
    divisors = df.iloc[
        points_row, 1:
    ]  # Get the first row, which contains the value of the assignments
    divisors = pd.to_numeric(divisors.tolist())
    divisors = np.append(divisors, np.sum(divisors))

    # Remove the rows up to the points_row
    df = df.iloc[points_row + 1 :, :]

    # Convert all but first column to numeric values
    df.iloc[:, 1:] = df.iloc[:, 1:].apply(pd.to_numeric, errors="coerce")

    # Create a total column that sums the scores for each student
    df["total"] = df.iloc[:, 1:].apply("sum", axis=1)

    # Loop through each row to calculate the percentage for each assignment
    if normalize:
        for c in range(1, df.shape[1]):
            df.iloc[:, c] = df.iloc[:, c] / divisors[c - 1]
    if replaceMissingWithZero:
        df.iloc[:, 1:] = (
            df.iloc[:, 1:].astype(float).fillna(0)
        )  # Fill in missing values with 0

    # Remove test student
    df = df.query('Student.str.contains("Test") == False')

    # Convert from wide to long
    hw = df.melt(value_vars=hwCols, var_name="Homework", value_name="Score")

    # Create canvas with two columns
    fig, axs = plt.subplots(figsize=(15, 5), ncols=2)
    fig.tight_layout(pad=2)

    # Boxplot on left side
    sns.boxplot(
        data=hw,
        x="Homework",
        y="Score",
        hue="Homework",
        ax=axs[0],
        showmeans=True,
        meanprops={"markeredgecolor": "black"},
    )
    axs[0].set_title("Boxplot of Homework Scores")

    # Stripplot on right side
    sns.stripplot(data=hw, x="Homework", y="Score", hue="Homework", ax=axs[1])
    axs[1].set_title("Stripplot of Homework Scores")


# Test out the function
# assignmentPlots('/Users/rnguymon/Downloads/htdy2.csv', 5, 9)


def assignmentPlots2(
    filePath, cols_to_exclude=[], normalize=True, replaceMissingWithZero=True
):
    """
    This function uses an export from canvas to create boxplots and stripplots
    of assignments.

    Parameters
    ----------
    filePath : Str
        Full path to the location of the Canvas gradebook file.
    cols_to_exclude: List
        List of abbreviated column names that you don't want to include. You may have to run the function once to see what the column names are.
    normalize : Bool
        Whether to convert the scores to a 0-1 scale so that assignments
        of differnt point values can be compared. The default is True.
    replaceMissingWithZero : Bool
        Whether to replace missing assignment scores with 0. The default is True.

    Returns
    -------
    A box and whisker plot on the left and a stripplot on the right.

    """
    df = pd.read_csv(filePath)

    # Find the row that contains Points Possible
    points_row = df.iloc[:, 0].str.contains("Points")
    points_row = points_row[points_row == True]
    points_row = points_row.index[0]
    points_row

    # Find the student column
    cols_to_keep = [c for c in df.columns if c.find("Student") >= 0]

    # Find the columns that contains the points possible values
    points_cols = pd.to_numeric(df.iloc[points_row, :], errors="coerce")
    points_cols = points_cols.index[points_cols.notna()]
    cols_to_keep.extend(points_cols)

    # Keep only the student and points columns
    df = df.loc[:, cols_to_keep]

    # Abbreviate column names
    df.columns = [abbreviate_col_names(c) for c in df.columns]

    # Get numbers to use as the divisors
    divisors = df.iloc[
        points_row, 1:
    ]  # Get the first row, which contains the value of the assignments
    divisors = pd.to_numeric(divisors)

    # Remove the rows up to the points_row
    df = df.iloc[points_row + 1 :, :]

    # Remove test student
    df = df.query('Student.str.contains("Test") == False')

    # Set student name as index
    df = df.set_index("Student")

    # Keep only columns that have submissions
    pct_complete = df.count(axis=0) / df.shape[0]
    asgn_with_submissions = pct_complete[pct_complete > 0].index
    df = df.loc[:, asgn_with_submissions]

    # Remove columns that the user doesn't want
    df = df.drop(columns=cols_to_exclude)

    # Put Yellowdig as the first column
    col_names = list(df.columns)
    for i, c in enumerate(col_names):
        if c == "YD":
            my_col = col_names.pop(i)
    col_names.insert(0, my_col)
    df = df.loc[:, col_names]

    # Keep only the matching divisors
    divisors = divisors[df.columns]

    # Convert all columns to numeric values
    df = df.apply(pd.to_numeric, errors="coerce")

    # Create a total column that sums the scores for each student
    df["Total"] = df.apply("sum", axis=1)
    divisors["Total"] = sum(divisors)

    # Loop through each row to calculate the percentage for each assignment
    if normalize:
        for c in range(df.shape[1]):
            df.iloc[:, c] = df.iloc[:, c] / divisors.iloc[c]
    if replaceMissingWithZero:
        df = df.astype(float).fillna(0)  # Fill in missing values with 0

    # Convert from wide to long
    hw = df.melt(var_name="Homework", value_name="Score")

    # Create canvas with two columns
    fig, axs = plt.subplots(figsize=(15, 5), ncols=2, sharey=True)

    # Boxplot on left side
    sns.boxplot(
        data=hw,
        x="Homework",
        y="Score",
        hue="Homework",
        ax=axs[0],
        showmeans=True,
        meanprops={"markeredgecolor": "black"},
    )
    axs[0].set_title("Boxplot of Homework Scores")

    # Stripplot on right side
    sns.stripplot(data=hw, x="Homework", y="Score", hue="Homework", ax=axs[1])
    axs[1].set_title("Stripplot of Homework Scores")
    fig.tight_layout(pad=2)


# %% relocate function
def relocate(df, old_index, new_index):
    """
    This function relocates one column of a dataframe based on index number.

    Parameters
    ----------
    df : Pandas dataframe
        This is the dataframe object that has a column you would like to
        relocate.
    old_index : INT
        This is th eindex number of the column that you want to relocate.
    new_index : INT
        This is the destination index number of the relocated column.

    Returns
    -------
    df : Pandas dataframe
        The same dataframe with the relocated column.

    """
    # Convert column names to a list, col_names
    col_names = df.columns.tolist()
    # Remove the column and insert it into a new location
    col_names.insert(new_index, col_names.pop(old_index))
    # Slice the dataframe using the col_names list
    df = df.loc[:, col_names]
    # Return the dataframe
    return df


# %% relocate by name function
def relocate_by_name(df, col_name, new_index):
    """
    This function relocates one column of a dataframe based on the new column's name and the new index number.

    Parameters
    ----------
    df : Pandas dataframe
        This is the dataframe object that has a column you would like to
        relocate.
    col_name : STR
        This is the name of the column that you want to relocate.
    new_index : INT
        This is the destination index number of the relocated column.

    Returns
    -------
    df : Pandas dataframe
        The same dataframe with the relocated column.

    """
    # Convert column names to a list, col_names
    col_names = list(df.columns)

    # Loop through the column names and remove the specified column
    for i, c in enumerate(col_names):
        if c == col_name:
            my_col = col_names.pop(i)

    # Insert the column into the new location
    col_names.insert(new_index, my_col)

    # Slice the dataframe using the col_names list
    df = df.loc[:, col_names]

    # Return the dataframe
    return df


# %% SQlite ERD Function


def sqlite_erd(db, render=True, include_actions=False, return_markdown=False):
    """
    Render a Mermaid ER diagram for a SQLite database (inline in Jupyter).

    - Opens the SQLite DB (or uses an existing sqlite3.Connection)
    - Introspects tables, PKs, and FKs (incl. composites) via PRAGMA
    - Emits Mermaid `erDiagram` with one-to-many relationships

    Args:
        db (str | sqlite3.Connection): Path to a SQLite file or an existing connection.
        render (bool): If True, display inline and return None (prevents double-render).
        include_actions (bool): If True, append ON DELETE/UPDATE to relationship labels.
        return_markdown (bool): When render=False, return IPython Markdown instead of raw text.

    Returns:
        None | str | IPython.display.Markdown:
            - None if render=True (diagram already displayed once).
            - If render=False:
                * Mermaid code block string (default), or
                * Markdown object when return_markdown=True.
    """
    # Auto-import everything needed
    import sqlite3
    from contextlib import contextmanager
    from IPython.display import Markdown, display

    @contextmanager
    def _connect(maybe_conn):
        if isinstance(maybe_conn, sqlite3.Connection):
            yield maybe_conn, False
        else:
            conn = sqlite3.connect(str(maybe_conn))
            try:
                yield conn, True
            finally:
                conn.close()

    def _get_tables(conn):
        rows = conn.execute(
            "SELECT name FROM sqlite_master "
            "WHERE type='table' AND name NOT LIKE 'sqlite_%' "
            "ORDER BY name"
        ).fetchall()
        return [r[0] for r in rows]

    def _get_columns(conn, table):
        # PRAGMA table_info: cid, name, type, notnull, dflt_value, pk (0 or ordinal >0)
        cols = conn.execute(f"PRAGMA table_info('{table}')").fetchall()
        return [
            {"name": name, "type": (col_type or "TEXT").strip(), "pk_ord": int(pk_ord)}
            for _, name, col_type, _, _, pk_ord in cols
        ]

    def _get_foreign_keys(conn, table):
        # PRAGMA foreign_key_list: id, seq, table(parent), from, to, on_update, on_delete, match
        rows = conn.execute(f"PRAGMA foreign_key_list('{table}')").fetchall()
        groups = {}
        for fk_id, seq, parent, from_col, to_col, on_update, on_delete, match in rows:
            g = groups.setdefault(
                fk_id,
                {
                    "parent": parent,
                    "pairs": [],
                    "on_update": on_update,
                    "on_delete": on_delete,
                    "match": match,
                },
            )
            g["pairs"].append((from_col, to_col))
        for g in groups.values():
            g["pairs"].sort(key=lambda p: (p[0] or "", p[1] or ""))
        return list(groups.values())

    def _format_table_block(name, columns):
        lines = [f"    {name} {{"]  # Mermaid table header
        for c in columns:
            tag = " PK" if c["pk_ord"] else ""
            lines.append(f"        {c['type']} {c['name']}{tag}")
        lines.append("    }")
        return "\n".join(lines)

    def _format_relationship(parent, child, pairs, on_update=None, on_delete=None):
        # Parent ||--o{ Child : "Child.child_cols → Parent.parent_cols [actions]"
        child_cols = ", ".join([p[0] for p in pairs])
        parent_cols = ", ".join([p[1] if p[1] else "(PK)" for p in pairs])
        label = f"{child}.{child_cols} → {parent}.{parent_cols}"
        if include_actions:
            actions = []
            if on_delete and on_delete.upper() != "NO ACTION":
                actions.append(f"ON DELETE {on_delete}")
            if on_update and on_update.upper() != "NO ACTION":
                actions.append(f"ON UPDATE {on_update}")
            if actions:
                label += " [" + ", ".join(actions) + "]"
        return f'    {parent} ||--o{{ {child} : "{label}"'

    def _build_mermaid(conn):
        tables = _get_tables(conn)

        table_blocks = []
        for t in tables:
            cols = _get_columns(conn, t)
            table_blocks.append(_format_table_block(t, cols))

        rels, seen = [], set()
        for child in tables:
            for fk in _get_foreign_keys(conn, child):
                parent = fk["parent"]
                if parent not in tables:
                    continue
                key = (parent, child, tuple(fk["pairs"]))
                if key in seen:
                    continue
                seen.add(key)
                rels.append(
                    _format_relationship(
                        parent, child, fk["pairs"], fk["on_update"], fk["on_delete"]
                    )
                )

        parts = ["```mermaid", "erDiagram"]
        parts.extend(table_blocks)
        if rels:
            parts.append("")
            parts.extend(rels)
        parts.append("```")
        return "\n".join(parts)

    with _connect(db) as (conn, _created):
        conn.execute("PRAGMA foreign_keys=ON;")
        mermaid_code_block = _build_mermaid(conn)

    if render:
        display(Markdown(mermaid_code_block))  # display once
        return None  # <-- prevent Jupyter from auto-rendering a return value
    else:
        return Markdown(mermaid_code_block) if return_markdown else mermaid_code_block


# %% Function to return code cells as formatted text in a Word document
import nbformat
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from pygments import lex
from pygments.lexers import PythonLexer
from pygments.token import Token


TOKEN_COLORS = {
    # keywords: for, if, return, import, etc.
    Token.Keyword: RGBColor(0x00, 0x80, 0x00),  # green-ish
    # builtins: len, range, print, dict, list, etc.
    Token.Name.Builtin: RGBColor(0x00, 0x80, 0x80),  # teal
    # functions & classes
    Token.Name.Function: RGBColor(0x00, 0x00, 0xFF),  # blue
    Token.Name.Class: RGBColor(0x00, 0x00, 0xFF),  # blue
    # attributes/method-like names (we’ll also use this color for df.plot)
    Token.Name.Attribute: RGBColor(0x00, 0x00, 0xFF),  # blue-ish
    # strings
    Token.String: RGBColor(0xBA, 0x21, 0x21),  # dark red
    # comments
    Token.Comment: RGBColor(0x40, 0x80, 0x80),  # gray-green
    # numbers
    Token.Number: RGBColor(0x00, 0x80, 0x80),  # teal
    # mathematical operators in purple
    Token.Operator: RGBColor(0x80, 0x00, 0x80),
    Token.Operator.Word: RGBColor(0x80, 0x00, 0x80),
    # punctuation (commas, parens, etc.) stay black by default
    Token.Punctuation: RGBColor(0x00, 0x00, 0x00),
}

DEFAULT_TEXT_COLOR = RGBColor(0x00, 0x00, 0x00)  # black


def get_color_for_token(ttype):
    """
    Walk up the token hierarchy until we find a matching color.
    Fallback = black (DEFAULT_TEXT_COLOR).
    """
    current = ttype
    while current is not Token and current not in TOKEN_COLORS:
        current = current.parent
    return TOKEN_COLORS.get(current, DEFAULT_TEXT_COLOR)


def add_gray_background(cell, fill_hex="F7F7F7"):
    """
    Shade a table cell with a light gray background.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(
        r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(
            nsdecls("w"), fill_hex
        )
    )
    tc_pr.append(shading)


def process_token_into_runs(
    para, ttype, value, font_name, font_size_pt, prev_ttype, prev_value
):
    """
    Turn a single Pygments token into one or more Word runs.
    - Keeps syntax colors based on token type.
    - Forces '.' tokens to black.
    - If a Token.Name comes right after '.', color it like an attribute/method.
    """
    # Handle multi-line tokens by splitting on newline
    lines = value.split("\n")

    for line_index, line in enumerate(lines):
        if not line:
            # Still need to preserve line breaks
            if line_index < len(lines) - 1:
                para.add_run().add_break()
            continue

        run = para.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)

        # Decide on color
        if value == "." and ttype in (Token.Operator, Token.Punctuation):
            # Literal '.' → always black
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif (
            ttype is Token.Name
            and prev_ttype in (Token.Operator, Token.Punctuation)
            and prev_value == "."
        ):
            # Name immediately following '.' → treat like attribute/method
            attr_color = TOKEN_COLORS.get(Token.Name.Attribute, DEFAULT_TEXT_COLOR)
            run.font.color.rgb = attr_color

        else:
            # Normal syntax coloring
            color = get_color_for_token(ttype)
            run.font.color.rgb = color

        # Add line break if not the last fragment
        if line_index < len(lines) - 1:
            para.add_run().add_break()


def notebook_code_to_word(ipynb_path, docx_path, font_name="Consolas", font_size_pt=10):
    """
    Convert code cells from a Jupyter notebook into a Word document with:
      - JupyterLab Light-ish syntax colors
      - gray background for each code cell
      - '.' explicitly black
      - names immediately after '.' (df.plot) styled as methods.
    """

    # --- Color config (approx JupyterLab Light) -----------------------------

    nb = nbformat.read(ipynb_path, as_version=4)
    doc = Document()

    for cell in nb.cells:
        if cell.cell_type != "code":
            continue

        # Create a 1x1 table as a "code block" with gray background
        table = doc.add_table(rows=1, cols=1)
        table.allow_autofit = True
        code_cell = table.rows[0].cells[0]

        # Shade the background
        add_gray_background(code_cell, fill_hex="F7F7F7")

        # Use the first paragraph in the cell for our code
        para = code_cell.paragraphs[0]

        prev_ttype = None
        prev_value = None

        # Lex the code with Pygments
        for ttype, value in lex(cell.source, PythonLexer()):
            if not value:
                continue

            process_token_into_runs(
                para, ttype, value, font_name, font_size_pt, prev_ttype, prev_value
            )

            prev_ttype, prev_value = ttype, value

        # Blank line after each code-cell table
        doc.add_paragraph()

    doc.save(docx_path)
    print(f"Saved syntax-colored code cells to: {docx_path}")
