"""
Structured DOCX Parser for Doctra

This module provides comprehensive DOCX parsing capabilities for extracting
text, tables, images, and structured content from Microsoft Word documents.
"""

from __future__ import annotations
import json
import os
import re
import sys
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from contextlib import ExitStack
from PIL import Image
from tqdm import tqdm

# Optional imports for Excel functionality
try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.utils.dataframe import dataframe_to_rows
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.worksheet.hyperlink import Hyperlink
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.ns import qn
except ImportError:
    print("Warning: python-docx not installed. DOCX parsing will not be available.")
    print("Install with: pip install python-docx")
    Document = None
    DocumentType = None
    Table = None
    Paragraph = None
    CT_Tbl = None
    CT_P = None
    qn = None

# Removed IMAGE_SUBDIRS import - using simple images folder structure
from doctra.exporters.markdown_writer import write_markdown
from doctra.exporters.html_writer import write_html
from doctra.exporters.markdown_table import render_markdown_table
from doctra.exporters.image_saver import save_box_image
from doctra.engines.vlm.outlines_types import TabularArtifact
from doctra.engines.vlm.service import VLMStructuredExtractor


class StructuredDOCXParser:
    """
    Comprehensive DOCX parser for extracting all types of content.
    
    Processes DOCX documents to extract text, tables, images, and figures.
    Supports structured data extraction and optional VLM processing for
    enhanced content analysis.

    :param vlm: VLM engine instance (VLMStructuredExtractor). If None, VLM processing is disabled.
    :param extract_images: Whether to extract embedded images (default: True)
    :param preserve_formatting: Whether to preserve text formatting in output (default: True)
    :param table_detection: Whether to detect and extract tables (default: True)
    """

    def __init__(
        self,
        *,
        vlm: Optional[VLMStructuredExtractor] = None,
        extract_images: bool = True,
        preserve_formatting: bool = True,
        table_detection: bool = True,
        export_excel: bool = True,
    ):
        """
        Initialize the StructuredDOCXParser with processing configuration.

        :param vlm: VLM engine instance (VLMStructuredExtractor). If None, VLM processing is disabled.
        :param extract_images: Whether to extract embedded images (default: True)
        :param preserve_formatting: Whether to preserve text formatting in output (default: True)
        :param table_detection: Whether to detect and extract tables (default: True)
        :param export_excel: Whether to export tables to Excel file (default: True)
        """
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        self.extract_images = extract_images
        self.preserve_formatting = preserve_formatting
        self.table_detection = table_detection
        self.export_excel = export_excel
        
        # Initialize VLM engine - use provided instance or None
        if vlm is None:
            self.vlm = None
        elif isinstance(vlm, VLMStructuredExtractor):
            self.vlm = vlm
        else:
            raise TypeError(
                f"vlm must be an instance of VLMStructuredExtractor or None, "
                f"got {type(vlm).__name__}"
            )

    def parse(self, docx_path: str) -> None:
        """
        Parse a DOCX document and extract all content.

        :param docx_path: Path to the DOCX file to parse
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        docx_path = Path(docx_path)
        output_dir = Path(f"outputs/{docx_path.stem}")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        print(f"📄 Processing DOCX: {docx_path.name}")
        
        try:
            doc = Document(docx_path)
            
            document_data = self._extract_document_structure(doc)
            
            images_data = []
            if self.extract_images:
                images_data = self._extract_images(doc, output_dir)
            
            tables_data = [elem for elem in document_data['elements'] if elem['type'] == 'table']
            
            if self.vlm is not None and images_data:
                total_steps = len(images_data)
            else:
                total_steps = 1
            
            progress_bar = tqdm(total=total_steps, desc="Processing DOCX", unit="image")
            
            vlm_extracted_data = []
            if self.vlm is not None and images_data:
                vlm_extracted_data = self._process_vlm_data(images_data, output_dir, progress_bar)
            else:
                progress_bar.update(1)
            
            progress_bar.close()
            
            self._generate_markdown_output(document_data, images_data, output_dir, vlm_extracted_data)
            self._generate_html_output(document_data, images_data, output_dir, vlm_extracted_data)
            
            if self.export_excel:
                if vlm_extracted_data:
                    self._generate_excel_output_with_vlm(tables_data, vlm_extracted_data, output_dir)
                else:
                    self._generate_excel_output(tables_data, output_dir)
            
            print(f"✅ DOCX parsing completed successfully!")
            print(f"📊 Extracted: {len(document_data.get('paragraphs', []))} paragraphs, "
                  f"{len(tables_data)} tables, {len(images_data)} images")
            
        except Exception as e:
            print(f"❌ Error parsing DOCX: {e}")
            raise

    def _extract_document_structure(self, doc: DocumentType) -> Dict[str, Any]:
        """Extract the overall document structure."""
        document_data = {
            'elements': [],  # Mixed list of paragraphs, tables, and other elements
            'paragraphs': [],
            'headings': [],
            'lists': [],
            'metadata': {}
        }
        
        document_data['metadata'] = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
            'subject': doc.core_properties.subject or '',
            'created': str(doc.core_properties.created) if doc.core_properties.created else '',
            'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
        }
        
        self._extract_document_elements_in_order(doc, document_data)
        
        return document_data

    def _extract_document_elements_in_order(self, doc: DocumentType, document_data: Dict):
        """Extract document elements (paragraphs and tables) in their original order."""
        elements = []
        paragraph_index = 0
        table_index = 0
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element and para.text.strip():
                        para_data = {
                            'type': 'paragraph',
                            'index': paragraph_index,
                            'text': para.text.strip(),
                            'style': para.style.name if para.style else 'Normal',
                            'is_heading': para.style.name.startswith('Heading') if para.style else False,
                            'level': self._get_heading_level(para.style.name) if para.style else 0,
                            'formatting': self._extract_formatting(para) if self.preserve_formatting else {}
                        }
                        
                        elements.append(para_data)
                        document_data['paragraphs'].append(para_data)
                        
                        # Categorize headings
                        if para_data['is_heading']:
                            document_data['headings'].append(para_data)
                        
                        paragraph_index += 1
                        break
            
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_data = {
                            'type': 'table',
                            'index': table_index,
                            'rows': len(table.rows),
                            'cols': len(table.columns),
                            'data': [],
                            'markdown': ''
                        }
                        
                        for row_idx, row in enumerate(table.rows):
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                row_data.append(cell_text)
                            table_data['data'].append(row_data)
                        
                        if table_data['data']:
                            headers = table_data['data'][0] if table_data['data'] else []
                            rows = table_data['data'][1:] if len(table_data['data']) > 1 else []
                            table_data['markdown'] = render_markdown_table(headers, rows)
                        
                        elements.append(table_data)
                        table_index += 1
                        break
        
        document_data['elements'] = elements

    def _extract_tables(self, doc: DocumentType, output_dir: Path) -> List[Dict[str, Any]]:
        """Extract all tables from the document."""
        tables_data = []
        
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': [],
                'markdown': ''
            }
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data['data'].append(row_data)
            
            if table_data['data']:
                headers = table_data['data'][0] if table_data['data'] else []
                rows = table_data['data'][1:] if len(table_data['data']) > 1 else []
                table_data['markdown'] = render_markdown_table(headers, rows)
                print(f"📊 Table {table_idx + 1}: {len(table_data['data'])} rows, {len(table_data['data'][0]) if table_data['data'] else 0} columns")
            
            tables_data.append(table_data)
        
        return tables_data

    def _extract_images(self, doc: DocumentType, output_dir: Path) -> List[Dict[str, Any]]:
        """Extract embedded images from the document."""
        images_data = []
        images_dir = output_dir / "images"
        images_dir.mkdir(exist_ok=True)
        
        try:
            for rel in doc.part.rels.values():
                if hasattr(rel, 'target_ref'):
                    content_type = getattr(rel, 'target_content_type', 'unknown')
                    is_image = False
                    if "image" in rel.target_ref or "media" in rel.target_ref:
                        is_image = True
                    elif content_type and "image/" in content_type:
                        is_image = True
                    elif rel.target_ref.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')):
                        is_image = True
                    
                    if is_image:
                        try:
                            image_blob = rel.target_part.blob
                            if image_blob:
                                original_filename = rel.target_ref
                                clean_filename = Path(original_filename).name
                                
                                image_data = {
                                    'filename': clean_filename,
                                    'original_path': original_filename,
                                    'type': clean_filename.split('.')[-1].lower(),
                                    'path': str(images_dir / clean_filename)
                                }
                                
                                target_path = Path(image_data['path'])
                                target_path.parent.mkdir(parents=True, exist_ok=True)
                                
                                with open(target_path, 'wb') as f:
                                    f.write(image_blob)
                                
                                images_data.append(image_data)
                        except Exception as img_error:
                            pass  # Silently skip problematic images
            
        except Exception as e:
            pass  # Silently skip if relationships can't be accessed
        
        return images_data

    def _process_vlm_data(self, images_data: List, output_dir: Path, progress_bar=None) -> List[Dict]:
        """Process images with VLM to extract structured data."""
        vlm_extracted_data = []
        if images_data:
            for i, img_data in enumerate(images_data):
                try:
                    if progress_bar:
                        progress_bar.set_description(f"Processing image {i+1}/{len(images_data)}: {img_data['filename']}")
                    
                    result = self.vlm.extract_table_or_chart(img_data['path'])
                    
                    if hasattr(result, 'title') and hasattr(result, 'description'):
                        vlm_data = {
                            'title': result.title,
                            'description': result.description,
                            'headers': result.headers,
                            'rows': result.rows,
                            'type': 'TabularArtifact',
                            'source_image': img_data['filename'],
                            'page': f"Image {i+1}"
                        }
                        vlm_extracted_data.append(vlm_data)
                    elif isinstance(result, str):
                        # Try to parse JSON string and create proper structure
                        try:
                            parsed_data = json.loads(result)
                            vlm_data = {
                                'title': parsed_data.get('title', f"Extracted from {img_data['filename']}"),
                                'description': parsed_data.get('description', ''),
                                'headers': parsed_data.get('headers', []),
                                'rows': parsed_data.get('rows', []),
                                'type': 'TabularArtifact',
                                'source_image': img_data['filename'],
                                'page': f"Image {i+1}"
                            }
                            vlm_extracted_data.append(vlm_data)
                        except json.JSONDecodeError:
                            # Fallback for non-JSON string
                            vlm_data = {
                                'title': f"Extracted from {img_data['filename']}",
                                'description': result[:300] if len(result) > 300 else result,
                                'headers': [],
                                'rows': [],
                                'type': 'TabularArtifact',
                                'source_image': img_data['filename'],
                                'page': f"Image {i+1}",
                                'raw_response': result
                            }
                            vlm_extracted_data.append(vlm_data)
                    
                    # Update progress bar after each image
                    if progress_bar:
                        progress_bar.update(1)
                        
                except Exception as img_error:
                    # Still update progress bar even if image processing fails
                    if progress_bar:
                        progress_bar.update(1)
                    pass  # Silently skip problematic images
        
        return vlm_extracted_data

    def _safe_sheet_name(self, raw_title: str) -> str:
        """
        Create a safe Excel sheet name from a raw title.
        
        Ensures the sheet name is valid for Excel by removing invalid characters,
        handling length limits, and avoiding duplicates.
        """
        import re
        
        # Excel invalid characters
        invalid_chars = r'[:\\/*?\[\]]'
        max_length = 31
        
        name = (raw_title or "Untitled").strip()
        name = re.sub(invalid_chars, "_", name)
        name = re.sub(r"\s+", " ", name)
        name = name[:max_length] if name else "Sheet"
        
        return name

    def _generate_markdown_output(self, document_data: Dict, images_data: List, output_dir: Path, vlm_extracted_data: List = None):
        """Generate markdown output."""
        markdown_content = []
        
        if document_data['metadata']['title']:
            markdown_content.append(f"# {document_data['metadata']['title']}")
        
        for element in document_data['elements']:
            if element['type'] == 'paragraph':
                if element['is_heading']:
                    level = element['level']
                    markdown_content.append(f"{'#' * level} {element['text']}")
                else:
                    markdown_content.append(element['text'])
            elif element['type'] == 'table':
                if element['markdown']:
                    markdown_content.append(f"\n## Table {element['index'] + 1}")
                    markdown_content.append(element['markdown'])
        
        if vlm_extracted_data:
            for i, vlm_table in enumerate(vlm_extracted_data):
                if vlm_table['rows']:
                    markdown_content.append(f"\n## {vlm_table['title']}")
                    if vlm_table['description']:
                        markdown_content.append(f"*{vlm_table['description']}*")
                    
                    if vlm_table['headers'] and vlm_table['rows']:
                        vlm_markdown = render_markdown_table(vlm_table['headers'], vlm_table['rows'])
                        markdown_content.append(vlm_markdown)
        else:
            for img in images_data:
                relative_path = f"images/{img['filename']}"
                markdown_content.append(f"\n![{img['filename']}]({relative_path})")
        
        write_markdown(markdown_content, str(output_dir), "document.md")

    def _generate_html_output(self, document_data: Dict, images_data: List, output_dir: Path, vlm_extracted_data: List = None):
        """Generate HTML output."""
        html_content = []
        
        if document_data['metadata']['title']:
            html_content.append(f"<h1>{document_data['metadata']['title']}</h1>")
        
        for element in document_data['elements']:
            if element['type'] == 'paragraph':
                if element['is_heading']:
                    level = element['level']
                    html_content.append(f"<h{level}>{element['text']}</h{level}>")
                else:
                    html_content.append(f"<p>{element['text']}</p>")
            elif element['type'] == 'table':
                if element['data']:
                    html_content.append(f"<h2>Table {element['index'] + 1}</h2>")
                    html_table = self._generate_html_table(element['data'])
                    html_content.append(html_table)
        
        if vlm_extracted_data:
            for i, vlm_table in enumerate(vlm_extracted_data):
                if vlm_table['rows']:
                    html_content.append(f"<h2>{vlm_table['title']}</h2>")
                    if vlm_table['description']:
                        html_content.append(f"<p><em>{vlm_table['description']}</em></p>")
                    
                    if vlm_table['headers'] and vlm_table['rows']:
                        table_data = [vlm_table['headers']] + vlm_table['rows']
                        vlm_html_table = self._generate_html_table(table_data)
                        html_content.append(vlm_html_table)
        else:
            for img in images_data:
                relative_path = f"images/{img['filename']}"
                html_content.append(f'<img src="{relative_path}" alt="{img["filename"]}" />')
        
        write_html(html_content, str(output_dir), "document.html")

    def _generate_excel_output(self, tables_data: List, output_dir: Path):
        """Generate Excel output with all tables and Table of Contents."""
        if not tables_data:
            print("⚠️  No tables found to export to Excel")
            return
        
        if not EXCEL_AVAILABLE:
            print("⚠️  Excel export requires pandas and openpyxl: Missing dependencies")
            print("Install with: pip install pandas openpyxl")
            return
        
        try:
            wb = Workbook()
            wb.remove(wb.active)
            
            HEADER_FILL = PatternFill(fill_type="solid", start_color="FF2E7D32", end_color="FF2E7D32")
            HEADER_FONT = Font(color="FFFFFFFF", bold=True)
            HEADER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
            
            toc_data = []
            sheet_index = 1
            sheet_mapping = {}
            
            for i, table in enumerate(tables_data):
                if table['data']:
                    table_title = table.get('title', f"Table {i+1}")
                    sheet_name = self._safe_sheet_name(table_title)
                    ws = wb.create_sheet(title=sheet_name)
                    
                    for row_idx, row_data in enumerate(table['data']):
                        for col_idx, cell_value in enumerate(row_data):
                            ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell_value)
                    
                    if table['data']:
                        ncols = len(table['data'][0]) if table['data'] else 0
                        for col_idx in range(1, ncols + 1):
                            cell = ws.cell(row=1, column=col_idx)
                            cell.fill = HEADER_FILL
                            cell.font = HEADER_FONT
                            cell.alignment = HEADER_ALIGN
                        ws.freeze_panes = "A2"
                    
                    for column in ws.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws.column_dimensions[column_letter].width = adjusted_width
                    
                    toc_data.append([
                        sheet_index,
                        table_title,
                        f"Original table from document",
                        len(table['data']),
                        len(table['data'][0]) if table['data'] else 0,
                        "Document"
                    ])
                    sheet_mapping[table_title] = sheet_name
                    sheet_index += 1
            
            if toc_data:
                toc_ws = wb.create_sheet(title="Table_of_Contents", index=0)
                
                toc_headers = ["Sheet #", "Table Name", "Description", "Rows", "Columns", "Source"]
                for col_idx, header in enumerate(toc_headers):
                    cell = toc_ws.cell(row=1, column=col_idx + 1, value=header)
                    cell.fill = HEADER_FILL
                    cell.font = HEADER_FONT
                    cell.alignment = HEADER_ALIGN
                
                for row_idx, row_data in enumerate(toc_data):
                    for col_idx, cell_value in enumerate(row_data):
                        cell = toc_ws.cell(row=row_idx + 2, column=col_idx + 1, value=cell_value)
                        
                        if col_idx == 1 and cell_value in sheet_mapping:
                            sheet_name = sheet_mapping[cell_value]
                            
                            if ' ' in sheet_name or any(char in sheet_name for char in ['[', ']', '*', '?', ':', '\\', '/']):
                                hyperlink_ref = f"#'{sheet_name}'!A1"
                            else:
                                hyperlink_ref = f"#{sheet_name}!A1"
                            
                            cell.hyperlink = Hyperlink(ref=hyperlink_ref, target=hyperlink_ref)
                            cell.font = Font(color="0000FF", underline="single")
                        
                        if col_idx == 2:
                            cell.alignment = Alignment(wrap_text=True, vertical="top")
                
                toc_ws.column_dimensions['A'].width = 10
                toc_ws.column_dimensions['B'].width = 30
                toc_ws.column_dimensions['C'].width = 60
                toc_ws.column_dimensions['D'].width = 10
                toc_ws.column_dimensions['E'].width = 10
                toc_ws.column_dimensions['F'].width = 15
                
                for row_idx in range(2, len(toc_data) + 2):
                    toc_ws.row_dimensions[row_idx].height = 30
            
            excel_path = output_dir / "tables.xlsx"
            wb.save(excel_path)
            
        except Exception as e:
            print(f"❌ Error creating Excel file: {e}")

    def _generate_excel_output_with_vlm(self, tables_data: List, vlm_extracted_data: List, output_dir: Path):
        """Generate Excel output with both original tables and VLM extracted data, including table of contents."""
        if not tables_data and not vlm_extracted_data:
            print("⚠️  No tables found to export to Excel")
            return

        if not EXCEL_AVAILABLE:
            print("⚠️  Excel export requires pandas and openpyxl: Missing dependencies")
            print("Install with: pip install pandas openpyxl")
            return
        
        try:
            wb = Workbook()
            wb.remove(wb.active)


            # Define styling constants (matching PDF parser)
            HEADER_FILL = PatternFill(fill_type="solid", start_color="FF2E7D32", end_color="FF2E7D32")  # Green
            HEADER_FONT = Font(color="FFFFFFFF", bold=True)
            HEADER_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)

            toc_data = []
            sheet_index = 1
            sheet_mapping = {}

            for i, table in enumerate(tables_data):
                if table['data']:
                    table_title = table.get('title', f"Table {i+1}")
                    sheet_name = self._safe_sheet_name(table_title)
                    ws = wb.create_sheet(title=sheet_name)
                    
                    for row_idx, row_data in enumerate(table['data']):
                        for col_idx, cell_value in enumerate(row_data):
                            ws.cell(row=row_idx + 1, column=col_idx + 1, value=cell_value)

                    if table['data']:
                        ncols = len(table['data'][0]) if table['data'] else 0
                        for col_idx in range(1, ncols + 1):
                            cell = ws.cell(row=1, column=col_idx)
                            cell.fill = HEADER_FILL
                            cell.font = HEADER_FONT
                            cell.alignment = HEADER_ALIGN
                        ws.freeze_panes = "A2"

                    for column in ws.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws.column_dimensions[column_letter].width = adjusted_width

                    toc_data.append([
                        sheet_index,
                        table_title,
                        f"Original table from document",
                        len(table['data']),
                        len(table['data'][0]) if table['data'] else 0,
                        "Document"
                    ])
                    sheet_mapping[table_title] = sheet_name
                    sheet_index += 1

            for i, vlm_table in enumerate(vlm_extracted_data):
                if vlm_table['rows']:
                    table_title = vlm_table['title']
                    sheet_name = self._safe_sheet_name(table_title)
                    ws = wb.create_sheet(title=sheet_name)
                    
                    for col_idx, header in enumerate(vlm_table['headers']):
                        cell = ws.cell(row=1, column=col_idx + 1, value=header)
                        cell.fill = HEADER_FILL
                        cell.font = HEADER_FONT
                        cell.alignment = HEADER_ALIGN
                    
                    for row_idx, row_data in enumerate(vlm_table['rows']):
                        for col_idx, cell_value in enumerate(row_data):
                            ws.cell(row=row_idx + 2, column=col_idx + 1, value=cell_value)

                    ws.freeze_panes = "A2"

                    for column in ws.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws.column_dimensions[column_letter].width = adjusted_width

                    toc_data.append([
                        sheet_index,
                        table_title,
                        vlm_table['description'],
                        len(vlm_table['rows']),
                        len(vlm_table['headers']),
                        "VLM Extracted"
                    ])
                    sheet_mapping[table_title] = sheet_name
                    sheet_index += 1

            if toc_data:
                toc_ws = wb.create_sheet(title="Table_of_Contents", index=0)
                
                toc_headers = ["Sheet #", "Table Name", "Description", "Rows", "Columns", "Source"]
                for col_idx, header in enumerate(toc_headers):
                    cell = toc_ws.cell(row=1, column=col_idx + 1, value=header)
                    cell.fill = HEADER_FILL
                    cell.font = HEADER_FONT
                    cell.alignment = HEADER_ALIGN
                
                for row_idx, row_data in enumerate(toc_data):
                    for col_idx, cell_value in enumerate(row_data):
                        cell = toc_ws.cell(row=row_idx + 2, column=col_idx + 1, value=cell_value)
                        
                        if col_idx == 1 and cell_value in sheet_mapping:
                            sheet_name = sheet_mapping[cell_value]
                            
                            if ' ' in sheet_name or any(char in sheet_name for char in ['[', ']', '*', '?', ':', '\\', '/']):
                                hyperlink_ref = f"#'{sheet_name}'!A1"
                            else:
                                hyperlink_ref = f"#{sheet_name}!A1"
                            
                            cell.hyperlink = Hyperlink(ref=hyperlink_ref, target=hyperlink_ref)
                            cell.font = Font(color="0000FF", underline="single")
                        
                        if col_idx == 2:
                            cell.alignment = Alignment(wrap_text=True, vertical="top")
                
                toc_ws.column_dimensions['A'].width = 10
                toc_ws.column_dimensions['B'].width = 30
                toc_ws.column_dimensions['C'].width = 60
                toc_ws.column_dimensions['D'].width = 10
                toc_ws.column_dimensions['E'].width = 10
                toc_ws.column_dimensions['F'].width = 15
                
                for row_idx in range(2, len(toc_data) + 2):
                    toc_ws.row_dimensions[row_idx].height = 30

            excel_path = output_dir / "tables.xlsx"
            wb.save(excel_path)

        except Exception as e:
            print(f"❌ Error creating Excel file: {e}")


    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if style_name.startswith('Heading'):
            try:
                return int(style_name.split()[-1])
            except:
                return 1
        return 0

    def _extract_formatting(self, paragraph: Paragraph) -> Dict[str, Any]:
        """Extract formatting information from paragraph."""
        formatting = {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_size': None,
            'font_name': None
        }
        
        try:
            for run in paragraph.runs:
                if run.bold:
                    formatting['bold'] = True
                if run.italic:
                    formatting['italic'] = True
                if run.underline:
                    formatting['underline'] = True
                if run.font.size:
                    formatting['font_size'] = run.font.size.pt
                if run.font.name:
                    formatting['font_name'] = run.font.name
        except:
            pass
        
        return formatting

    def _generate_html_table(self, table_data: List[List[str]]) -> str:
        """Generate HTML table from table data."""
        if not table_data:
            return ""
        
        html = ["<table border='1'>"]
        
        for row_idx, row in enumerate(table_data):
            html.append("<tr>")
            for cell in row:
                tag = "th" if row_idx == 0 else "td"
                html.append(f"<{tag}>{cell}</{tag}>")
            html.append("</tr>")
        
        html.append("</table>")
        return '\n'.join(html)
