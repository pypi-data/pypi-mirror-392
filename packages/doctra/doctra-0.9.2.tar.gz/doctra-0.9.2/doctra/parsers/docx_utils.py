"""
DOCX Processing Utilities

Helper functions for DOCX document processing and content extraction.
"""

from __future__ import annotations
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    DocumentType = None
    Table = None
    Paragraph = None
    CT_Tbl = None
    CT_P = None
    qn = None


def extract_document_metadata(doc: DocumentType) -> Dict[str, Any]:
    """
    Extract metadata from a DOCX document.
    
    :param doc: The DOCX document object
    :return: Dictionary containing document metadata
    """
    metadata = {
        'title': doc.core_properties.title or '',
        'author': doc.core_properties.author or '',
        'subject': doc.core_properties.subject or '',
        'keywords': doc.core_properties.keywords or '',
        'comments': doc.core_properties.comments or '',
        'created': str(doc.core_properties.created) if doc.core_properties.created else '',
        'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
        'last_modified_by': doc.core_properties.last_modified_by or '',
        'revision': doc.core_properties.revision or 0,
        'version': doc.core_properties.version or '',
    }
    
    return metadata


def extract_paragraph_structure(paragraph: Paragraph) -> Dict[str, Any]:
    """
    Extract detailed structure from a paragraph.
    
    :param paragraph: The paragraph object
    :return: Dictionary containing paragraph structure
    """
    structure = {
        'text': paragraph.text,
        'style_name': paragraph.style.name if paragraph.style else 'Normal',
        'is_heading': False,
        'heading_level': 0,
        'alignment': str(paragraph.alignment) if paragraph.alignment else 'left',
        'runs': [],
        'formatting': {}
    }
    
    if paragraph.style and paragraph.style.name.startswith('Heading'):
        structure['is_heading'] = True
        try:
            structure['heading_level'] = int(paragraph.style.name.split()[-1])
        except:
            structure['heading_level'] = 1
    
    for run in paragraph.runs:
        run_data = {
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font.name else None,
            'font_size': run.font.size.pt if run.font.size else None,
            'font_color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
        }
        structure['runs'].append(run_data)
    
    structure['formatting'] = {
        'bold': any(run.bold for run in paragraph.runs),
        'italic': any(run.italic for run in paragraph.runs),
        'underline': any(run.underline for run in paragraph.runs),
    }
    
    return structure


def extract_table_structure(table: Table) -> Dict[str, Any]:
    """
    Extract detailed structure from a table.
    
    :param table: The table object
    :return: Dictionary containing table structure
    """
    table_data = {
        'rows': len(table.rows),
        'columns': len(table.columns),
        'data': [],
        'headers': [],
        'has_header_row': False,
        'merged_cells': [],
        'table_style': table.style.name if table.style else 'Table Grid'
    }
    
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for col_idx, cell in enumerate(row.cells):
            cell_data = {
                'text': cell.text.strip(),
                'row_span': 1,
                'col_span': 1,
                'is_merged': False
            }
            
            if cell._tc in [tc for tc in row._tr.tc_lst]:
                pass
            
            row_data.append(cell_data)
        table_data['data'].append(row_data)
    
    if table_data['data']:
        table_data['headers'] = table_data['data'][0]
        table_data['has_header_row'] = True
    
    return table_data


def extract_list_structure(paragraph: Paragraph) -> Optional[Dict[str, Any]]:
    """
    Extract list structure from a paragraph.
    
    :param paragraph: The paragraph object
    :return: Dictionary containing list structure or None if not a list
    """
    if not paragraph.style or not paragraph.style.name.startswith('List'):
        return None
    
    list_data = {
        'text': paragraph.text,
        'style': paragraph.style.name,
        'level': 0,
        'is_ordered': 'Number' in paragraph.style.name,
        'is_bulleted': 'Bullet' in paragraph.style.name,
    }
    
    level_match = re.search(r'(\d+)', paragraph.style.name)
    if level_match:
        list_data['level'] = int(level_match.group(1))
    
    return list_data


def clean_text(text: str) -> str:
    """
    Clean and normalize text content.
    
    :param text: Raw text to clean
    :return: Cleaned text
    """
    text = re.sub(r'\s+', ' ', text)
    
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    return text.strip()


def extract_document_outline(doc: DocumentType) -> List[Dict[str, Any]]:
    """
    Extract document outline (headings hierarchy).
    
    :param doc: The DOCX document object
    :return: List of heading dictionaries
    """
    outline = []
    
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split()[-1])
                outline.append({
                    'level': level,
                    'text': para.text,
                    'style': para.style.name
                })
            except:
                pass
    
    return outline


def detect_document_sections(doc: DocumentType) -> List[Dict[str, Any]]:
    """
    Detect document sections based on headings and content.
    
    :param doc: The DOCX document object
    :return: List of section dictionaries
    """
    sections = []
    current_section = None
    
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith('Heading'):
            # Save previous section
            if current_section:
                sections.append(current_section)
            
            # Start new section
            try:
                level = int(para.style.name.split()[-1])
                current_section = {
                    'title': para.text,
                    'level': level,
                    'content': [],
                    'start_paragraph': len(sections)
                }
            except:
                pass
        elif current_section and para.text.strip():
            current_section['content'].append(para.text)
    
    if current_section:
        sections.append(current_section)
    
    return sections


def extract_hyperlinks(doc: DocumentType) -> List[Dict[str, str]]:
    """
    Extract hyperlinks from the document.
    
    :param doc: The DOCX document object
    :return: List of hyperlink dictionaries
    """
    hyperlinks = []
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xml:
                # Look for hyperlink elements
                hyperlink_match = re.search(r'<w:hyperlink[^>]*r:id="([^"]*)"', run._element.xml)
                if hyperlink_match:
                    rel_id = hyperlink_match.group(1)
                    try:
                        rel = doc.part.rels[rel_id]
                        hyperlinks.append({
                            'text': run.text,
                            'url': rel.target_ref,
                            'type': 'external' if rel.target_ref.startswith('http') else 'internal'
                        })
                    except:
                        pass
    
    return hyperlinks


def extract_images_metadata(doc: DocumentType) -> List[Dict[str, Any]]:
    """
    Extract metadata about embedded images.
    
    :param doc: The DOCX document object
    :return: List of image metadata dictionaries
    """
    images = []
    
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_info = {
                    'filename': rel.target_ref,
                    'content_type': rel.target_content_type,
                    'size': len(rel.target_part.blob),
                    'extension': rel.target_ref.split('.')[-1].lower(),
                    'relationship_id': rel.rId
                }
                images.append(image_info)
    except Exception as e:
        print(f"Warning: Could not extract image metadata: {e}")
    
    return images


def validate_docx_file(file_path: str) -> Tuple[bool, str]:
    """
    Validate if a file is a valid DOCX file.
    
    :param file_path: Path to the file to validate
    :return: Tuple of (is_valid, error_message)
    """
    try:
        if not Path(file_path).exists():
            return False, "File does not exist"
        
        if not file_path.lower().endswith('.docx'):
            return False, "File is not a DOCX file"
        
        # Try to open the document
        doc = Document(file_path)
        
        # Basic validation - check if we can access core properties
        _ = doc.core_properties.title
        
        return True, "Valid DOCX file"
        
    except Exception as e:
        return False, f"Invalid DOCX file: {str(e)}"


def get_document_statistics(doc: DocumentType) -> Dict[str, int]:
    """
    Get basic statistics about the document.
    
    :param doc: The DOCX document object
    :return: Dictionary containing document statistics
    """
    stats = {
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'headings': 0,
        'lists': 0,
        'images': 0,
        'characters': 0,
        'words': 0
    }
    
    for para in doc.paragraphs:
        if para.style:
            if para.style.name.startswith('Heading'):
                stats['headings'] += 1
            elif para.style.name.startswith('List'):
                stats['lists'] += 1
        
        stats['characters'] += len(para.text)
        stats['words'] += len(para.text.split())
    
    try:
        stats['images'] = len([rel for rel in doc.part.rels.values() if "image" in rel.target_ref])
    except:
        pass
    
    return stats
