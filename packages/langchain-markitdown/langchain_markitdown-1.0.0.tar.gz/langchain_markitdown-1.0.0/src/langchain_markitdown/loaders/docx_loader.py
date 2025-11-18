from __future__ import annotations

from typing import TYPE_CHECKING, Any, Dict, List, Optional, Tuple

from langchain_core.documents import Document
from langchain_text_splitters import MarkdownHeaderTextSplitter

from ..base_loader import BaseMarkitdownLoader

if TYPE_CHECKING:
    from markitdown import MarkItDown


class DocxLoader(BaseMarkitdownLoader):
    """DOCX loader with optional page/header splitting support."""

    DEFAULT_HEADERS: List[Tuple[str, str]] = [
        ("#", "Header 1"),
        ("##", "Header 2"),
        ("###", "Header 3"),
    ]

    def __init__(
        self,
        file_path: str,
        split_by_page: bool = False,
        *,
        verbose: bool = False,
        converter: Optional["MarkItDown"] = None,
    ):
        super().__init__(file_path, verbose=verbose, converter=converter)
        self.split_by_page = split_by_page

    def load(
        self,
        headers_to_split_on: Optional[List[Tuple[str, str]]] = None,
    ) -> List[Document]:
        """Load a DOCX file and convert it to LangChain documents.

        If ``split_by_page`` is enabled, the loader attempts to split using the
        per-page information returned by MarkItDown. When ``headers_to_split_on``
        is provided, header-based splitting will be applied to the available
        content regardless of the ``split_by_page`` flag.
        """
        metadata: Dict[str, Any] = {
            "source": self.file_path,
            "success": False,
            "conversion_success": False,
        }

        try:
            metadata.update(self._file_metadata())
            result = self._convert_to_markdown()
            metadata["success"] = True
            metadata["conversion_success"] = True
            metadata.update(self._extract_conversion_metadata(result))

            metadata.update(self._extract_docx_metadata())

            headers = headers_to_split_on or self.DEFAULT_HEADERS
            text_content = self._get_text_content(result)

            if self.split_by_page:
                return self._split_pages(result, metadata, headers)

            if headers_to_split_on:
                splitter = self._build_splitter(headers)
                return self._split_content_with_headers(
                    splitter, text_content, metadata, content_type="document_section"
                )

            metadata["content_type"] = "document_full"
            return [Document(page_content=text_content, metadata=metadata)]
        except FileNotFoundError as exc:
            metadata["error"] = "File not found."
            raise ValueError(
                f"Markitdown conversion failed for {self.file_path}: File not found",
            ) from exc
        except Exception as exc:
            metadata["error"] = str(exc)
            raise ValueError(
                f"Failed to load and convert DOCX file: {exc}",
            ) from exc

    def _extract_docx_metadata(self) -> Dict[str, Any]:
        """Collect DOCX-level metadata via python-docx when available."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return {"metadata_extraction_error": "python-docx not installed"}

        try:
            doc = DocxDocument(self.file_path)
            core_props = doc.core_properties
        except Exception as exc:  # pragma: no cover - edge of python-docx
            return {"metadata_extraction_error": str(exc)}

        metadata: Dict[str, Any] = {}

        def set_if_present(attr_name: str, key: Optional[str] = None) -> None:
            value = getattr(core_props, attr_name, None)
            if value:
                metadata[key or attr_name] = str(value)

        for attr in [
            "author",
            "title",
            "subject",
            "keywords",
            "created",
            "modified",
            "last_modified_by",
            "revision",
            "category",
        ]:
            set_if_present(attr)

        return metadata

    def _split_pages(
        self,
        result: Any,
        metadata: Dict[str, Any],
        headers_to_split_on: List[Tuple[str, str]],
    ) -> List[Document]:
        pages = getattr(result, "pages", None)
        splitter = self._build_splitter(headers_to_split_on)
        documents: List[Document] = []

        if isinstance(pages, list) and pages:
            for page_num, page_content in enumerate(pages, start=1):
                page_metadata = metadata.copy()
                page_metadata["page_number"] = page_num
                page_metadata["content_type"] = "document_page"
                documents.extend(
                    self._split_content_with_headers(
                        splitter,
                        page_content,
                        page_metadata,
                        content_type="document_page",
                    ),
                )
        else:
            documents = self._split_content_with_headers(
                splitter,
                self._get_text_content(result),
                {**metadata, "content_type": "document_page"},
            )

        return documents

    def _build_splitter(
        self,
        headers_to_split_on: List[Tuple[str, str]],
    ) -> MarkdownHeaderTextSplitter:
        return MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            return_each_line=True,
        )

    def _split_content_with_headers(
        self,
        splitter: MarkdownHeaderTextSplitter,
        content: str,
        metadata: Dict[str, Any],
        *,
        content_type: str = "document_section",
    ) -> List[Document]:
        splits = splitter.split_text(content)

        if not splits:
            return [Document(page_content=content, metadata=metadata.copy())]

        updated_documents: List[Document] = []
        for split in splits:
            split.metadata.update({**metadata, "content_type": content_type})
            updated_documents.append(split)

        return updated_documents
