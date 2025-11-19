"""Functions for generating tables and overviews for creating reports"""

from docx import Document # type: ignore
from docx.shared import Cm # type: ignore

from ..models import *
from ..visualizations import svg_task_hierarchy

import subprocess, tempfile

__all__ = ('model_to_docx',
           )


#Export to docx report
def create_title_page(plan: Model,doc):
    doc.add_heading(plan.title, 0)
    if not plan.summary is None:
        doc.add_paragraph(plan.summary)
    
def create_actor_pages(plan: Model, doc: Document):
    doc.add_page_break()
    doc.add_heading('Actors', 1)

    for actor in plan.actors:
        doc.add_heading(actor.title,2)

        table = doc.add_table(rows=0, cols=2)
        table.style = doc.styles['Light Grid Accent 1']
        
        cells = table.add_row().cells
        cells[0].text = 'ID'
        cells[1].text = actor.get_id()
        
        cells = table.add_row().cells
        cells[0].text = 'Title'
        cells[1].text = '-' if actor.title is None else actor.title 

        cells = table.add_row().cells
        cells[0].text = 'Type'
        cells[1].text = '-' if actor.type is None else actor.type

        cells = table.add_row().cells
        cells[0].text = 'Description'
        cells[1].text = '-' if actor.description is None else actor.description

        cells = table.add_row().cells
        cells[0].text = 'Affiliations'
        cells[1].text = ', '.join(actor.groups) if actor.groups else '-'

        if not actor.annotations is None:
            for annotation, value in actor.annotations.items():
                cells = table.add_row().cells
                cells[0].text = annotation
                cells[1].text = str(value)


def create_task_pages(plan: Model, doc: Document):
    doc.add_page_break()
    doc.add_heading('Tasks', 1)

    for num, task in enumerate(plan.tasks,1):
        doc.add_heading(f'{num}. {task.title}',2)
        if not task.description is None:
            par = doc.add_paragraph(task.description)
    
        #Show decomposition
        svg = svg_task_hierarchy(num, task)
        with open(f'tmp.png',"w") as output:
            subprocess.run(["cairosvg","-f","png","-"], stdout= output, input=svg, encoding = 'utf-8')
        doc.add_picture('tmp.png',width=Cm(15))

        def collect_rows(prefix,acc,task):
            rows.append((prefix, task))
            for sub_num, sub_task in enumerate(task.tasks,1):
                collect_rows(f'{prefix}.{sub_num}',acc,sub_task)

        rows: list[tuple[str,Task]] = []
        collect_rows(str(num),rows,task)
     
        #Add decomposition as table
        table = doc.add_table(rows=1, cols=2)
        table.style = doc.styles['Light Grid']
        header = table.rows[0].cells
        header[0].text = 'Number'
        header[1].text = 'Title'
        for sub_num, sub_task in rows:
            cells = table.add_row().cells
            cells[0].text = sub_num
            cells[1].text = '-' if sub_task.title is None else sub_task.title

        # Add details for each subtask
        for sub_num, sub_task in rows:
            doc.add_heading(f'{sub_num}. {"-" if sub_task.title is None else sub_task.title}',3)
            table = doc.add_table(rows=0, cols=2)
            table.style = doc.styles['Light Grid Accent 1']
       
            cells = table.add_row().cells
            cells[0].text = 'Number'
            cells[1].text = sub_num

            cells = table.add_row().cells
            cells[0].text = 'Title'
            cells[1].text = '-' if sub_task.title is None else sub_task.title 

            cells = table.add_row().cells
            cells[0].text = 'Description'
            cells[1].text = '-' if sub_task.description is None else sub_task.description

            cells = table.add_row().cells
            cells[0].text = 'Parameter type'
            cells[1].text = '-' if sub_task.parameter_type is None else str(sub_task.parameter_type)

            #List all required information space
            doc.add_heading(f'Required information spaces',4)
            table = doc.add_table(rows=0, cols=4)
            table.style = doc.styles['Light Grid Accent 1']
            
            cells = table.add_row().cells
            cells[0].text = 'Name'
            cells[1].text = 'Type'
            cells[2].text = 'Read/Write'
            cells[3].text = 'Binding'

            for req in sub_task.info_space_requirements:
                cells = table.add_row().cells
                cells[0].text = req.name
                cells[1].text = '-' if req.type is None else str(req.type)

                if req.read and req.write:
                    cells[2].text = 'Both'
                elif req.read:
                    cells[2].text = 'Read'
                elif req.write:
                    cells[2].text = 'Write'


def create_info_space_pages(plan: Model,doc):
    doc.add_page_break()
    doc.add_heading('Information spaces', 1)

    for ifs in plan.info_spaces:
        doc.add_heading('-' if ifs.title is None else ifs.title,2)

        table = doc.add_table(rows=0, cols=2)
        table.style = doc.styles['Light Grid Accent 1']
        
        cells = table.add_row().cells
        cells[0].text = 'ID'
        cells[1].text = ifs.get_id()
        
        cells = table.add_row().cells
        cells[0].text = 'Title'
        cells[1].text = '-' if ifs.title is None else ifs.title 

        cells = table.add_row().cells
        cells[0].text = 'Type'
        cells[1].text = '-' if ifs.type is None else str(ifs.type)

        cells = table.add_row().cells
        cells[0].text = 'Description'
        cells[1].text = '-' if ifs.description is None else ifs.description

        if not ifs.annotations is None:
            for annotation, value in ifs.annotations.items():
                cells = table.add_row().cells
                cells[0].text = annotation
                cells[1].text = str(value)

async def model_to_docx(plan: Model):

    doc = Document()

    create_title_page(plan,doc)
    create_actor_pages(plan,doc)
    create_task_pages(plan,doc)
    create_info_space_pages(plan,doc)

    with tempfile.TemporaryFile() as tmp:
        doc.save(tmp)
        tmp.flush()
        tmp.seek(0)
        content = tmp.read()     
    return content
